"""
Diagram Renderer for IB Report Formatter
Renders flow diagrams as high-quality matplotlib images inserted into Word.

Requires: matplotlib (optional dependency, graceful fallback to text placeholder).
"""

import logging
import os
import tempfile
from typing import Optional, Dict, List, Tuple, TYPE_CHECKING

from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from md_parser import Diagram

logger = logging.getLogger(__name__)

_NAVY = "#003366"
_DARK_GRAY = "#404040"
_MEDIUM_GRAY = "#888888"
_LIGHT_BG = "#F5F5F5"


def _matplotlib_available() -> bool:
    try:
        import matplotlib
        return True
    except ImportError:
        return False


class DiagramRenderer:
    """Renders Diagram objects as matplotlib images in Word documents."""

    FONT = "Malgun Gothic"

    def __init__(self, doc: "DocxDocument", theme_colors: Optional[Dict[str, str]] = None):
        self.doc = doc
        self.colors = theme_colors or {}
        self._navy = self.colors.get("navy", _NAVY)

    def render(self, diagram: "Diagram") -> bool:
        """Render diagram as image. Returns True on success."""
        if not diagram.boxes:
            return False

        if not _matplotlib_available():
            self._placeholder(diagram.title)
            return False

        image_path = self._render_image(diagram)
        if image_path:
            try:
                self._insert(image_path, diagram.title, diagram.notes)
                return True
            except Exception as e:
                logger.warning("Failed to insert diagram: %s", e)
                self._placeholder(diagram.title)
                return False
            finally:
                try:
                    os.unlink(image_path)
                except OSError:
                    pass
        else:
            self._placeholder(diagram.title)
            return False

    def _render_image(self, diagram: "Diagram") -> Optional[str]:
        """Render the diagram to a temporary PNG."""
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            from matplotlib.patches import FancyBboxPatch

            # ── Font setup ──────────────────────────────────────────────
            korean_font = self._find_korean_font()
            if korean_font:
                plt.rcParams["font.family"] = korean_font
            plt.rcParams["axes.unicode_minus"] = False

            boxes = diagram.boxes
            arrows = diagram.arrows

            # ── Layout calculation ──────────────────────────────────────
            xs = [b.pos[0] for b in boxes]
            ys = [b.pos[1] for b in boxes]
            x_min, x_max = min(xs), max(xs)
            y_min, y_max = min(ys), max(ys)
            x_range = max(x_max - x_min, 1)
            y_range = max(y_max - y_min, 1)

            # Figure size — large for readability
            fig_w = max(16, x_range * 4.0 + 5)
            fig_h = max(10, y_range * 3.5 + 4)
            fig, ax = plt.subplots(figsize=(fig_w, fig_h))
            fig.patch.set_facecolor("white")

            margin = 1.0
            ax.set_xlim(x_min - margin, x_max + margin)
            ax.set_ylim(y_min - margin * 0.8, y_max + margin * 0.8)
            ax.set_aspect("equal")
            ax.axis("off")

            # ── Box dimensions ──────────────────────────────────────────
            # Dynamic sizing based on spacing
            min_x_gap = float("inf")
            min_y_gap = float("inf")
            for i, b1 in enumerate(boxes):
                for b2 in boxes[i + 1:]:
                    dx = abs(b1.pos[0] - b2.pos[0])
                    dy = abs(b1.pos[1] - b2.pos[1])
                    if dx > 0.01:
                        min_x_gap = min(min_x_gap, dx)
                    if dy > 0.01:
                        min_y_gap = min(min_y_gap, dy)

            box_w = min(min_x_gap * 0.65, 2.2) if min_x_gap < float("inf") else 2.2
            box_h = min(min_y_gap * 0.40, 0.9) if min_y_gap < float("inf") else 0.9
            box_w = max(box_w, 1.5)
            box_h = max(box_h, 0.65)

            # ── Colors per style ────────────────────────────────────────
            navy = self._navy
            style_map = {
                "default": {"fill": navy, "edge": navy, "text": "white"},
                "highlight": {"fill": "#4472C4", "edge": "#4472C4", "text": "white"},
                "subtle": {"fill": _LIGHT_BG, "edge": "#BBBBBB", "text": _DARK_GRAY},
            }

            # ── Build box lookup ────────────────────────────────────────
            box_pos = {}
            for b in boxes:
                box_pos[b.id] = (b.pos[0], b.pos[1])

            # ── Draw arrows first (behind boxes) ───────────────────────
            for arrow in arrows:
                if arrow.from_id not in box_pos or arrow.to_id not in box_pos:
                    continue
                x1, y1 = box_pos[arrow.from_id]
                x2, y2 = box_pos[arrow.to_id]

                # Edge intersection points
                sx, sy = self._edge_pt(x1, y1, x2, y2, box_w, box_h)
                ex, ey = self._edge_pt(x2, y2, x1, y1, box_w, box_h)

                # Arrow style
                ls = "--" if arrow.style == "dashed" else "-"
                arrowstyle = "<->" if arrow.style == "both" else "->"
                color = _MEDIUM_GRAY

                ax.annotate(
                    "", xy=(ex, ey), xytext=(sx, sy),
                    arrowprops=dict(
                        arrowstyle=arrowstyle, color=color,
                        lw=2.0, linestyle=ls,
                        shrinkA=0, shrinkB=0,
                        connectionstyle="arc3,rad=0.0",
                    ),
                    zorder=1,
                )

                # Arrow label
                if arrow.label:
                    mx = (sx + ex) / 2
                    my = (sy + ey) / 2
                    # Perpendicular offset
                    dx, dy = ex - sx, ey - sy
                    dist = max((dx ** 2 + dy ** 2) ** 0.5, 0.001)
                    off_x = -dy / dist * 0.28
                    off_y = dx / dist * 0.28
                    ax.text(
                        mx + off_x, my + off_y,
                        arrow.label,
                        ha="center", va="center",
                        fontsize=16, color=_DARK_GRAY,
                        linespacing=1.3,
                        bbox=dict(
                            boxstyle="round,pad=0.25",
                            facecolor="white", edgecolor="#DDDDDD",
                            alpha=0.92, linewidth=0.5,
                        ),
                        zorder=4,
                    )

            # ── Draw boxes ──────────────────────────────────────────────
            for b in boxes:
                cx, cy = b.pos[0], b.pos[1]
                s = style_map.get(b.style, style_map["default"])

                # Rounded rectangle
                rect = FancyBboxPatch(
                    (cx - box_w / 2, cy - box_h / 2), box_w, box_h,
                    boxstyle="round,pad=0.06",
                    facecolor=s["fill"], edgecolor=s["edge"],
                    linewidth=1.8, zorder=2,
                )
                ax.add_patch(rect)

                # Drop shadow (subtle)
                shadow = FancyBboxPatch(
                    (cx - box_w / 2 + 0.02, cy - box_h / 2 - 0.02),
                    box_w, box_h,
                    boxstyle="round,pad=0.06",
                    facecolor="#00000010", edgecolor="none",
                    linewidth=0, zorder=1,
                )
                ax.add_patch(shadow)

                # Label text — split lines, first line bold & larger
                lines = b.label.split("\n")
                total_lines = len(lines)
                line_height = 0.22  # spacing between lines

                # Vertical start position to center the text block
                start_y = cy + (total_lines - 1) * line_height / 2

                for i, line_text in enumerate(lines):
                    ly = start_y - i * line_height
                    weight = "bold" if i == 0 else "normal"
                    size = 20 if i == 0 else 16
                    ax.text(
                        cx, ly, line_text,
                        ha="center", va="center",
                        fontsize=size, fontweight=weight,
                        color=s["text"], zorder=3,
                    )

            # ── Title ───────────────────────────────────────────────────
            if diagram.title:
                ax.set_title(
                    diagram.title,
                    fontsize=24, fontweight="bold",
                    color=_DARK_GRAY, pad=25,
                )

            # ── Notes at bottom ─────────────────────────────────────────
            if diagram.notes:
                note_text = "    ".join(f"※ {n}" for n in diagram.notes)
                fig.text(
                    0.5, 0.02, note_text,
                    ha="center", va="bottom",
                    fontsize=14, color=_MEDIUM_GRAY,
                    style="italic",
                )

            plt.tight_layout(rect=[0, 0.05, 1, 0.95])

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False, mode="wb") as f:
                temp_path = f.name

            fig.savefig(temp_path, dpi=200, bbox_inches="tight",
                        facecolor="white", edgecolor="none")
            plt.close(fig)
            return temp_path

        except Exception as e:
            logger.warning("Diagram rendering failed: %s", e)
            import traceback
            traceback.print_exc()
            return None

    # ── Helpers ──────────────────────────────────────────────────────────

    @staticmethod
    def _edge_pt(cx, cy, tx, ty, bw, bh):
        """Point on box edge where line from center toward target exits."""
        dx = tx - cx
        dy = ty - cy
        if dx == 0 and dy == 0:
            return cx, cy
        hw, hh = bw / 2, bh / 2
        if abs(dx) * hh > abs(dy) * hw:
            t = hw / abs(dx)
        else:
            t = hh / abs(dy)
        return cx + dx * t, cy + dy * t

    @staticmethod
    def _find_korean_font() -> Optional[str]:
        try:
            import matplotlib.font_manager as fm
            for name in ["Malgun Gothic", "NanumGothic", "AppleGothic"]:
                if fm.findfont(name, fallback_to_default=False):
                    return name
        except Exception:
            pass
        return None

    def _insert(self, file_path: str, title: str, notes: List[str]):
        """Insert diagram image centered with caption and notes."""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(file_path, width=Inches(6.0))

        # Caption
        if title:
            from ib_renderer import STYLE, FontStyler
            cap = self.doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(title)
            FontStyler.apply_run_style(
                r, font_size=STYLE.SMALL_SIZE, italic=True, color=STYLE.DARK_GRAY
            )

        self.doc.add_paragraph()

    def _placeholder(self, title: str):
        """Text placeholder when matplotlib is unavailable."""
        from ib_renderer import STYLE, FontStyler
        p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)
        run = p.add_run(f"[Diagram: {title}]")
        FontStyler.apply_run_style(run, italic=True, color=STYLE.RED)
