"""Tests for word_parser module."""

import base64
import io
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches

from md_parser import ElementType
from md_to_word import IBReportConverter
from word_parser import NumberingTracker, StyleDetector, parse_word_file

PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z8eQAAAAASUVORK5CYII="
)


def _save_doc(doc: Document, output_path: Path) -> Path:
    """Save a temporary docx fixture and return its path."""
    doc.save(str(output_path))
    return output_path


def _write_png(output_path: Path) -> Path:
    """Write a tiny PNG fixture used for image extraction tests."""
    output_path.write_bytes(PNG_BYTES)
    return output_path


def _inject_custom_properties(output_path: Path, properties) -> Path:
    """Inject a minimal docProps/custom.xml part for metadata extraction tests."""
    custom_props_xml = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
    ]

    for index, (name, value) in enumerate(properties.items(), start=2):
        custom_props_xml.append(
            f'<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="{index}" name="{name}">'
            f"<vt:lpwstr>{value}</vt:lpwstr>"
            "</property>"
        )

    custom_props_xml.append("</Properties>")

    docx_bytes = output_path.read_bytes()
    buffer = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as source_zip:
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as target_zip:
            for info in source_zip.infolist():
                target_zip.writestr(info, source_zip.read(info.filename))
            target_zip.writestr("docProps/custom.xml", "".join(custom_props_xml))

    output_path.write_bytes(buffer.getvalue())
    return output_path


def _read_zip_member(docx_path: Path, member_name: str) -> str:
    """Read a UTF-8-ish XML member from a DOCX zip."""
    with zipfile.ZipFile(docx_path, "r") as archive:
        return archive.read(member_name).decode("utf-8")


def _rewrite_document_xml(docx_path: Path, transform) -> Path:
    """Rewrite word/document.xml in a DOCX using a string transform."""
    buffer = io.BytesIO()
    with zipfile.ZipFile(docx_path, "r") as source_zip:
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as target_zip:
            for info in source_zip.infolist():
                data = source_zip.read(info.filename)
                if info.filename == "word/document.xml":
                    data = transform(data.decode("utf-8")).encode("utf-8")
                target_zip.writestr(info, data)
    docx_path.write_bytes(buffer.getvalue())
    return docx_path


def test_parse_preserves_paragraph_table_paragraph_order(tmp_path):
    """Paragraph-table-paragraph order should survive Word parsing."""
    doc = Document()
    doc.add_paragraph("Opening paragraph")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "100"

    doc.add_paragraph("Closing paragraph")
    docx_path = _save_doc(doc, tmp_path / "ordered.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    assert [element.element_type for element in model.elements[:3]] == [
        ElementType.PARAGRAPH,
        ElementType.TABLE,
        ElementType.PARAGRAPH,
    ]
    assert model.elements[0].raw_text == "Opening paragraph"
    assert model.elements[2].raw_text == "Closing paragraph"


def test_parse_callout_table_in_document_order(tmp_path):
    """A single-cell callout table should become one blockquote in place."""
    doc = Document()
    doc.add_paragraph("Before callout")

    callout = doc.add_table(rows=1, cols=1)
    callout.cell(0, 0).text = "요약: 핵심 포인트"

    doc.add_paragraph("After callout")
    docx_path = _save_doc(doc, tmp_path / "callout.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    assert [element.element_type for element in model.elements[:3]] == [
        ElementType.PARAGRAPH,
        ElementType.BLOCKQUOTE,
        ElementType.PARAGRAPH,
    ]
    assert model.elements[1].content.title == "요약"
    assert model.elements[1].content.text == "핵심 포인트"


def test_parse_image_paragraph_with_extraction(tmp_path):
    """Image-only paragraphs should become image elements with extracted files."""
    image_path = _write_png(tmp_path / "tiny.png")
    image_dir = tmp_path / "images"

    doc = Document()
    doc.add_paragraph("Chart below")
    doc.add_picture(str(image_path))
    docx_path = _save_doc(doc, tmp_path / "images.docx")

    model = parse_word_file(
        str(docx_path),
        extract_images=True,
        image_output_dir=str(image_dir),
    )

    image_elements = [element for element in model.elements if element.element_type == ElementType.IMAGE]
    assert len(image_elements) == 1
    assert image_elements[0].content.path.startswith("image_")
    assert list(image_dir.glob("image_*"))


def test_parse_image_paragraph_uses_word_alt_text(tmp_path):
    """Word drawing description should win over filename-derived alt text."""
    image_path = _write_png(tmp_path / "tiny.png")

    doc = Document()
    doc.add_picture(str(image_path))
    doc.inline_shapes[0]._inline.docPr.set("descr", "Revenue bridge")
    docx_path = _save_doc(doc, tmp_path / "image_alt.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    image_element = next(element for element in model.elements if element.element_type == ElementType.IMAGE)
    assert image_element.content.alt_text == "Revenue bridge"


def test_parse_image_paragraph_embeds_base64_without_extraction(tmp_path):
    """Base64 image data should be preserved even when files are not extracted."""
    image_path = _write_png(tmp_path / "tiny.png")

    doc = Document()
    doc.add_picture(str(image_path))
    docx_path = _save_doc(doc, tmp_path / "image_base64.docx")

    model = parse_word_file(
        str(docx_path),
        extract_images=False,
        embed_images_base64=True,
    )

    image_element = next(element for element in model.elements if element.element_type == ElementType.IMAGE)
    assert image_element.content.path == ""
    assert image_element.content.base64_data is not None
    assert image_element.content.mime_type == "image/png"


def test_parse_word_subscript_runs(tmp_path):
    """Word subscript formatting should be captured into TextRun metadata."""
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("H")
    sub_run = para.add_run("2")
    sub_run.font.subscript = True
    para.add_run("O")
    docx_path = _save_doc(doc, tmp_path / "subscript.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    paragraph = next(element.content for element in model.elements if element.element_type == ElementType.PARAGRAPH)
    subscript_runs = [run for run in paragraph.runs if run.subscript]
    assert len(subscript_runs) == 1
    assert subscript_runs[0].text == "2"


def test_parse_word_run_colors(tmp_path):
    """Word run RGB colors should be captured into TextRun metadata."""
    doc = Document()
    para = doc.add_paragraph()
    colored = para.add_run("Loss")
    colored.font.color.rgb = RGBColor(192, 0, 0)
    docx_path = _save_doc(doc, tmp_path / "run_color.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    paragraph = next(element.content for element in model.elements if element.element_type == ElementType.PARAGRAPH)
    assert paragraph.runs[0].color_hex == "#C00000"


def test_parse_word_theme_run_colors(tmp_path):
    """Theme colors should resolve through theme1.xml instead of falling back to dummy val."""
    doc = Document()
    para = doc.add_paragraph()
    themed = para.add_run("Theme")
    themed.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
    docx_path = _save_doc(doc, tmp_path / "theme_color.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    paragraph = next(element.content for element in model.elements if element.element_type == ElementType.PARAGRAPH)
    assert paragraph.runs[0].color_hex == "#4F81BD"


def test_parse_word_theme_tint_colors(tmp_path):
    """Theme tint attributes should lighten the resolved theme color."""
    doc = Document()
    para = doc.add_paragraph()
    themed = para.add_run("Tint")
    themed.font.color.theme_color = MSO_THEME_COLOR_INDEX.ACCENT_1
    docx_path = _save_doc(doc, tmp_path / "theme_tint.docx")

    _rewrite_document_xml(
        docx_path,
        lambda xml: xml.replace(
            'w:themeColor="accent1"/>',
            'w:themeColor="accent1" w:themeTint="80"/>',
        ),
    )

    model = parse_word_file(str(docx_path), extract_images=False)

    paragraph = next(element.content for element in model.elements if element.element_type == ElementType.PARAGRAPH)
    assert paragraph.runs[0].color_hex == "#A7C0DE"


def test_generic_single_cell_table_stays_table(tmp_path):
    """Generic 1x1 tables should not be auto-promoted to callouts."""
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "Quarterly repayment note appears below."
    docx_path = _save_doc(doc, tmp_path / "generic_table.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    assert len(model.elements) == 1
    assert model.elements[0].element_type == ElementType.TABLE


def test_parse_table_alignment_and_nested_list_level(tmp_path):
    """Generic Word tables/lists should preserve alignment and nesting hints."""
    doc = Document()

    doc.add_paragraph("Parent bullet", style="List Bullet")
    child = doc.add_paragraph("Child bullet", style="List Bullet")
    child.paragraph_format.left_indent = Inches(0.5)

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "2026A"
    table.cell(1, 0).text = "Revenue"
    value_para = table.cell(1, 1).paragraphs[0]
    value_para.text = "100"
    value_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    docx_path = _save_doc(doc, tmp_path / "generic_features.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    bullet_items = [element.content for element in model.elements if element.element_type == ElementType.BULLET_LIST]
    table_element = next(element for element in model.elements if element.element_type == ElementType.TABLE)

    assert bullet_items[0].indent_level == 0
    assert bullet_items[1].indent_level == 1
    assert table_element.content.alignments == ["left", "right"]


def test_parse_word_native_numbered_list_values(tmp_path):
    """Word-native numbered lists should recover sequential numbering values."""
    doc = Document()
    doc.add_paragraph("First item", style="List Number")
    doc.add_paragraph("Second item", style="List Number")
    doc.add_paragraph("Third item", style="List Number")
    docx_path = _save_doc(doc, tmp_path / "numbered_values.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    numbers = [
        element.content[0]
        for element in model.elements
        if element.element_type == ElementType.NUMBERED_LIST
    ]
    assert numbers == ["1", "2", "3"]


def test_parse_word_native_numbering_continues_across_paragraphs(tmp_path):
    """A numbered list should continue across an intervening body paragraph."""
    doc = Document()
    doc.add_paragraph("First item", style="List Number")
    doc.add_paragraph("Body explanation")
    doc.add_paragraph("Second item", style="List Number")
    docx_path = _save_doc(doc, tmp_path / "numbering_continues.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    numbered = [
        element.content[0]
        for element in model.elements
        if element.element_type == ElementType.NUMBERED_LIST
    ]
    assert numbered == ["1", "2"]


def test_detect_heading_level_follows_custom_style_base_chain(tmp_path):
    """Custom styles based on Heading styles should still map to heading levels."""
    doc = Document()
    custom_style = doc.styles.add_style("Sector Header", WD_STYLE_TYPE.PARAGRAPH)
    custom_style.base_style = doc.styles["Heading 2"]
    doc.add_paragraph("Custom heading", style=custom_style)
    docx_path = _save_doc(doc, tmp_path / "custom_heading.docx")

    model = parse_word_file(str(docx_path), extract_images=False)

    assert model.elements[0].element_type == ElementType.HEADING_2
    assert model.elements[0].content.text == "Custom heading"


def test_detect_list_type_uses_style_id_for_localized_word_styles():
    """Localized Office can keep English style IDs even when the name is translated."""

    class DummyStyle:
        name = "목록 번호"
        style_id = "ListNumber"

    class DummyParagraph:
        style = DummyStyle()
        text = "항목"
        _p = type("P", (), {"xpath": staticmethod(lambda _expr: [])})()

    assert StyleDetector.detect_list_type(DummyParagraph()) == "number"


def test_numbering_tracker_formats_non_decimal_values():
    """Numbering tracker should support non-decimal Word numbering formats."""
    assert NumberingTracker._format_value(3, "upperRoman") == "III"
    assert NumberingTracker._format_value(4, "lowerLetter") == "d"


def test_parse_ib_generated_doc_recovers_metadata_and_footnotes(tmp_path):
    """IB-generated docs should round-trip cover metadata and ENDNOTES cleanly."""
    markdown = """---
title: "업그레이드 보고서"
subtitle: "부제목"
company: "Korea Development Bank"
ticker: "KDB"
sector: "Banking"
analyst: "Analyst A"
recipient: "테스트 고객"
report_type: "DCM RESEARCH"
date: "2026-03-08"
---

# 본문 제목

본문 문장입니다.

## Citations
1. Source A
"""
    md_path = tmp_path / "ib_roundtrip.md"
    md_path.write_text(markdown, encoding="utf-8")
    docx_path = tmp_path / "ib_roundtrip.docx"

    converter = IBReportConverter(str(md_path), output_path=str(docx_path))
    converter.convert()

    model = parse_word_file(str(docx_path), extract_images=False)

    assert model.metadata.title == "업그레이드 보고서"
    assert model.metadata.subtitle == "부제목"
    assert model.metadata.ticker == "KDB"
    assert model.metadata.company == "Korea Development Bank"
    assert model.metadata.sector == "Banking"
    assert model.metadata.analyst == "Analyst A"
    assert model.metadata.extra["recipient"] == "테스트 고객"
    assert model.metadata.extra["date"] == "2026-03-08"
    assert model.metadata.extra["report_type"] == "DCM RESEARCH"
    assert model.footnotes == {1: "Source A"}
    assert all(element.raw_text != "TABLE OF CONTENTS" for element in model.elements)
    assert all(element.raw_text != "면책 조항" for element in model.elements)


def test_parse_ib_generated_doc_recovers_native_word_footnotes(tmp_path):
    """Native Word footnotes should round-trip without falling back to ENDNOTES."""
    markdown = """# 본문 제목

본문 문장^1^입니다.

## Citations
1. Source A
"""
    md_path = tmp_path / "native_footnote.md"
    md_path.write_text(markdown, encoding="utf-8")
    docx_path = tmp_path / "native_footnote.docx"

    converter = IBReportConverter(str(md_path), output_path=str(docx_path))
    converter.convert()

    footnotes_xml = _read_zip_member(docx_path, "word/footnotes.xml")
    document_xml = _read_zip_member(docx_path, "word/document.xml")

    assert 'w:footnoteReference w:id="1"' in document_xml
    assert 'w:footnote w:id="1"' in footnotes_xml
    assert "ENDNOTES" not in document_xml

    model = parse_word_file(str(docx_path), extract_images=False)

    assert model.footnotes == {1: "Source A"}
    paragraph = next(
        element.content
        for element in model.elements
        if element.element_type == ElementType.PARAGRAPH and "본문 문장" in element.content.text
    )
    assert any(run.superscript and run.text == "1" for run in paragraph.runs)


def test_parse_custom_doc_properties_maps_known_and_unknown_fields(tmp_path):
    """Custom docProps should populate standard metadata and extra fields."""
    doc = Document()
    doc.add_paragraph("Body")
    docx_path = _save_doc(doc, tmp_path / "custom_props.docx")
    _inject_custom_properties(
        docx_path,
        {
            "Company": "Acme Corp",
            "Ticker": "ACME",
            "Recipient": "Client A",
            "Report Type": "INITIATION",
            "Deal Size": "KRW 500bn",
        },
    )

    model = parse_word_file(str(docx_path), extract_images=False)

    assert model.metadata.company == "Acme Corp"
    assert model.metadata.ticker == "ACME"
    assert model.metadata.extra["recipient"] == "Client A"
    assert model.metadata.extra["report_type"] == "INITIATION"
    assert model.metadata.extra["deal_size"] == "KRW 500bn"
