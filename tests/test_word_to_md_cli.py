"""End-to-end tests for the Word-to-Markdown CLI path."""

import base64
import subprocess
import sys
from pathlib import Path
from types import SimpleNamespace

from docx import Document

from word_to_md import WordToMarkdownConverter, run_conversion

PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z8eQAAAAASUVORK5CYII="
)


def _write_png(output_path):
    """Write a tiny PNG fixture used for image extraction tests."""
    output_path.write_bytes(PNG_BYTES)
    return output_path


def test_run_conversion_strip_and_no_frontmatter(tmp_path):
    """CLI conversion should honor strip and no-frontmatter options."""
    doc = Document()
    doc.core_properties.title = "CLI Test"
    doc.add_heading("Test Heading", level=1)
    para = doc.add_paragraph()
    para.add_run("Bold").bold = True
    para.add_run(" paragraph")

    docx_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.md"
    doc.save(str(docx_path))

    args = SimpleNamespace(
        output_file=str(output_path),
        strip=True,
        no_frontmatter=True,
        extract_images=False,
        embed_images_base64=False,
        image_dir=None,
        verbose=False,
    )

    exit_code = run_conversion(docx_path, args)
    rendered = output_path.read_text(encoding="utf-8")

    assert exit_code == 0
    assert not rendered.startswith("---")
    assert "# Test Heading" in rendered
    assert "**" not in rendered
    assert "Bold paragraph" in rendered


def test_converter_extract_images_writes_files_and_markdown_paths(tmp_path, monkeypatch):
    """Image extraction should keep markdown image links relative to the output file."""
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    # Guard against resolving image links from the process CWD instead of the .md location.
    monkeypatch.chdir(workspace)

    source_dir = tmp_path / "source"
    source_dir.mkdir()
    output_dir = tmp_path / "exports"
    image_path = _write_png(source_dir / "tiny.png")

    doc = Document()
    doc.add_picture(str(image_path))
    docx_path = source_dir / "image_input.docx"
    output_path = output_dir / "image_output.md"
    image_dir = output_dir / "image_output_images"
    doc.save(str(docx_path))

    converter = WordToMarkdownConverter(
        docx_file_path=str(docx_path),
        output_path=str(output_path),
        extract_images=True,
    )

    saved_path = converter.convert()
    rendered = output_path.read_text(encoding="utf-8")

    assert saved_path == str(output_path)
    assert "![Image 1](image_output_images/image_1.png)" in rendered
    assert list(image_dir.glob("image_*"))


def test_converter_embeds_images_base64_without_extraction(tmp_path):
    """CLI converter should inline images when base64 embedding is requested."""
    image_path = _write_png(tmp_path / "tiny.png")

    doc = Document()
    doc.add_picture(str(image_path))
    docx_path = tmp_path / "image_input.docx"
    output_path = tmp_path / "image_output.md"
    doc.save(str(docx_path))

    converter = WordToMarkdownConverter(
        docx_file_path=str(docx_path),
        output_path=str(output_path),
        extract_images=False,
        embed_images_base64=True,
    )

    converter.convert()
    rendered = output_path.read_text(encoding="utf-8")

    assert "data:image/png;base64," in rendered
    assert "![Image 1]" in rendered or "![Image]" in rendered or "![Tiny]" in rendered


def test_cli_stdin_output_uses_safe_save_for_nested_parent_dirs(tmp_path):
    """Stdin DOCX output should create missing parent dirs via the safe-save path."""
    doc = Document()
    doc.add_paragraph("Piped body")
    docx_path = tmp_path / "stdin_source.docx"
    doc.save(str(docx_path))

    output_path = tmp_path / "nested" / "exports" / "stdin_output.md"

    # Regression guard: stdin mode used to bypass safe_save and fail on nested output paths.
    result = subprocess.run(
        [
            sys.executable,
            "word_to_md.py",
            "-",
            str(output_path),
            "--no-frontmatter",
        ],
        input=docx_path.read_bytes(),
        capture_output=True,
        cwd=str(Path(__file__).resolve().parent.parent),
        timeout=30,
    )

    assert result.returncode == 0
    assert output_path.exists()
    assert "Piped body" in output_path.read_text(encoding="utf-8")


def test_run_conversion_preserves_word_native_numbering(tmp_path):
    """End-to-end conversion should keep Word-native numbered list values."""
    doc = Document()
    doc.add_paragraph("First item", style="List Number")
    doc.add_paragraph("Second item", style="List Number")
    doc.add_paragraph("Third item", style="List Number")

    docx_path = tmp_path / "numbering.docx"
    output_path = tmp_path / "numbering.md"
    doc.save(str(docx_path))

    args = SimpleNamespace(
        output_file=str(output_path),
        strip=False,
        no_frontmatter=True,
        extract_images=False,
        embed_images_base64=False,
        image_dir=None,
        verbose=False,
        batch=False,
    )

    exit_code = run_conversion(docx_path, args)
    rendered = output_path.read_text(encoding="utf-8")

    assert exit_code == 0
    assert "1. First item" in rendered
    assert "2. Second item" in rendered
    assert "3. Third item" in rendered


def test_run_conversion_keeps_continued_numbering_after_body_paragraph(tmp_path):
    """Numbered lists should not restart at 1 after an intervening paragraph."""
    doc = Document()
    doc.add_paragraph("First item", style="List Number")
    doc.add_paragraph("Body explanation")
    doc.add_paragraph("Second item", style="List Number")

    docx_path = tmp_path / "continued_numbering.docx"
    output_path = tmp_path / "continued_numbering.md"
    doc.save(str(docx_path))

    args = SimpleNamespace(
        output_file=str(output_path),
        strip=False,
        no_frontmatter=True,
        extract_images=False,
        embed_images_base64=False,
        image_dir=None,
        verbose=False,
        batch=False,
    )

    exit_code = run_conversion(docx_path, args)
    rendered = output_path.read_text(encoding="utf-8")

    assert exit_code == 0
    assert "1. First item" in rendered
    assert "2. Second item" in rendered
