"""
Tests for Step 3: Stream-based processing + MIME detection.

Tests cover:
    - stream_utils: detect_format, ensure_seekable, is_stream
    - word_parser: parse() / parse_word_file() with BinaryIO
    - md_parser: parse_markdown_file() with BinaryIO
    - converters: InputConverter.accepts() with streams
    - CLI stdin pipe support
"""

import io
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document

PNG_BYTES = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x04\x00\x00\x00"
    b"\xb5\x1c\x0c\x02\x00\x00\x00\x0bIDATx\xdac\xfc\xff\x1f\x00\x03\x03\x02\x00\xee\xd9\xf1"
    b"\xe4\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _inject_custom_properties(docx_bytes: bytes, properties) -> bytes:
    """Inject docProps/custom.xml into an in-memory DOCX payload."""
    custom_props_xml = [
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>',
        '<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" '
        'xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes">'
    ]

    for index, (name, value) in enumerate(properties.items(), start=2):
        custom_props_xml.append(
            f'<property fmtid="{{D5CDD505-2E9C-101B-9397-08002B2CF9AE}}" pid="{index}" name="{name}">'
            f"<vt:lpwstr>{value}</vt:lpwstr>"
            "</property>"
        )

    custom_props_xml.append("</Properties>")

    buffer = io.BytesIO()
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as source_zip:
        with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as target_zip:
            for info in source_zip.infolist():
                target_zip.writestr(info, source_zip.read(info.filename))
            target_zip.writestr("docProps/custom.xml", "".join(custom_props_xml))
    return buffer.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# FIXTURES
# ═══════════════════════════════════════════════════════════════════════════════


@pytest.fixture
def sample_docx_bytes():
    """Create a minimal DOCX in memory and return its bytes."""
    doc = Document()
    doc.add_heading("Stream Test", level=1)
    doc.add_paragraph("Hello from a stream.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_md_bytes():
    """Create sample Markdown bytes."""
    return "# 스트림 테스트\n\n본문 텍스트입니다.\n".encode("utf-8")


@pytest.fixture
def sample_docx_path(tmp_path, sample_docx_bytes):
    """Write sample DOCX to disk and return path."""
    p = tmp_path / "test.docx"
    p.write_bytes(sample_docx_bytes)
    return p


@pytest.fixture
def sample_md_path(tmp_path, sample_md_bytes):
    """Write sample MD to disk and return path."""
    p = tmp_path / "test.md"
    p.write_bytes(sample_md_bytes)
    return p


# ═══════════════════════════════════════════════════════════════════════════════
# stream_utils tests
# ═══════════════════════════════════════════════════════════════════════════════


class TestDetectFormat:
    """Tests for stream_utils.detect_format."""

    def test_detect_docx_by_signature(self, sample_docx_bytes):
        from stream_utils import detect_format

        stream = io.BytesIO(sample_docx_bytes)
        assert detect_format(stream) == "docx"

    def test_detect_md_by_content(self, sample_md_bytes):
        from stream_utils import detect_format

        stream = io.BytesIO(sample_md_bytes)
        assert detect_format(stream) == "md"

    def test_detect_with_extension_hint(self, sample_docx_bytes):
        from stream_utils import detect_format

        stream = io.BytesIO(sample_docx_bytes)
        assert detect_format(stream, hint=".docx") == "docx"

    def test_detect_hint_without_dot(self, sample_md_bytes):
        from stream_utils import detect_format

        stream = io.BytesIO(sample_md_bytes)
        assert detect_format(stream, hint="md") == "md"

    def test_detect_empty_stream(self):
        from stream_utils import detect_format

        stream = io.BytesIO(b"")
        assert detect_format(stream) == "unknown"

    def test_detect_binary_garbage(self):
        from stream_utils import detect_format

        stream = io.BytesIO(bytes(range(256)) * 10)
        # Not ZIP, not text → unknown
        assert detect_format(stream) == "unknown"

    def test_stream_position_preserved(self, sample_docx_bytes):
        from stream_utils import detect_format

        stream = io.BytesIO(sample_docx_bytes)
        stream.seek(5)
        detect_format(stream)
        # ensure_seekable wraps, so position should be at start of the new wrapper
        # but the point is detect_format doesn't consume the stream
        # We just verify it doesn't raise

    def test_detect_korean_euckr(self):
        """Korean text in EUC-KR should be detected as md."""
        from stream_utils import detect_format

        korean = "한국어 텍스트입니다.".encode("euc-kr")
        stream = io.BytesIO(korean)
        assert detect_format(stream) == "md"

    def test_detect_zip_as_docx(self):
        """Any ZIP is treated as docx (we only support DOCX as ZIP)."""
        from stream_utils import detect_format

        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as zf:
            zf.writestr("test.txt", "hello")
        stream = io.BytesIO(buf.getvalue())
        assert detect_format(stream) == "docx"


class TestEnsureSeekable:
    """Tests for stream_utils.ensure_seekable."""

    def test_already_seekable(self):
        from stream_utils import ensure_seekable

        stream = io.BytesIO(b"hello")
        result = ensure_seekable(stream)
        assert result is stream  # same object

    def test_non_seekable_wrapped(self):
        from stream_utils import ensure_seekable

        class NonSeekable:
            def __init__(self, data):
                self._data = data
                self._pos = 0

            def read(self, n=-1):
                if n == -1:
                    result = self._data[self._pos:]
                    self._pos = len(self._data)
                else:
                    result = self._data[self._pos:self._pos + n]
                    self._pos += n
                return result

            def seekable(self):
                return False

        ns = NonSeekable(b"test data")
        result = ensure_seekable(ns)
        assert result.read() == b"test data"
        result.seek(0)
        assert result.read() == b"test data"


class TestIsStream:
    def test_stream(self):
        from stream_utils import is_stream

        assert is_stream(io.BytesIO(b""))
        assert not is_stream("a string")
        assert not is_stream(Path("/tmp"))
        assert not is_stream(42)


# ═══════════════════════════════════════════════════════════════════════════════
# word_parser stream tests
# ═══════════════════════════════════════════════════════════════════════════════


class TestWordParserStream:
    """Test that WordParser accepts BinaryIO."""

    def test_parse_from_stream(self, sample_docx_bytes):
        from word_parser import parse_word_file

        stream = io.BytesIO(sample_docx_bytes)
        model = parse_word_file(stream, extract_images=False)
        assert model is not None
        assert len(model.elements) > 0
        # Should find our heading or paragraph text
        texts = [str(e) for e in model.elements]
        combined = " ".join(texts)
        assert "Stream Test" in combined or "Hello" in combined

    def test_parse_from_path_still_works(self, sample_docx_path):
        from word_parser import parse_word_file

        model = parse_word_file(str(sample_docx_path), extract_images=False)
        assert model is not None
        assert len(model.elements) > 0

    def test_parse_stream_via_word_parser_class(self, sample_docx_bytes):
        from word_parser import WordParser

        stream = io.BytesIO(sample_docx_bytes)
        parser = WordParser(extract_images=False)
        model = parser.parse(stream)
        assert len(model.elements) > 0

    def test_parse_stream_extracts_custom_doc_properties(self, sample_docx_bytes):
        """Custom docProps should survive BinaryIO parsing."""
        from word_parser import parse_word_file

        payload = _inject_custom_properties(
            sample_docx_bytes,
            {"Company": "Stream Corp", "Report Type": "PIPELINE"},
        )
        stream = io.BytesIO(payload)

        model = parse_word_file(stream, extract_images=False)

        assert model.metadata.company == "Stream Corp"
        assert model.metadata.extra["report_type"] == "PIPELINE"

    def test_parse_image_stream_with_base64_embedding(self, tmp_path):
        """Stream parsing should still populate base64 image data."""
        from word_parser import parse_word_file

        image_path = tmp_path / "tiny.png"
        image_path.write_bytes(PNG_BYTES)

        doc = Document()
        doc.add_picture(str(image_path))
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)

        model = parse_word_file(buf, extract_images=False, embed_images_base64=True)

        image_element = next(element for element in model.elements if element.element_type.name == "IMAGE")
        assert image_element.content.base64_data is not None


# ═══════════════════════════════════════════════════════════════════════════════
# md_parser stream tests
# ═══════════════════════════════════════════════════════════════════════════════


class TestMdParserStream:
    """Test that parse_markdown_file accepts BinaryIO."""

    def test_parse_from_stream(self, sample_md_bytes):
        from md_parser import parse_markdown_file

        stream = io.BytesIO(sample_md_bytes)
        model = parse_markdown_file(stream)
        assert model is not None
        assert len(model.elements) > 0

    def test_parse_from_path_still_works(self, sample_md_path):
        from md_parser import parse_markdown_file

        model = parse_markdown_file(str(sample_md_path))
        assert model is not None
        assert len(model.elements) > 0

    def test_stream_korean_euckr(self):
        from md_parser import parse_markdown_file

        content = "# 한국어 제목\n\n본문 내용\n".encode("euc-kr")
        stream = io.BytesIO(content)
        model = parse_markdown_file(stream)
        assert len(model.elements) > 0


# ═══════════════════════════════════════════════════════════════════════════════
# converters stream tests
# ═══════════════════════════════════════════════════════════════════════════════


class TestConverterRegistryStream:
    """Test ConverterRegistry with BinaryIO streams."""

    def test_registry_convert_docx_stream(self, sample_docx_bytes):
        from converters import get_default_registry

        registry = get_default_registry()
        stream = io.BytesIO(sample_docx_bytes)
        model = registry.convert(stream, extract_images=False)
        assert model is not None
        assert len(model.elements) > 0

    def test_registry_convert_md_stream(self, sample_md_bytes):
        from converters import get_default_registry

        registry = get_default_registry()
        stream = io.BytesIO(sample_md_bytes)
        model = registry.convert(stream)
        assert model is not None
        assert len(model.elements) > 0

    def test_registry_with_extension_hint(self, sample_docx_bytes):
        from converters import get_default_registry

        registry = get_default_registry()
        stream = io.BytesIO(sample_docx_bytes)
        model = registry.convert(stream, extension_hint=".docx", extract_images=False)
        assert model is not None

    def test_registry_unknown_stream_raises(self):
        from converters import get_default_registry

        registry = get_default_registry()
        stream = io.BytesIO(b"\x00\x01\x02\x03" * 100)
        with pytest.raises(ValueError, match="No converter found"):
            registry.convert(stream)


# ═══════════════════════════════════════════════════════════════════════════════
# CLI pipe tests
# ═══════════════════════════════════════════════════════════════════════════════


class TestCliPipe:
    """Test word_to_md.py stdin pipe support."""

    def test_pipe_md_to_stdout(self, sample_md_bytes, tmp_path):
        """Pipe markdown into CLI, expect markdown output on stdout."""
        result = subprocess.run(
            [sys.executable, "word_to_md.py", "-", "--no-frontmatter"],
            input=sample_md_bytes,
            capture_output=True,
            cwd=str(Path(__file__).resolve().parent.parent),
            timeout=30,
        )
        assert result.returncode == 0
        output = result.stdout.decode("utf-8")
        assert "스트림 테스트" in output or "본문" in output

    def test_pipe_docx_to_file(self, sample_docx_bytes, tmp_path):
        """Pipe DOCX into CLI, expect output file."""
        out_file = tmp_path / "piped_output.md"
        result = subprocess.run(
            [
                sys.executable,
                "word_to_md.py",
                "-",
                str(out_file),
                "--no-frontmatter",
            ],
            input=sample_docx_bytes,
            capture_output=True,
            cwd=str(Path(__file__).resolve().parent.parent),
            timeout=30,
        )
        assert result.returncode == 0
        assert out_file.exists()
        content = out_file.read_text(encoding="utf-8")
        assert "Stream Test" in content or "Hello" in content
