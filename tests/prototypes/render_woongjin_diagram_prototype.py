"""
Standalone prototype renderer for the Woongjin group structure diagram.

This script does not touch the main md->word pipeline. It extracts the fenced
tree block from tests/웅진_계열사.md and renders only that section into a
separate Word document for visual review.
"""

import re
from pathlib import Path
from typing import List

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SOURCE_PATH = Path(__file__).resolve().parent.parent / "웅진_계열사.md"
OUTPUT_PATH = Path(__file__).resolve().parent / "웅진_계열사_도식_프로토타입.docx"

TREE_PREFIX_RE = re.compile(r"^([\s│├└─]+)(.*)$")
VALUE_SPLIT_RE = re.compile(r"^(.*?)(\d+(?:\.\d+)?%|실질지배)(\s+──\s+)(.*)$")

NAVY = RGBColor(0, 51, 102)
DARK = RGBColor(54, 60, 67)
MID = RGBColor(120, 128, 140)
PANEL_BG = "F7F8FA"
PANEL_BORDER = "D6DBE1"


def extract_diagram_lines(markdown_text: str) -> List[str]:
    """Extract the first fenced code block from the markdown sample."""
    lines = markdown_text.splitlines()
    inside = False
    diagram_lines: List[str] = []

    for line in lines:
        if line.strip().startswith("```") and not inside:
            inside = True
            continue
        if line.strip().startswith("```") and inside:
            break
        if inside:
            diagram_lines.append(line.rstrip("\n"))

    return diagram_lines


def add_panel_border(table) -> None:
    """Style the single-cell table as a soft bordered panel."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tbl_borders = OxmlElement("w:tblBorders")

    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "6")
        border.set(qn("w:color"), PANEL_BORDER)
        tbl_borders.append(border)

    for edge in ("insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "nil")
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)
    if tbl.tblPr is None:
        tbl.insert(0, tbl_pr)


def set_cell_margins(cell) -> None:
    """Apply comfortable padding to the panel cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")

    for edge, width in (("top", "150"), ("left", "220"), ("bottom", "150"), ("right", "220")):
        margin = OxmlElement(f"w:{edge}")
        margin.set(qn("w:w"), width)
        margin.set(qn("w:type"), "dxa")
        tc_mar.append(margin)

    tc_pr.append(tc_mar)


def style_run(run, font_name: str, font_size: float, color: RGBColor, bold: bool = False) -> None:
    """Apply a consistent style to a run."""
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.color.rgb = color

    r_pr = run._element.get_or_add_rPr()
    if r_pr.rFonts is None:
        r_pr.get_or_add_rFonts()
    r_pr.rFonts.set(qn("w:eastAsia"), font_name)


def add_styled_tree_line(paragraph, line: str) -> None:
    """Render one tree line with light semantic emphasis."""
    if not line.strip():
        spacer = paragraph.add_run(" ")
        style_run(spacer, "Consolas", 9.5, MID)
        return

    prefix_match = TREE_PREFIX_RE.match(line)
    if prefix_match:
        prefix, remainder = prefix_match.groups()
    else:
        prefix, remainder = "", line

    if prefix:
        run = paragraph.add_run(prefix)
        style_run(run, "Consolas", 9.5, MID)

    value_match = VALUE_SPLIT_RE.match(remainder)
    if value_match:
        before, value, divider, after = value_match.groups()
        if before:
            run = paragraph.add_run(before)
            style_run(run, "Consolas", 9.5, DARK)
        value_run = paragraph.add_run(value)
        style_run(value_run, "Consolas", 9.5, NAVY, bold=True)
        divider_run = paragraph.add_run(divider)
        style_run(divider_run, "Consolas", 9.5, MID)
        tail_run = paragraph.add_run(after)
        style_run(tail_run, "Consolas", 9.5, DARK)
        return

    is_root = "사업지주회사" in remainder
    run = paragraph.add_run(remainder)
    style_run(run, "Consolas", 9.8 if is_root else 9.5, NAVY if is_root else DARK, bold=is_root)


def render_diagram(diagram_lines: List[str]) -> Path:
    """Render the extracted diagram lines into a standalone docx."""
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("웅진 계열회사 도식 프로토타입")
    style_run(title_run, "Arial", 16, NAVY, bold=True)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("ASCII tree block rendered as a review-only diagram panel")
    style_run(note_run, "Calibri", 9.5, MID)

    panel = doc.add_table(rows=1, cols=1)
    panel.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = panel.rows[0].cells[0]
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), PANEL_BG)
    cell._tc.get_or_add_tcPr().append(shd)
    set_cell_margins(cell)
    add_panel_border(panel)

    first = True
    for line in diagram_lines:
        paragraph = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        add_styled_tree_line(paragraph, line)

    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


def main() -> None:
    markdown = SOURCE_PATH.read_text(encoding="utf-8")
    diagram_lines = extract_diagram_lines(markdown)
    output = render_diagram(diagram_lines)
    print(output)


if __name__ == "__main__":
    main()
