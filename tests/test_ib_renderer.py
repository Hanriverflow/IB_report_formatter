"""
Tests for ib_renderer module.

Tests rendering functionality including:
- Table number formatting
- Callout box styling
- Image rendering
- Text formatting (bold, italic, superscript)
- Separator rendering
- Callout content formatting
"""

import base64

import pytest

from docx import Document

from md_parser import (
    Blockquote,
    CodeBlock,
    Diagram,
    DiagramBox,
    DocumentMetadata,
    DocumentModel,
    Element,
    ElementType,
    Heading,
    Image,
    ListItem,
    Paragraph,
    Table,
    TableCell,
    TableRow,
    TableType,
    TextRun,
)
from word_parser import parse_word_file

PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Z8eQAAAAASUVORK5CYII="
)


def _semantic_bookmark_names(doc: Document):
    """Collect ib-report semantic bookmark names from a rendered document."""
    bookmark_starts = doc.element.xpath(".//*[local-name()='bookmarkStart']")
    names = []
    for bookmark in bookmark_starts:
        name = bookmark.get(
            "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}name",
            "",
        )
        if name.startswith("_ibrep_"):
            names.append(name)
    return names

# ═══════════════════════════════════════════════════════════════════════════════
# TABLE RENDERER TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestTableRendererFormatting:
    """Tests for TableRenderer._format_financial_number static method."""

    @pytest.fixture
    def format_number(self):
        """Import the format method."""
        from ib_renderer import TableRenderer

        return TableRenderer._format_financial_number

    def test_format_plain_number(self, format_number):
        """Format plain integers with thousand separators."""
        assert format_number("1234567") == "1,234,567"
        assert format_number("1000") == "1,000"
        assert format_number("999") == "999"
        assert format_number("100") == "100"

    def test_format_decimal_number(self, format_number):
        """Format decimal numbers."""
        assert format_number("1234.56") == "1,234.56"
        assert format_number("1000000.99") == "1,000,000.99"

    def test_format_negative_parentheses(self, format_number):
        """Format negative numbers in parentheses."""
        assert format_number("(1234567)") == "(1,234,567)"
        assert format_number("(500)") == "(500)"

    def test_format_negative_minus(self, format_number):
        """Format negative numbers with minus sign."""
        assert format_number("-1234567") == "-1,234,567"
        assert format_number("-500") == "-500"

    def test_format_with_suffix(self, format_number):
        """Format numbers with Korean unit suffixes."""
        assert format_number("1234567억") == "1,234,567억"
        assert format_number("100조") == "100조"
        assert format_number("50%") == "50%"
        assert format_number("12.5%") == "12.5%"

    def test_skip_already_formatted(self, format_number):
        """Skip numbers already with thousand separators."""
        assert format_number("1,234,567") == "1,234,567"
        assert format_number("1,000") == "1,000"

    def test_skip_non_numeric(self, format_number):
        """Skip non-numeric content."""
        assert format_number("Revenue") == "Revenue"
        assert format_number("N/A") == "N/A"
        assert format_number("") == ""
        assert format_number("-") == "-"

    def test_content_aware_widths_give_more_space_to_text_column(self):
        """Text-heavy columns should receive more width than compact numeric columns."""
        from docx import Document

        from ib_renderer import DocumentStyler, TableRenderer

        doc = Document()
        DocumentStyler(doc).create_styles()

        table = Table(
            rows=[
                TableRow(
                    cells=[
                        TableCell(content="항목"),
                        TableCell(content="2024"),
                        TableCell(content="2025"),
                    ],
                    is_header=True,
                ),
                TableRow(
                    cells=[
                        TableCell(
                            content="매출원가 상승과 판관비 증가가 동시에 반영된 장문 설명 열",
                        ),
                        TableCell(content="10.2%", is_numeric=True),
                        TableCell(content="11.8%", is_numeric=True),
                    ]
                ),
                TableRow(
                    cells=[
                        TableCell(
                            content="환율 및 원재료 가격 변동의 복합 영향이 드러나는 설명 열",
                        ),
                        TableCell(content="9.8%", is_numeric=True),
                        TableCell(content="10.1%", is_numeric=True),
                    ]
                ),
            ],
            col_count=3,
        )

        TableRenderer(doc).render(table)

        rendered = doc.tables[0]
        first_width = rendered.columns[0].width
        second_width = rendered.columns[1].width
        third_width = rendered.columns[2].width

        assert rendered.autofit is False
        assert first_width > second_width
        assert first_width > third_width

    def test_width_estimator_caps_numeric_columns_and_uses_full_table_width(self):
        """Numeric columns should stay compact while the table still fills the line."""
        from docx import Document

        from ib_renderer import DocumentStyler, TableRenderer

        doc = Document()
        DocumentStyler(doc).create_styles()
        renderer = TableRenderer(doc)
        table = Table(
            rows=[
                TableRow(
                    cells=[
                        TableCell(content="구분"),
                        TableCell(content="1Q"),
                        TableCell(content="2Q"),
                        TableCell(content="3Q"),
                        TableCell(content="4Q"),
                    ],
                    is_header=True,
                ),
                TableRow(
                    cells=[
                        TableCell(content="설명 텍스트가 긴 핵심 지표"),
                        TableCell(content="101", is_numeric=True),
                        TableCell(content="102", is_numeric=True),
                        TableCell(content="103", is_numeric=True),
                        TableCell(content="104", is_numeric=True),
                    ]
                ),
            ],
            col_count=5,
        )

        widths = renderer._estimate_column_widths(table, renderer._get_available_table_width_inches())

        assert len(widths) == 5
        assert abs(sum(widths) - renderer._get_available_table_width_inches()) < 0.05
        assert max(widths[1:]) <= 1.35
        assert widths[0] == max(widths)

    def test_column_kind_inference_checks_first_three_body_cells(self):
        """Column kind should be inferred from the first few body cells, not header text."""
        from docx import Document

        from ib_renderer import DocumentStyler, TableRenderer

        doc = Document()
        DocumentStyler(doc).create_styles()
        renderer = TableRenderer(doc)
        table = Table(
            rows=[
                TableRow(
                    cells=[TableCell(content="구분"), TableCell(content="값"), TableCell(content="코드")],
                    is_header=True,
                ),
                TableRow(
                    cells=[
                        TableCell(content="매출"),
                        TableCell(content="10.2%"),
                        TableCell(content="A-101"),
                    ]
                ),
                TableRow(
                    cells=[
                        TableCell(content="영업이익"),
                        TableCell(content="11.8%"),
                        TableCell(content="B-102"),
                    ]
                ),
                TableRow(
                    cells=[
                        TableCell(content="순이익"),
                        TableCell(content="9.7%"),
                        TableCell(content="C-103"),
                    ]
                ),
            ],
            col_count=3,
        )

        assert renderer._infer_column_kinds(table) == ["text", "numeric", "text"]


# ═══════════════════════════════════════════════════════════════════════════════
# CALLOUT RENDERER TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestCalloutStyles:
    """Tests for CalloutRenderer style configurations."""

    def test_callout_style_exists(self):
        """Verify callout styles are defined."""
        from ib_renderer import CalloutRenderer

        styles = CalloutRenderer._CALLOUT_STYLES

        # Check key styles exist
        assert "EXECUTIVE SUMMARY" in styles
        assert "요약" in styles
        assert "KEY INSIGHT" in styles
        assert "시사점" in styles
        assert "WARNING" in styles
        assert "주의" in styles
        assert "NOTE" in styles
        assert "참고" in styles

    def test_executive_summary_has_navy_background(self):
        """Executive Summary should have navy background."""
        from ib_renderer import STYLE, CalloutRenderer

        style = CalloutRenderer._CALLOUT_STYLES["EXECUTIVE SUMMARY"]
        bg_hex, _, _, _ = style

        assert bg_hex == STYLE.NAVY_HEX


# ═══════════════════════════════════════════════════════════════════════════════
# IMAGE RENDERER TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestImageRendererMimeTypes:
    """Tests for ImageRenderer MIME type handling."""

    def test_mime_to_extension(self):
        """Convert MIME types to file extensions."""
        from ib_renderer import ImageRenderer

        assert ImageRenderer._mime_to_extension("image/png") == ".png"
        assert ImageRenderer._mime_to_extension("image/jpeg") == ".jpg"
        assert ImageRenderer._mime_to_extension("image/gif") == ".gif"
        assert ImageRenderer._mime_to_extension("image/webp") == ".webp"
        assert ImageRenderer._mime_to_extension("image/svg+xml") == ".svg"

    def test_unknown_mime_defaults_to_png(self):
        """Unknown MIME types default to .png."""
        from ib_renderer import ImageRenderer

        assert ImageRenderer._mime_to_extension("image/unknown") == ".png"
        assert ImageRenderer._mime_to_extension("application/pdf") == ".png"

    def test_inserted_image_sets_word_alt_text(self, tmp_path):
        """Inserted images should carry alt text in Word metadata and round-trip back out."""
        from ib_renderer import IBDocumentRenderer

        image_path = tmp_path / "tiny.png"
        image_path.write_bytes(PNG_BYTES)

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.IMAGE,
                    content=Image(alt_text="Revenue bridge", path=str(image_path)),
                )
            ]
        )

        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        assert len(doc.inline_shapes) >= 1
        inline_shape = doc.inline_shapes[0]
        assert inline_shape._inline.docPr.get("descr") == "Revenue bridge"
        assert inline_shape._inline.docPr.get("title") == "Revenue bridge"

        output_path = tmp_path / "roundtrip_image.docx"
        doc.save(str(output_path))
        parsed = parse_word_file(str(output_path), extract_images=False)
        image_element = next(
            element for element in parsed.elements if element.element_type == ElementType.IMAGE
        )
        assert image_element.content.alt_text == "Revenue bridge"


# ═══════════════════════════════════════════════════════════════════════════════
# LATEX RENDERER TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestLaTeXRenderer:
    """Tests for LaTeXRenderer."""

    def test_latex_availability_check(self):
        """Check matplotlib availability detection."""
        from ib_renderer import LaTeXRenderer

        # Should return True or False, not raise
        result = LaTeXRenderer.is_available()
        assert isinstance(result, bool)

    @pytest.mark.skipif(
        not __import__("importlib.util", fromlist=[""]).find_spec("matplotlib"),
        reason="matplotlib not installed",
    )
    def test_render_simple_equation(self):
        """Render a simple LaTeX equation."""
        import os

        from ib_renderer import LaTeXRenderer

        if not LaTeXRenderer.is_available():
            pytest.skip("matplotlib not available")

        path = LaTeXRenderer.render_to_image("x^2 + y^2 = z^2")

        assert path is not None
        assert os.path.exists(path)
        assert path.endswith(".png")

        # Cleanup
        os.unlink(path)

    def test_display_text_conversion_for_korean_equation(self):
        """Korean text equations should become readable unicode fallback text."""
        from ib_renderer import LaTeXRenderer

        expression = (
            r"\text{K-ICS비율} = \frac{\text{가용자본}}{\text{요구자본}} "
            r"\times 100 \geq 50\%"
        )

        assert (
            LaTeXRenderer.to_display_text(expression)
            == "K-ICS비율 = (가용자본 / 요구자본) × 100 ≥ 50%"
        )

    def test_display_text_conversion_keeps_mixed_korean_math_readable(self):
        """Mixed Korean and LaTeX commands should render as readable unicode text."""
        from ib_renderer import LaTeXRenderer

        expression = r"\text{가용자본}_t + \alpha"

        assert LaTeXRenderer.to_display_text(expression) == "가용자본_t + α"

    def test_display_text_conversion_for_korean_summation(self):
        """Summation operators should not remain as literal LaTeX commands."""
        from ib_renderer import LaTeXRenderer

        expression = r"\text{요구자본} = \sum_{i=1}^{n} x_i"

        assert LaTeXRenderer.to_display_text(expression) == "요구자본 = ∑_i=1^n x_i"

    def test_non_ascii_equation_prefers_plain_text_renderer(self, monkeypatch):
        """Non-ASCII equations should use the unicode fallback renderer first."""
        from ib_renderer import LaTeXRenderer

        monkeypatch.setattr(
            LaTeXRenderer,
            "is_available",
            classmethod(lambda cls: True),
        )
        monkeypatch.setattr(
            LaTeXRenderer,
            "_render_plain_text_to_image",
            classmethod(lambda cls, display_text, fontsize, dpi: "plain.png"),
        )
        monkeypatch.setattr(
            LaTeXRenderer,
            "_render_mathtext_to_image",
            classmethod(
                lambda cls, expression, fontsize, dpi: (_ for _ in ()).throw(
                    AssertionError("mathtext path should not run for non-ASCII equations")
                )
            ),
        )

        result = LaTeXRenderer.render_to_image(r"\text{가용자본}")

        assert result == "plain.png"

    def test_inline_latex_runs_render_as_inline_images(self, tmp_path, monkeypatch):
        """Inline LaTeX runs should render as inline pictures rather than fallback text."""
        from ib_renderer import DocumentStyler, TextRenderer

        image_path = tmp_path / "latex-inline.png"
        image_path.write_bytes(PNG_BYTES)
        monkeypatch.setattr(
            "ib_renderer.LaTeXRenderer.render_to_image",
            lambda expression, fontsize=14, dpi=150: str(image_path),
        )

        doc = Document()
        DocumentStyler(doc).create_styles()
        paragraph = doc.add_paragraph()
        TextRenderer.render_runs(
            paragraph,
            [
                TextRun(text="Formula "),
                TextRun(text="x^2", is_latex=True),
                TextRun(text=" end"),
            ],
        )

        assert len(doc.inline_shapes) == 1
        assert "Formula " in paragraph.text
        assert " end" in paragraph.text
        assert "[x^2]" not in paragraph.text

    def test_inline_latex_paragraph_renders_picture(self, tmp_path, monkeypatch):
        """Paragraph inline LaTeX should become an inline picture in the rendered document."""
        from ib_renderer import IBDocumentRenderer

        image_path = tmp_path / "latex-paragraph.png"
        image_path.write_bytes(PNG_BYTES)
        monkeypatch.setattr(
            "ib_renderer.LaTeXRenderer.render_to_image",
            lambda expression, fontsize=14, dpi=150: str(image_path),
        )

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.PARAGRAPH,
                    content=Paragraph(
                        text="Formula $x^2$ end",
                        runs=[
                            TextRun(text="Formula "),
                            TextRun(text="x^2", is_latex=True),
                            TextRun(text=" end"),
                        ],
                        has_inline_latex=True,
                    ),
                )
            ]
        )

        doc = IBDocumentRenderer().render(model)

        assert len(doc.inline_shapes) >= 1


# ═══════════════════════════════════════════════════════════════════════════════
# STYLE CONFIGURATION TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestIBStyle:
    """Tests for IBStyle configuration."""

    def test_style_colors_defined(self):
        """Verify essential colors are defined."""
        from ib_renderer import STYLE

        assert STYLE.NAVY is not None
        assert STYLE.DARK_GRAY is not None
        assert STYLE.WHITE is not None
        assert STYLE.RED is not None

    def test_style_fonts_defined(self):
        """Verify fonts are defined."""
        from ib_renderer import STYLE

        assert STYLE.HEADING_FONT == "Arial"
        assert STYLE.BODY_FONT == "Calibri"
        assert STYLE.KOREAN_FONT == "Malgun Gothic"

    def test_font_policy_prefers_macos_font(self):
        """macOS should prefer a mac-native Korean font first."""
        from ib_renderer import FontPolicy

        assert FontPolicy.resolve_korean_font("Darwin") == "Apple SD Gothic Neo"

    def test_font_policy_keeps_windows_default(self):
        """Windows should keep the existing default Korean font."""
        from ib_renderer import STYLE, FontPolicy

        assert FontPolicy.resolve_korean_font("Windows") == STYLE.KOREAN_FONT

    def test_style_hex_colors(self):
        """Verify hex color strings are valid."""
        from ib_renderer import STYLE

        # Should be 6-character hex strings
        assert len(STYLE.NAVY_HEX) == 6
        assert len(STYLE.LIGHT_GRAY_HEX) == 6
        assert all(c in "0123456789ABCDEF" for c in STYLE.NAVY_HEX.upper())


# ═══════════════════════════════════════════════════════════════════════════════
# DISCLAIMER RENDERER TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestDisclaimerRendererFormatting:
    """Tests for disclaimer paragraph formatting behavior."""

    def test_split_content_lines_removes_blank_lines(self):
        """Split helper should keep only non-empty lines."""
        from ib_renderer import DisclaimerRenderer

        content = "첫 문장\n\n둘째 문장\n  \n셋째 문장"
        lines = DisclaimerRenderer._split_content_lines(content)

        assert lines == ["첫 문장", "둘째 문장", "셋째 문장"]

    def test_disclaimer_content_uses_plain_paragraph_text(self):
        """Disclaimer content should not inject manual newline characters."""
        from docx import Document

        from ib_renderer import DisclaimerRenderer, DocumentStyler

        doc = Document()
        DocumentStyler(doc).create_styles()
        renderer = DisclaimerRenderer(doc)
        renderer.render("Korea Development Bank")

        target_lines = [
            "본 자료는 해당 문서에 최대한 정확하고 완전한 정보를 담고자 노력하였으나,",
            "본 자료는 당행의 저작물로서 모든 저작권은 당행에게 있으며, 당행의 동의 없이",
            "본 제안서의 내용은 현재의 시장상황 및 발행구조에 대한 기초정보에 근거한 것으로",
        ]

        for line in target_lines:
            paragraph = next((p for p in doc.paragraphs if line in p.text), None)
            assert paragraph is not None
            assert "\n" not in paragraph.text


class TestCoverRendererDisclaimerTable:
    """Tests for cover disclaimer table rendering behavior."""

    def test_cover_includes_single_cell_disclaimer_table(self):
        """Cover page should include a disclaimer table with requested legal sentence."""
        from docx import Document

        from ib_renderer import CoverRenderer, DocumentStyler
        from md_parser import DocumentMetadata

        doc = Document()
        DocumentStyler(doc).create_styles()

        metadata = DocumentMetadata(
            title="테스트 보고서",
            subtitle="테스트 부제",
            company="Korea Development Bank",
            ticker="TEST",
            sector="SECTOR",
            analyst="DCM Team 1",
        )

        CoverRenderer(doc).render(metadata)

        one_cell_tables = [
            table for table in doc.tables if len(table.rows) == 1 and len(table.rows[0].cells) == 1
        ]
        assert one_cell_tables

        disclaimer_table = None
        for table in one_cell_tables:
            text = table.rows[0].cells[0].text
            if "당행은 해당 문서에 최대한 정확하고 완전한 정보를 담고자 노력하였으나" in text:
                disclaimer_table = table
                break

        assert disclaimer_table is not None


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT FORMATTING RENDERING TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestTextRendererFormatting:
    """Tests for TextRenderer.render_text_with_formatting."""

    @pytest.fixture
    def doc(self):
        from ib_renderer import DocumentStyler

        d = Document()
        DocumentStyler(d).create_styles()
        return d

    def test_bold_markers_rendered(self, doc):
        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        TextRenderer.render_text_with_formatting(p, "Normal **Bold** text")
        runs = p.runs
        # Should have 3 runs: "Normal", "Bold", "text"
        assert len(runs) >= 2
        bold_runs = [r for r in runs if r.font.bold]
        assert len(bold_runs) >= 1
        assert "Bold" in bold_runs[0].text

    def test_italic_markers_rendered(self, doc):
        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        TextRenderer.render_text_with_formatting(p, "Normal *italic* text")
        runs = p.runs
        italic_runs = [r for r in runs if r.font.italic]
        assert len(italic_runs) >= 1
        assert "italic" in italic_runs[0].text

    def test_superscript_markers_rendered(self, doc):
        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        TextRenderer.render_text_with_formatting(p, "footnote^1^")
        runs = p.runs
        super_runs = [r for r in runs if r.font.superscript]
        assert len(super_runs) >= 1
        assert "1" in super_runs[0].text

    def test_subscript_markers_rendered(self, doc):
        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        TextRenderer.render_text_with_formatting(p, "H~2~O")
        sub_runs = [r for r in p.runs if r.font.subscript]
        assert len(sub_runs) >= 1
        assert "2" in sub_runs[0].text

    def test_bold_and_italic_combined(self, doc):
        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        TextRenderer.render_text_with_formatting(p, "**Bold** and *italic* mixed")
        bold_runs = [r for r in p.runs if r.font.bold]
        italic_runs = [r for r in p.runs if r.font.italic]
        assert len(bold_runs) >= 1
        assert len(italic_runs) >= 1

    def test_plain_text_no_extra_runs(self, doc):
        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        TextRenderer.render_text_with_formatting(p, "Plain text only")
        assert len(p.runs) == 1
        assert p.runs[0].text.strip() == "Plain text only"

    def test_default_color_applied(self, doc):
        from docx.shared import RGBColor

        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        white = RGBColor(255, 255, 255)
        TextRenderer.render_text_with_formatting(
            p, "White text", default_color=white
        )
        assert p.runs[0].font.color.rgb == white

    def test_color_span_rendered(self, doc):
        from docx.shared import RGBColor

        from ib_renderer import TextRenderer

        p = doc.add_paragraph()
        TextRenderer.render_text_with_formatting(
            p, '<span style="color:#C00000">Loss</span>'
        )
        assert p.runs[0].font.color.rgb == RGBColor(192, 0, 0)


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE CELL RUNS RENDERING TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestTableCellRunsRendering:
    """Tests that table cells use TextRun formatting when available."""

    def _build_model_with_table(self, cell_runs=None, cell_text="Content"):
        """Helper to build a DocumentModel with a single table."""
        cells = [TableCell(content=cell_text, runs=cell_runs or [])]
        header = TableRow(cells=[TableCell(content="Header")])
        data = TableRow(cells=cells)
        table = Table(rows=[header, data], col_count=1)
        return DocumentModel(
            elements=[Element(element_type=ElementType.TABLE, content=table)]
        )

    def test_cell_with_runs_uses_runs(self):
        from ib_renderer import DocumentStyler, IBDocumentRenderer

        runs = [
            TextRun(text="Bold", bold=True),
            TextRun(text=" normal"),
        ]
        model = self._build_model_with_table(cell_runs=runs)

        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        table = self._find_content_table(doc, "Bold")
        assert table is not None, "Table with 'Bold' not found"
        data_cell = table.rows[1].cells[0]
        cell_runs = data_cell.paragraphs[0].runs

        bold_found = any(r.font.bold for r in cell_runs)
        assert bold_found, "Cell runs should preserve bold formatting"

    def _find_content_table(self, doc, search_text):
        """Find the table containing search_text (skip cover/disclaimer tables)."""
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    if search_text in cell.text:
                        return t
        return None

    def test_cell_without_runs_falls_back_to_text(self):
        from ib_renderer import DocumentStyler, IBDocumentRenderer

        model = self._build_model_with_table(cell_text="**Important** value")

        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        table = self._find_content_table(doc, "Important")
        assert table is not None, "Table with 'Important' not found"
        data_cell = table.rows[1].cells[0]
        text = data_cell.text
        assert "Important" in text
        assert "**" not in text  # Markers should be parsed, not literal

    def test_header_cell_with_runs_uses_runs(self):
        from ib_renderer import DocumentStyler, IBDocumentRenderer

        runs = [TextRun(text="Revenue"), TextRun(text=" (M)", italic=True)]
        header = TableRow(cells=[TableCell(content="Revenue (M)", runs=runs)])
        data = TableRow(cells=[TableCell(content="100")])
        table = Table(rows=[header, data], col_count=1)
        model = DocumentModel(
            elements=[Element(element_type=ElementType.TABLE, content=table)]
        )

        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        content_table = self._find_content_table(doc, "Revenue")
        assert content_table is not None, "Table with 'Revenue' not found"
        header_cell = content_table.rows[0].cells[0]
        header_runs = header_cell.paragraphs[0].runs
        italic_found = any(r.font.italic for r in header_runs)
        assert italic_found, "Header runs should preserve italic"

    def test_table_cells_use_compact_paragraph_spacing_and_center_vertical_align(self):
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        from ib_renderer import IBDocumentRenderer

        model = self._build_model_with_table(
            cell_runs=[TextRun(text="100")],
            cell_text="100",
        )

        doc = IBDocumentRenderer().render(model)
        table = self._find_content_table(doc, "100")
        assert table is not None, "Table with '100' not found"

        header_cell = table.rows[0].cells[0]
        data_cell = table.rows[1].cells[0]

        assert header_cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        assert data_cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.CENTER
        assert header_cell.paragraphs[0].paragraph_format.space_after.pt == 0
        assert data_cell.paragraphs[0].paragraph_format.space_after.pt == 0

    def test_table_body_alignment_infers_text_left_and_numeric_right(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        from ib_renderer import IBDocumentRenderer

        table = Table(
            rows=[
                TableRow(
                    cells=[
                        TableCell(content="항목"),
                        TableCell(content="2024"),
                        TableCell(content="비고"),
                    ],
                    is_header=True,
                ),
                TableRow(
                    cells=[
                        TableCell(content="매출총이익"),
                        TableCell(content="1,250"),
                        TableCell(content="정상화 완료"),
                    ]
                ),
                TableRow(
                    cells=[
                        TableCell(content="영업이익"),
                        TableCell(content="980"),
                        TableCell(content="일회성 비용 반영"),
                    ]
                ),
                TableRow(
                    cells=[
                        TableCell(content="순이익"),
                        TableCell(content="815"),
                        TableCell(content="현금흐름 안정"),
                    ]
                ),
            ],
            col_count=3,
        )
        model = DocumentModel(
            elements=[Element(element_type=ElementType.TABLE, content=table)]
        )

        doc = IBDocumentRenderer().render(model)
        content_table = self._find_content_table(doc, "매출총이익")
        assert content_table is not None

        assert content_table.rows[1].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert content_table.rows[1].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT
        assert content_table.rows[1].cells[2].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT

    def test_digit_mixed_codes_stay_left_aligned_when_column_is_textual(self):
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        from ib_renderer import IBDocumentRenderer

        table = Table(
            rows=[
                TableRow(
                    cells=[
                        TableCell(content="코드"),
                        TableCell(content="수량"),
                    ],
                    is_header=True,
                ),
                TableRow(cells=[TableCell(content="A-101"), TableCell(content="12")]),
                TableRow(cells=[TableCell(content="B-102"), TableCell(content="15")]),
                TableRow(cells=[TableCell(content="C-103"), TableCell(content="18")]),
            ],
            col_count=2,
        )
        model = DocumentModel(
            elements=[Element(element_type=ElementType.TABLE, content=table)]
        )

        doc = IBDocumentRenderer().render(model)
        content_table = self._find_content_table(doc, "A-101")
        assert content_table is not None

        assert content_table.rows[1].cells[0].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.LEFT
        assert content_table.rows[1].cells[1].paragraphs[0].alignment == WD_ALIGN_PARAGRAPH.RIGHT


# ═══════════════════════════════════════════════════════════════════════════════
# LIST INDENT TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestListIndentRendering:
    """Tests that deep list nesting stays within a bounded indentation range."""

    def test_deep_bullet_indent_is_compressed(self):
        from ib_renderer import IBDocumentRenderer, STYLE

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.BULLET_LIST,
                    content=ListItem(text="Deep bullet", indent_level=5),
                )
            ]
        )

        doc = IBDocumentRenderer().render(model)
        paragraph = next(p for p in doc.paragraphs if "Deep bullet" in p.text)

        assert paragraph.paragraph_format.left_indent is not None
        assert paragraph.paragraph_format.left_indent.inches < 1.5
        assert paragraph.paragraph_format.left_indent.inches > 1.0
        assert paragraph.paragraph_format.first_line_indent == -STYLE.BULLET_INDENT

    def test_numbered_indent_keeps_depth_without_linear_growth(self):
        from ib_renderer import ListRenderer

        level_three = ListRenderer._resolve_indent(3).inches
        level_five = ListRenderer._resolve_indent(5).inches
        level_six = ListRenderer._resolve_indent(6).inches

        assert level_three == 1.0
        assert level_five > level_three
        assert level_five < 1.5
        assert level_six > level_five
        assert (level_six - level_five) < 0.25


# ═══════════════════════════════════════════════════════════════════════════════
# SEPARATOR RENDERING TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestSeparatorRendering:
    """Tests that SEPARATOR elements produce a visible horizontal rule."""

    def test_separator_creates_paragraph(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(element_type=ElementType.PARAGRAPH, content=Paragraph(text="Before")),
                Element(element_type=ElementType.SEPARATOR, content=None),
                Element(element_type=ElementType.PARAGRAPH, content=Paragraph(text="After")),
            ]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        # Should have paragraphs for: cover, TOC, "Before", separator, "After", disclaimer
        all_text = [p.text for p in doc.paragraphs]
        assert "Before" in all_text
        assert "After" in all_text

    def test_separator_has_border(self):
        from lxml import etree

        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[Element(element_type=ElementType.SEPARATOR, content=None)]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        # Find a paragraph with bottom border
        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        found_border = False
        for p in doc.paragraphs:
            borders = p._p.findall(".//w:pBdr/w:bottom", ns)
            if borders:
                found_border = True
                break
        assert found_border, "SEPARATOR should produce a paragraph with bottom border"

    def test_separator_adds_semantic_bookmark(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[Element(element_type=ElementType.SEPARATOR, content=None)]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        assert any(name.startswith("_ibrep_SEPARATOR_") for name in _semantic_bookmark_names(doc))

    def test_explicit_separator_marker_becomes_page_break_in_auto_mode(self):
        from ib_renderer import IBDocumentRenderer

        baseline_doc = IBDocumentRenderer().render(DocumentModel())
        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.SEPARATOR,
                    content=None,
                    raw_text="## ---",
                )
            ]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        baseline_page_breaks = baseline_doc.element.xpath(
            ".//*[local-name()='br' and @*[local-name()='type']='page']"
        )
        doc_page_breaks = doc.element.xpath(".//*[local-name()='br' and @*[local-name()='type']='page']")

        baseline_borders = [
            paragraph for paragraph in baseline_doc.paragraphs if paragraph._p.xpath(".//*[local-name()='pBdr']")
        ]
        doc_borders = [paragraph for paragraph in doc.paragraphs if paragraph._p.xpath(".//*[local-name()='pBdr']")]

        assert len(doc_page_breaks) == len(baseline_page_breaks) + 1
        assert len(doc_borders) == len(baseline_borders)

    def test_page_break_separator_adds_semantic_bookmark(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.SEPARATOR,
                    content=None,
                    raw_text="## ---",
                )
            ]
        )
        doc = IBDocumentRenderer().render(model)

        assert any(name.startswith("_ibrep_SEPARATOR_") for name in _semantic_bookmark_names(doc))

    def test_separator_mode_page_break_overrides_plain_separator(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[Element(element_type=ElementType.SEPARATOR, content=None, raw_text="---")]
        )
        renderer = IBDocumentRenderer(separator_mode="page-break")
        doc = renderer.render(model)

        page_breaks = doc.element.xpath(".//*[local-name()='br' and @*[local-name()='type']='page']")
        assert len(page_breaks) >= 1

    def test_empty_element_no_crash(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[Element(element_type=ElementType.EMPTY, content=None)]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)
        # Should not crash — that's the test
        assert doc is not None


class TestTOCPreviewRendering:
    """Tests that TOC preview entries are visible before field update."""

    def test_toc_preview_contains_heading_entries(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            metadata=DocumentMetadata(title="샘플 보고서"),
            elements=[
                Element(
                    element_type=ElementType.HEADING_1,
                    content=Heading(level=1, text="샘플 보고서"),
                ),
                Element(
                    element_type=ElementType.HEADING_2,
                    content=Heading(level=2, text="I. 개요"),
                ),
                Element(
                    element_type=ElementType.HEADING_3,
                    content=Heading(level=3, text="1. 세부 항목"),
                ),
            ],
        )

        doc = IBDocumentRenderer().render(model)
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        assert "TABLE OF CONTENTS" in texts
        assert "샘플 보고서" in texts
        assert "I. 개요" in texts
        assert "1. 세부 항목" in texts
        assert not any("Update Field" in text for text in texts)

    def test_toc_title_and_preview_use_toc_font(self):
        from ib_renderer import IBDocumentRenderer, STYLE

        model = DocumentModel(
            metadata=DocumentMetadata(title="샘플 보고서"),
            elements=[
                Element(
                    element_type=ElementType.HEADING_1,
                    content=Heading(level=1, text="샘플 보고서"),
                ),
                Element(
                    element_type=ElementType.HEADING_2,
                    content=Heading(level=2, text="I. 개요"),
                ),
            ],
        )

        doc = IBDocumentRenderer().render(model)
        toc_title = next(p for p in doc.paragraphs if p.text.strip() == "TABLE OF CONTENTS")
        preview_entry = next(p for p in doc.paragraphs if p.text.strip() == "I. 개요")

        assert toc_title.runs[0].font.name == STYLE.TOC_FONT
        assert preview_entry.runs[0].font.name == STYLE.TOC_FONT

    def test_word_toc_styles_use_toc_font(self):
        from docx import Document

        from ib_renderer import DocumentStyler, STYLE

        doc = Document()
        DocumentStyler(doc).create_styles()

        assert doc.styles["TOC 1"].font.name == STYLE.TOC_FONT
        assert doc.styles["TOC 2"].font.name == STYLE.TOC_FONT
        assert doc.styles["TOC 3"].font.name == STYLE.TOC_FONT
        assert doc.styles["TOC 4"].font.name == STYLE.TOC_FONT

    def test_code_block_renders_as_shaded_monospace_block(self):
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.CODE_BLOCK,
                    content=CodeBlock(code="A\n└── B", language="text"),
                )
            ]
        )

        doc = IBDocumentRenderer().render(model)
        code_table = next((table for table in doc.tables if "└── B" in table.cell(0, 0).text), None)

        assert code_table is not None
        cell = code_table.cell(0, 0)
        run = cell.paragraphs[0].runs[0]
        assert run.font.name == "Consolas"
        assert cell.vertical_alignment == WD_CELL_VERTICAL_ALIGNMENT.TOP
        assert not cell._tc.xpath(".//*[local-name()='tblBorders']/*[@*[local-name()='val']='single']")

    def test_code_block_adds_semantic_bookmark(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.CODE_BLOCK,
                    content=CodeBlock(code="print('alpha')", language="python"),
                )
            ]
        )

        doc = IBDocumentRenderer().render(model)

        assert any(name.startswith("_ibrep_CODE_BLOCK_python_") for name in _semantic_bookmark_names(doc))

    def test_diagram_adds_semantic_bookmark(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.DIAGRAM,
                    content=Diagram(
                        diagram_type="flow",
                        title="Approval Flow",
                        boxes=[DiagramBox(id="start", label="Start", pos=[0, 0])],
                    ),
                )
            ]
        )

        doc = IBDocumentRenderer().render(model)

        assert any(name.startswith("_ibrep_DIAGRAM_flow_") for name in _semantic_bookmark_names(doc))


class TestCoverMetadataInference:
    """Tests that inferred analysis metadata appears on the cover panel."""

    def test_cover_metadata_panel_renders_inferred_analysis_rows(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            metadata=DocumentMetadata(
                title="일동제약 주식회사 수익성 변화 분석 보고서",
                company="일동제약 주식회사",
                extra={
                    "date": "2026년 3월 20일",
                    "analysis_period": "제9기→제10기",
                    "analysis_basis": "연결재무제표 기준",
                },
            )
        )

        doc = IBDocumentRenderer().render(model)
        cover_table = doc.tables[0]
        rows = [[cell.text.strip() for cell in row.cells] for row in cover_table.rows]

        assert ["REPORT DATE", "2026년 3월 20일"] in rows
        assert ["ANALYSIS PERIOD", "제9기→제10기"] in rows
        assert ["ANALYSIS BASIS", "연결재무제표 기준"] in rows
        assert ["INSTITUTION", "일동제약 주식회사"] in rows
        assert not any(row[0] == "PREPARED BY" for row in rows)
        assert not any(row[0] == "SECTOR" for row in rows)

    def test_cover_uses_cover_font_for_title_and_metadata(self):
        from ib_renderer import IBDocumentRenderer, STYLE

        model = DocumentModel(
            metadata=DocumentMetadata(
                title="주식회사 웅진 자회사 보유구조 및 영업관계 분석",
                company="주식회사 웅진",
                extra={"date": "2026년 3월 23일"},
            )
        )

        doc = IBDocumentRenderer().render(model)
        title_paragraph = next(
            p for p in doc.paragraphs if p.text.strip() == "자회사 보유구조 및 영업관계 분석"
        )
        metadata_table = doc.tables[0]
        label_run = metadata_table.rows[0].cells[0].paragraphs[0].runs[0]
        value_run = metadata_table.rows[0].cells[1].paragraphs[0].runs[0]

        assert title_paragraph.runs[0].font.name == STYLE.COVER_FONT
        assert label_run.font.name == STYLE.COVER_FONT
        assert value_run.font.name == STYLE.COVER_FONT

    def test_cover_hides_placeholder_sector_and_splits_company_from_title(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            metadata=DocumentMetadata(
                title="일동제약 주식회사 수익성 변화 분석 보고서",
                company="일동제약 주식회사",
                sector="SECTOR",
            )
        )

        doc = IBDocumentRenderer().render(model)
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        assert "SECTOR" not in texts
        assert "DCM RESEARCH" not in texts
        assert "일동제약 주식회사" in texts
        assert "수익성 변화 분석 보고서" in texts
        assert "일동제약 주식회사 수익성 변화 분석 보고서" not in texts

    def test_subject_company_only_affects_cover_not_disclaimer_brand(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            metadata=DocumentMetadata(
                title="일동제약 주식회사 수익성 변화 분석 보고서",
                company="Korea Development Bank",
                extra={"subject_company": "일동제약 주식회사"},
            )
        )

        doc = IBDocumentRenderer().render(model)
        cover_rows = [[cell.text.strip() for cell in row.cells] for row in doc.tables[0].rows]
        texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        assert ["INSTITUTION", "일동제약 주식회사"] in cover_rows
        assert any("Korea Development Bank. All rights reserved." in text for text in texts)


# ═══════════════════════════════════════════════════════════════════════════════
# CALLOUT CONTENT FORMATTING TESTS
# ═══════════════════════════════════════════════════════════════════════════════


class TestCalloutContentFormatting:
    """Tests that callout content renders with inline formatting."""

    def test_callout_bold_content(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.BLOCKQUOTE,
                    content=Blockquote(
                        title="KEY INSIGHT",
                        text="This is **important** information.",
                    ),
                )
            ]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        # Find the callout table
        callout_table = None
        for t in doc.tables:
            if len(t.rows) == 1 and len(t.rows[0].cells) == 1:
                cell_text = t.rows[0].cells[0].text
                if "important" in cell_text:
                    callout_table = t
                    break

        assert callout_table is not None
        cell = callout_table.rows[0].cells[0]

        # Content paragraph should have formatted runs
        content_paras = [p for p in cell.paragraphs if "important" in p.text]
        assert len(content_paras) >= 1

        bold_runs = [r for r in content_paras[0].runs if r.font.bold]
        assert len(bold_runs) >= 1, "Bold markers in callout content should be rendered"

    def test_callout_italic_content(self):
        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.BLOCKQUOTE,
                    content=Blockquote(
                        title="NOTE",
                        text="See *appendix* for details.",
                    ),
                )
            ]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        for t in doc.tables:
            if len(t.rows) == 1 and len(t.rows[0].cells) == 1:
                cell = t.rows[0].cells[0]
                for p in cell.paragraphs:
                    if "appendix" in p.text:
                        italic_runs = [r for r in p.runs if r.font.italic]
                        assert len(italic_runs) >= 1, "Italic in callout should render"
                        return
        pytest.fail("Callout with italic content not found")

    def test_callout_navy_bg_white_text(self):
        from docx.shared import RGBColor

        from ib_renderer import IBDocumentRenderer

        model = DocumentModel(
            elements=[
                Element(
                    element_type=ElementType.BLOCKQUOTE,
                    content=Blockquote(
                        title="EXECUTIVE SUMMARY",
                        text="**Key finding**: revenue up.",
                    ),
                )
            ]
        )
        renderer = IBDocumentRenderer()
        doc = renderer.render(model)

        white = RGBColor(255, 255, 255)
        for t in doc.tables:
            if len(t.rows) == 1 and len(t.rows[0].cells) == 1:
                cell = t.rows[0].cells[0]
                for p in cell.paragraphs:
                    if "finding" in p.text or "revenue" in p.text:
                        for r in p.runs:
                            if r.font.color.rgb == white:
                                return  # Found white text — pass
        pytest.fail("Executive summary content should have white text")
