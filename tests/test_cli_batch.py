"""Batch-mode smoke tests for the CLI entry points."""

from types import SimpleNamespace

from docx import Document

from md_to_word import run_batch_conversion as run_md_batch_conversion
from word_to_md import run_batch_conversion as run_word_batch_conversion


def test_md_to_word_batch_conversion_creates_outputs(tmp_path):
    """Batch mode should convert every markdown file in a directory."""
    input_dir = tmp_path / "md_inputs"
    output_dir = tmp_path / "docx_outputs"
    input_dir.mkdir()

    (input_dir / "one.md").write_text("# One\n\n본문입니다.\n", encoding="utf-8")
    (input_dir / "two.md").write_text("# Two\n\n둘째 문서입니다.\n", encoding="utf-8")

    args = SimpleNamespace(
        output_file=str(output_dir),
        no_cover=True,
        no_toc=True,
        no_disclaimer=True,
        format=False,
        deepresearch_cleaner="off",
        cite_mode="footnote",
        drop_unknown_markers=False,
        cleaner_report=False,
        verbose=False,
        batch=True,
    )

    exit_code = run_md_batch_conversion(input_dir, args)

    assert exit_code == 0
    assert len(list(output_dir.glob("*.docx"))) == 2


def test_word_to_md_batch_conversion_creates_outputs(tmp_path):
    """Batch mode should convert every Word file in a directory."""
    input_dir = tmp_path / "docx_inputs"
    output_dir = tmp_path / "md_outputs"
    input_dir.mkdir()

    for name in ("one", "two"):
        doc = Document()
        doc.add_heading(name.title(), level=1)
        doc.add_paragraph("본문입니다.")
        doc.save(str(input_dir / f"{name}.docx"))

    args = SimpleNamespace(
        output_file=str(output_dir),
        strip=False,
        no_frontmatter=True,
        extract_images=False,
        image_dir=None,
        verbose=False,
        batch=True,
    )

    exit_code = run_word_batch_conversion(input_dir, args)

    assert exit_code == 0
    assert len(list(output_dir.glob("*.md"))) == 2
