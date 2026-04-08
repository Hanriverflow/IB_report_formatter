"""
Word Parser Module for Word to Markdown Converter
Handles parsing of .docx files into DocumentModel.

Dependencies:
    Required: python-docx
    Optional: Pillow (for image processing)
"""

import base64
import io
import logging
import mimetypes
import re
import zipfile
from dataclasses import dataclass, field, replace
from enum import Enum, auto
from pathlib import Path
from typing import Dict, Iterator, List, Optional, Set, Tuple, Union
from xml.etree import ElementTree

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from md_parser import (
    Blockquote,
    DocumentMetadata,
    DocumentModel,
    Element,
    ElementType,
    Heading,
    Image,
    LaTeXEquation,
    ListItem,
    Paragraph,
    Table,
    TableParser,
    TableRow,
    TextRun,
)

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# PARSE CONTEXT
# ═══════════════════════════════════════════════════════════════════════════════


class DocumentProfile(Enum):
    """Document parsing profiles."""

    IB_GENERATED = auto()
    GENERIC = auto()


@dataclass
class ParseContext:
    """State shared across Word parsing passes."""

    profile: DocumentProfile = DocumentProfile.GENERIC
    skip_indices: Set[int] = field(default_factory=set)
    footnotes: Dict[int, str] = field(default_factory=dict)


@dataclass(frozen=True)
class ExtractedImageAsset:
    """Parsed image data keyed by Word relationship ID."""

    path: str = ""
    base64_data: Optional[str] = None
    mime_type: str = "image/png"
    default_alt_text: str = "Image"


@dataclass(frozen=True)
class ThemeColorResolver:
    """Resolves Word theme color references to concrete RGB hex strings."""

    color_map: Dict[str, str] = field(default_factory=dict)

    _THEME_PART_PATH = "word/theme/theme1.xml"
    _DRAWINGML_NS = {"a": "http://schemas.openxmlformats.org/drawingml/2006/main"}
    _ALIAS_MAP = {
        "dk1": "dark1",
        "lt1": "light1",
        "dk2": "dark2",
        "lt2": "light2",
        "hlink": "hyperlink",
        "folhlink": "followedhyperlink",
        "text1": "dark1",
        "background1": "light1",
        "text2": "dark2",
        "background2": "light2",
        "followedhyperlink": "followedhyperlink",
    }

    @classmethod
    def from_docx_bytes(cls, source_bytes: bytes) -> "ThemeColorResolver":
        """Parse a DOCX theme part into a theme-color lookup table."""
        try:
            with zipfile.ZipFile(io.BytesIO(source_bytes)) as archive:
                if cls._THEME_PART_PATH not in archive.namelist():
                    return cls()
                xml_bytes = archive.read(cls._THEME_PART_PATH)
        except (OSError, KeyError, zipfile.BadZipFile):
            return cls()

        try:
            root = ElementTree.fromstring(xml_bytes)
        except ElementTree.ParseError:
            return cls()

        clr_scheme = root.find(".//a:clrScheme", cls._DRAWINGML_NS)
        if clr_scheme is None:
            return cls()

        color_map: Dict[str, str] = {}
        for child in list(clr_scheme):
            local_name = cls._local_name(child.tag).lower()
            value = cls._extract_scheme_color(child)
            if not local_name or value is None:
                continue
            normalized_name = cls._normalize_theme_name(local_name)
            color_map[normalized_name] = value

        return cls(color_map=color_map)

    def resolve(
        self,
        theme_color: str,
        tint: Optional[str] = None,
        shade: Optional[str] = None,
    ) -> Optional[str]:
        """Resolve a theme color reference, optionally applying tint/shade."""
        if not theme_color:
            return None

        normalized_name = self._normalize_theme_name(theme_color)
        base_hex = self.color_map.get(normalized_name)
        if base_hex is None:
            return None

        rgb = self._hex_to_rgb(base_hex)
        if tint:
            rgb = self._apply_tint(rgb, tint)
        if shade:
            rgb = self._apply_shade(rgb, shade)
        return self._rgb_to_hex(rgb)

    @classmethod
    def _extract_scheme_color(cls, element) -> Optional[str]:
        """Extract a color value from a theme scheme node."""
        for child in list(element):
            local_name = cls._local_name(child.tag).lower()
            if local_name == "srgbclr":
                value = (child.attrib.get("val") or "").strip()
                return cls._normalize_hex(value)
            if local_name == "sysclr":
                value = (child.attrib.get("lastClr") or child.attrib.get("val") or "").strip()
                return cls._normalize_hex(value)
        return None

    @classmethod
    def _normalize_theme_name(cls, value: str) -> str:
        """Normalize theme color names and aliases to a common key."""
        normalized = value.strip().replace("_", "").replace("-", "").lower()
        return cls._ALIAS_MAP.get(normalized, normalized)

    @staticmethod
    def _local_name(tag: str) -> str:
        """Strip namespace from an XML tag."""
        return tag.split("}", 1)[-1] if "}" in tag else tag

    @staticmethod
    def _normalize_hex(value: str) -> Optional[str]:
        """Normalize a raw hex string to #RRGGBB."""
        candidate = value.strip().lstrip("#")
        if len(candidate) != 6 or not re.fullmatch(r"[0-9A-Fa-f]{6}", candidate):
            return None
        return f"#{candidate.upper()}"

    @staticmethod
    def _hex_to_rgb(value: str) -> Tuple[int, int, int]:
        """Convert #RRGGBB into an RGB tuple."""
        normalized = value.lstrip("#")
        return (
            int(normalized[0:2], 16),
            int(normalized[2:4], 16),
            int(normalized[4:6], 16),
        )

    @staticmethod
    def _rgb_to_hex(rgb: Tuple[int, int, int]) -> str:
        """Convert an RGB tuple into #RRGGBB."""
        return "#{:02X}{:02X}{:02X}".format(*rgb)

    @staticmethod
    def _apply_tint(rgb: Tuple[int, int, int], tint: str) -> Tuple[int, int, int]:
        """Apply Word theme tint to an RGB tuple."""
        tint_value = int(tint, 16) / 255.0
        return tuple(
            max(0, min(255, int(round(channel + (255 - channel) * tint_value))))
            for channel in rgb
        )

    @staticmethod
    def _apply_shade(rgb: Tuple[int, int, int], shade: str) -> Tuple[int, int, int]:
        """Apply Word theme shade to an RGB tuple."""
        shade_value = int(shade, 16) / 255.0
        return tuple(
            max(0, min(255, int(round(channel * shade_value))))
            for channel in rgb
        )


class NumberingTracker:
    """Track inferred numbering for Word lists whose numbers are formatting-only."""

    def __init__(self):
        self._counters: Dict[str, Dict[int, int]] = {}
        self._level_formats: Dict[str, Dict[int, str]] = {}

    def next_number(self, para, indent_level: int) -> str:
        """Return the next numbering label for a paragraph."""
        list_key = StyleDetector.get_numbering_key(para, indent_level)
        if list_key not in self._counters:
            self._counters[list_key] = {}
            self._level_formats[list_key] = {}

        counters = self._counters[list_key]
        level_formats = self._level_formats[list_key]
        for level in list(counters.keys()):
            if level > indent_level:
                del counters[level]

        counters[indent_level] = counters.get(indent_level, 0) + 1
        num_fmt = StyleDetector.resolve_numbering_format(para)
        if num_fmt:
            level_formats[indent_level] = num_fmt
        elif indent_level not in level_formats:
            level_formats[indent_level] = "decimal"

        return ".".join(
            self._format_value(counters[level], level_formats.get(level, "decimal"))
            for level in sorted(counters)
            if level <= indent_level
        )

    def break_sequence(self) -> None:
        """Reset inferred numbering between disconnected list blocks."""
        self._counters.clear()
        self._level_formats.clear()

    @staticmethod
    def _format_value(value: int, num_fmt: str) -> str:
        """Format a list counter using Word numbering semantics."""
        if num_fmt == "decimal":
            return str(value)
        if num_fmt == "upperLetter":
            return NumberingTracker._to_alpha(value).upper()
        if num_fmt == "lowerLetter":
            return NumberingTracker._to_alpha(value).lower()
        if num_fmt == "upperRoman":
            return NumberingTracker._to_roman(value).upper()
        if num_fmt == "lowerRoman":
            return NumberingTracker._to_roman(value).lower()
        return str(value)

    @staticmethod
    def _to_alpha(value: int) -> str:
        """Convert 1-based integer to alphabetic sequence."""
        result = []
        current = value
        while current > 0:
            current -= 1
            result.append(chr(ord("A") + (current % 26)))
            current //= 26
        return "".join(reversed(result)) or "A"

    @staticmethod
    def _to_roman(value: int) -> str:
        """Convert 1-based integer to Roman numeral sequence."""
        numerals = [
            (1000, "M"),
            (900, "CM"),
            (500, "D"),
            (400, "CD"),
            (100, "C"),
            (90, "XC"),
            (50, "L"),
            (40, "XL"),
            (10, "X"),
            (9, "IX"),
            (5, "V"),
            (4, "IV"),
            (1, "I"),
        ]
        remaining = value
        result = []
        for number, symbol in numerals:
            while remaining >= number:
                result.append(symbol)
                remaining -= number
        return "".join(result) or "I"


# ═══════════════════════════════════════════════════════════════════════════════
# METADATA EXTRACTOR
# ═══════════════════════════════════════════════════════════════════════════════


class MetadataExtractor:
    """Extracts document metadata from Word core properties and IB cover blocks."""

    _CUSTOM_PROPS_PATH = "docProps/custom.xml"
    _CUSTOM_PROPS_NS = {
        "cp": "http://schemas.openxmlformats.org/officeDocument/2006/custom-properties",
    }
    _IB_PANEL_FIELDS = {
        "REPORT DATE": ("extra", "date"),
        "PREPARED BY": ("analyst", None),
        "INSTITUTION": ("company", None),
        "SECTOR": ("sector", None),
        "PREPARED FOR": ("extra", "recipient"),
    }
    _IB_REPORT_TYPE_RE = re.compile(r"^[A-Z][A-Z0-9\s&/\-]{3,40}$")
    _NON_ALNUM_RE = re.compile(r"[^a-z0-9]+")
    _CUSTOM_METADATA_FIELD_MAP = {
        "title": ("title", None),
        "report_title": ("title", None),
        "subtitle": ("subtitle", None),
        "subject": ("subtitle", None),
        "company": ("company", None),
        "institution": ("company", None),
        "issuer": ("company", None),
        "ticker": ("ticker", None),
        "sector": ("sector", None),
        "analyst": ("analyst", None),
        "author": ("analyst", None),
        "date": ("extra", "date"),
        "report_date": ("extra", "date"),
        "recipient": ("extra", "recipient"),
        "prepared_for": ("extra", "recipient"),
        "report_type": ("extra", "report_type"),
    }

    @classmethod
    def extract(
        cls,
        doc: DocxDocument,
        source_bytes: Optional[bytes] = None,
    ) -> DocumentMetadata:
        """Extract metadata from document properties."""
        props = doc.core_properties
        metadata = DocumentMetadata()

        if props.title:
            metadata.title = props.title
        if props.author:
            metadata.analyst = props.author
        if props.subject:
            metadata.subtitle = props.subject
        if props.created:
            metadata.extra["date"] = props.created.strftime("%Y-%m-%d")
        if props.category:
            metadata.sector = props.category

        if source_bytes:
            cls._apply_custom_metadata(metadata, cls.extract_custom_properties(source_bytes))

        return metadata

    @classmethod
    def extract_custom_properties(cls, source_bytes: bytes) -> Dict[str, str]:
        """Extract custom document properties from docProps/custom.xml."""
        try:
            with zipfile.ZipFile(io.BytesIO(source_bytes)) as archive:
                if cls._CUSTOM_PROPS_PATH not in archive.namelist():
                    return {}
                xml_bytes = archive.read(cls._CUSTOM_PROPS_PATH)
        except (OSError, KeyError, zipfile.BadZipFile):
            return {}

        try:
            root = ElementTree.fromstring(xml_bytes)
        except ElementTree.ParseError:
            return {}

        properties: Dict[str, str] = {}
        prop_tag = "{%s}property" % cls._CUSTOM_PROPS_NS["cp"]
        for prop in root.findall(prop_tag):
            name = (prop.attrib.get("name") or "").strip()
            value = cls._extract_custom_property_value(prop)
            if name and value:
                properties[name] = value

        return properties

    @classmethod
    def _apply_custom_metadata(
        cls,
        metadata: DocumentMetadata,
        custom_properties: Dict[str, str],
    ) -> None:
        """Merge custom document properties into the shared metadata model."""
        for raw_key, value in custom_properties.items():
            normalized_key = cls._normalize_custom_key(raw_key)
            mapped = cls._CUSTOM_METADATA_FIELD_MAP.get(normalized_key)

            if mapped:
                field_name, extra_key = mapped
                if field_name == "extra" and extra_key:
                    metadata.extra[extra_key] = value
                else:
                    setattr(metadata, field_name, value)
                continue

            metadata.extra.setdefault(normalized_key, value)

    @staticmethod
    def _extract_custom_property_value(prop) -> str:
        """Read the text value from a custom property element."""
        for child in prop:
            if isinstance(child.tag, str):
                return "".join(child.itertext()).strip()
        return ""

    @classmethod
    def _normalize_custom_key(cls, key: str) -> str:
        """Normalize a custom metadata key to snake_case."""
        lowered = key.strip().lower()
        normalized = cls._NON_ALNUM_RE.sub("_", lowered).strip("_")
        return normalized

    @classmethod
    def apply_ib_cover_metadata(
        cls,
        metadata: DocumentMetadata,
        blocks: List[Union[DocxParagraph, DocxTable]],
        metadata_panel_index: Optional[int],
    ) -> None:
        """Recover metadata from an IB-generated cover page."""
        panel_values: Dict[str, str] = {}
        if metadata_panel_index is not None:
            table = blocks[metadata_panel_index]
            if isinstance(table, DocxTable):
                panel_values = cls._extract_metadata_panel_values(table)
                cls._apply_metadata_panel_values(metadata, panel_values)

        if metadata_panel_index is None:
            return

        cover_texts = cls._collect_cover_texts(blocks, metadata_panel_index)
        if not cover_texts:
            return

        report_type = next((text for text in cover_texts if cls._IB_REPORT_TYPE_RE.match(text)), "")
        if report_type:
            metadata.extra.setdefault("report_type", report_type)

        company = panel_values.get("INSTITUTION", metadata.company).strip()
        sector = panel_values.get("SECTOR", metadata.sector).strip()
        identity = cls._find_identity_line(cover_texts, company, sector, report_type)

        title, subtitle = cls._find_title_and_subtitle(
            cover_texts=cover_texts,
            report_type=report_type,
            identity=identity,
            sector=sector,
        )

        if title:
            metadata.title = title
        if subtitle:
            metadata.subtitle = subtitle
        if identity and company and identity != company:
            metadata.ticker = identity

    @classmethod
    def _extract_metadata_panel_values(cls, table: DocxTable) -> Dict[str, str]:
        """Extract key/value rows from the IB cover metadata panel."""
        values: Dict[str, str] = {}
        if len(table.columns) != 2:
            return values

        for row in table.rows:
            if len(row.cells) < 2:
                continue
            label = row.cells[0].text.strip().upper()
            value = row.cells[1].text.strip()
            if label and value:
                values[label] = value
        return values

    @classmethod
    def _apply_metadata_panel_values(
        cls, metadata: DocumentMetadata, panel_values: Dict[str, str]
    ) -> None:
        """Merge metadata panel values into DocumentMetadata."""
        for label, value in panel_values.items():
            target = cls._IB_PANEL_FIELDS.get(label)
            if not target:
                continue
            field_name, extra_key = target
            if field_name == "extra" and extra_key:
                metadata.extra[extra_key] = value
            else:
                setattr(metadata, field_name, value)

    @classmethod
    def _collect_cover_texts(
        cls,
        blocks: List[Union[DocxParagraph, DocxTable]],
        metadata_panel_index: int,
    ) -> List[str]:
        """Collect meaningful cover-paragraph text before the metadata panel."""
        texts: List[str] = []
        for block in blocks[:metadata_panel_index]:
            if not isinstance(block, DocxParagraph):
                continue
            text = block.text.strip()
            if text:
                texts.append(text)
        return texts

    @staticmethod
    def _find_identity_line(
        cover_texts: List[str],
        company: str,
        sector: str,
        report_type: str,
    ) -> str:
        """Find the cover identity line, usually ticker or company identifier."""
        excluded = {company.upper(), sector.upper(), report_type.upper()}
        for text in cover_texts:
            if text.upper() in excluded:
                continue
            if len(text) <= 24 and text == text.upper():
                return text
        return ""

    @staticmethod
    def _find_title_and_subtitle(
        cover_texts: List[str],
        report_type: str,
        identity: str,
        sector: str,
    ) -> Tuple[str, str]:
        """Infer title and subtitle from cover paragraphs."""
        filtered = [
            text
            for text in cover_texts
            if text not in {report_type, identity, sector}
        ]
        if not filtered:
            return "", ""

        title = max(filtered, key=len)
        subtitle = ""
        title_index = filtered.index(title)
        if title_index + 1 < len(filtered):
            candidate = filtered[title_index + 1]
            if candidate != title and len(candidate) <= 120:
                subtitle = candidate
        return title, subtitle


# ═══════════════════════════════════════════════════════════════════════════════
# STYLE DETECTOR
# ═══════════════════════════════════════════════════════════════════════════════


class StyleDetector:
    """Detects paragraph styles and formatting from Word documents."""

    HEADING_PATTERNS = [
        (re.compile(r"^Heading\s*1$", re.I), 1),
        (re.compile(r"^Heading\s*2$", re.I), 2),
        (re.compile(r"^Heading\s*3$", re.I), 3),
        (re.compile(r"^Heading\s*4$", re.I), 4),
        (re.compile(r"^제목\s*1$", re.I), 1),
        (re.compile(r"^제목\s*2$", re.I), 2),
        (re.compile(r"^제목\s*3$", re.I), 3),
    ]

    LIST_BULLET_PATTERNS = [
        re.compile(r"^List\s*Bullet", re.I),
        re.compile(r"^IB\s*Bullet", re.I),
    ]

    LIST_NUMBER_PATTERNS = [
        re.compile(r"^List\s*Number", re.I),
        re.compile(r"^List\s*Paragraph", re.I),
    ]

    LIST_BULLET_STYLE_IDS = {"ListBullet", "ListBullet2", "ListBullet3"}
    LIST_NUMBER_STYLE_IDS = {
        "ListNumber",
        "ListNumber2",
        "ListNumber3",
        "ListParagraph",
    }

    BULLET_TEXT_PATTERN = re.compile(r"^[•●■◦▪\-*]\s+(.+)$")
    NUMBERED_TEXT_PATTERN = re.compile(r"^(\d+)[\.\)]\s+(.+)$")
    LIST_LEVEL_STYLE_PATTERN = re.compile(r".*?(\d+)$")
    HEADING_LEVEL_HINT_RE = re.compile(r"(?:heading|제목)\s*([1-4])", re.I)

    @classmethod
    def detect_heading_level(cls, para) -> Optional[int]:
        """Detect if paragraph is a heading and return level (1-4)."""
        style_level = cls._detect_heading_level_from_style(para.style if para.style else None)
        if style_level:
            return style_level

        outline_level = cls._extract_outline_level(para)
        if outline_level:
            return outline_level

        if cls._is_heading_by_formatting(para):
            return 2
        return None

    @classmethod
    def detect_list_type(cls, para) -> Optional[str]:
        """Detect if paragraph is a list. Returns 'bullet' or 'number'."""
        num_fmt = cls.resolve_numbering_format(para)
        if num_fmt:
            return "bullet" if num_fmt == "bullet" else "number"

        style_name = para.style.name if para.style else ""
        style_id = para.style.style_id if para.style else ""

        if style_id in cls.LIST_BULLET_STYLE_IDS:
            return "bullet"
        if style_id in cls.LIST_NUMBER_STYLE_IDS:
            return "number"

        for pattern in cls.LIST_BULLET_PATTERNS:
            if pattern.match(style_name):
                return "bullet"
        for pattern in cls.LIST_NUMBER_PATTERNS:
            if pattern.match(style_name):
                return "number"

        text = para.text.strip()
        if cls.BULLET_TEXT_PATTERN.match(text):
            return "bullet"
        if cls.NUMBERED_TEXT_PATTERN.match(text):
            return "number"
        return None

    @classmethod
    def detect_list_level(cls, para) -> int:
        """Infer nested list level from numbering, style, or indentation."""
        level = cls._extract_numbering_level(para)
        if level is not None:
            return level

        style_name = para.style.name if para.style else ""
        style_match = cls.LIST_LEVEL_STYLE_PATTERN.match(style_name)
        if style_match:
            try:
                return max(0, int(style_match.group(1)) - 1)
            except ValueError:
                pass

        left_indent = para.paragraph_format.left_indent
        if left_indent is not None and left_indent.pt:
            approx_level = int(round(max(0.0, left_indent.pt) / 18.0)) - 1
            return max(0, approx_level)

        return 0

    @classmethod
    def extract_bullet_text(cls, text: str) -> str:
        """Remove the bullet marker from a paragraph's visible text."""
        match = cls.BULLET_TEXT_PATTERN.match(text.strip())
        if match:
            return match.group(1)
        return text.strip()

    @classmethod
    def extract_numbered_text(cls, text: str) -> Tuple[str, str]:
        """Extract list number and item text from a numbered list paragraph."""
        match = cls.NUMBERED_TEXT_PATTERN.match(text.strip())
        if match:
            return match.group(1), match.group(2)
        return "1", text.strip()

    @staticmethod
    def is_centered_caption(para) -> bool:
        """Return True if a paragraph looks like an image caption."""
        text = para.text.strip()
        if not text or len(text) > 120:
            return False
        alignment = para.alignment
        return bool(alignment == WD_ALIGN_PARAGRAPH.CENTER)

    @staticmethod
    def has_embedded_content(para) -> bool:
        """Return True if a paragraph contains non-text embedded content like images."""
        return bool(
            para._p.xpath(
                ".//*[local-name()='drawing' or local-name()='pict' or local-name()='object']"
            )
        )

    @staticmethod
    def _extract_numbering_level(para) -> Optional[int]:
        """Extract Word numbering level from paragraph XML."""
        for ilvl in para._p.xpath(".//*[local-name()='ilvl']"):
            value = ilvl.get(qn("w:val")) or ilvl.get("val")
            if value and str(value).isdigit():
                return int(value)
        return None

    @classmethod
    def _detect_heading_level_from_style(cls, style) -> Optional[int]:
        """Detect heading level from a style or its base-style chain."""
        for current_style in cls._iter_style_chain(style):
            for identifier in cls._style_identifiers(current_style):
                level = cls._match_heading_level(identifier)
                if level:
                    return level

            outline_level = cls._extract_outline_level_from_element(current_style.element)
            if outline_level:
                return outline_level

        return None

    @staticmethod
    def _iter_style_chain(style) -> Iterator:
        """Yield a style and each base style once."""
        seen: Set[int] = set()
        current = style
        while current is not None and id(current) not in seen:
            seen.add(id(current))
            yield current
            current = getattr(current, "base_style", None)

    @staticmethod
    def _style_identifiers(style) -> List[str]:
        """Return style identifiers that may encode heading semantics."""
        identifiers = []
        for attr in ("name", "style_id"):
            value = getattr(style, attr, "")
            if value:
                identifiers.append(str(value))
        return identifiers

    @classmethod
    def _match_heading_level(cls, identifier: str) -> Optional[int]:
        """Match a heading level from a style identifier string."""
        for pattern, level in cls.HEADING_PATTERNS:
            if pattern.match(identifier):
                return level

        match = cls.HEADING_LEVEL_HINT_RE.search(identifier)
        if match:
            return int(match.group(1))
        return None

    @classmethod
    def _extract_outline_level(cls, para) -> Optional[int]:
        """Extract paragraph or style outline level (0-based in Word)."""
        level = cls._extract_outline_level_from_element(para._p)
        if level:
            return level

        style = para.style if para.style else None
        for current_style in cls._iter_style_chain(style):
            level = cls._extract_outline_level_from_element(current_style.element)
            if level:
                return level
        return None

    @staticmethod
    def _extract_outline_level_from_element(element) -> Optional[int]:
        """Extract heading level from a paragraph/style XML element."""
        if element is None or not hasattr(element, "xpath"):
            return None

        for outline in element.xpath(".//*[local-name()='outlineLvl']"):
            value = outline.get(qn("w:val")) or outline.get("val")
            if value is not None and str(value).isdigit():
                return int(value) + 1
        return None

    @classmethod
    def get_numbering_key(cls, para, indent_level: int) -> str:
        """Build a stable key for a numbered list sequence."""
        for num_id in para._p.xpath(".//*[local-name()='numId']"):
            value = num_id.get(qn("w:val")) or num_id.get("val")
            if value:
                return f"num:{value}"

        style_id = para.style.style_id if para.style else ""
        if style_id:
            return f"style-id:{style_id}"

        style_name = para.style.name if para.style else "generic"
        return f"style-name:{style_name}:level:{indent_level}"

    @classmethod
    def resolve_numbering_format(cls, para) -> Optional[str]:
        """Resolve the Word numbering format for a paragraph if available."""
        numbering_info = cls._extract_numbering_info(para)
        if numbering_info is None:
            return None

        num_id, ilvl = numbering_info
        numbering_part = getattr(para.part, "numbering_part", None)
        if numbering_part is None:
            return None

        numbering_root = numbering_part._element
        abstract_num_id = None
        for num in numbering_root.xpath(".//*[local-name()='num']"):
            current_num_id = num.get(qn("w:numId")) or num.get("numId")
            if str(current_num_id) != num_id:
                continue
            abstract = next(
                (child for child in num if child.tag.endswith("abstractNumId")),
                None,
            )
            if abstract is not None:
                abstract_num_id = abstract.get(qn("w:val")) or abstract.get("val")
                break

        if abstract_num_id is None:
            return None

        for abstract in numbering_root.xpath(".//*[local-name()='abstractNum']"):
            current_id = abstract.get(qn("w:abstractNumId")) or abstract.get("abstractNumId")
            if str(current_id) != str(abstract_num_id):
                continue
            for level in abstract.xpath(".//*[local-name()='lvl']"):
                current_level = level.get(qn("w:ilvl")) or level.get("ilvl")
                if str(current_level) != str(ilvl):
                    continue
                num_fmt = next(
                    (child for child in level if child.tag.endswith("numFmt")),
                    None,
                )
                if num_fmt is not None:
                    value = num_fmt.get(qn("w:val")) or num_fmt.get("val")
                    if value:
                        return str(value)
        return None

    @staticmethod
    def _extract_numbering_info(para) -> Optional[Tuple[str, int]]:
        """Extract numId and ilvl from paragraph XML."""
        num_id = None
        ilvl = 0
        for element in para._p.xpath(".//*[local-name()='numId']"):
            value = element.get(qn("w:val")) or element.get("val")
            if value:
                num_id = str(value)
                break
        if num_id is None:
            return None

        for element in para._p.xpath(".//*[local-name()='ilvl']"):
            value = element.get(qn("w:val")) or element.get("val")
            if value and str(value).isdigit():
                ilvl = int(value)
                break
        return num_id, ilvl

    @classmethod
    def _is_heading_by_formatting(cls, para) -> bool:
        """Heuristic: all-bold short text might be a heading."""
        text = para.text.strip()
        if not text or len(text) > 80 or not para.runs:
            return False

        for run in para.runs:
            if run.text.strip() and not run.bold:
                return False
        return not text.endswith(".")


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE / CALLOUT EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════


class TableExtractor:
    """Extracts Table objects from Word tables."""

    NAVY_HEX = "003366"

    @classmethod
    def extract(
        cls,
        word_table,
        theme_color_resolver: Optional[ThemeColorResolver] = None,
    ) -> Table:
        """Convert a Word table to our Table model."""
        table = Table()
        table.col_count = len(word_table.columns)
        table.alignments = cls._detect_alignments(word_table, table.col_count)

        header_cells = [
            "\n".join(p.text for p in cell.paragraphs).strip()
            for cell in word_table.rows[0].cells
        ] if word_table.rows else []
        table.table_type = TableParser._detect_type(" ".join(header_cells).lower())

        for row_idx, word_row in enumerate(word_table.rows):
            is_header = row_idx == 0
            table_row = TableRow(is_header=is_header)

            for col_idx, cell in enumerate(word_row.cells):
                cell_text = "\n".join(p.text for p in cell.paragraphs).strip()
                parsed_cell = TableParser._parse_cell(
                    text=cell_text,
                    is_header=is_header,
                    col_idx=col_idx,
                    row_idx=row_idx,
                    total_rows=len(word_table.rows),
                    table_type=table.table_type,
                    header_cells=header_cells,
                )
                parsed_cell.runs = cls._extract_cell_runs(
                    cell,
                    theme_color_resolver=theme_color_resolver,
                )
                parsed_cell.alignment = (
                    table.alignments[col_idx] if col_idx < len(table.alignments) else "left"
                )
                table_row.cells.append(parsed_cell)

            table.rows.append(table_row)

        return table

    @classmethod
    def is_header_row_navy(cls, word_table) -> bool:
        """Check if first row has Navy background (IB-style header)."""
        try:
            if not word_table.rows or not word_table.rows[0].cells:
                return False
            cell = word_table.rows[0].cells[0]
            shd = cell._tc.get_or_add_tcPr().find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
            )
            if shd is not None:
                fill = str(
                    shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", "")
                )
                return fill.upper() == cls.NAVY_HEX.upper()
        except Exception:
            pass
        return False

    @staticmethod
    def _extract_cell_runs(
        cell,
        theme_color_resolver: Optional[ThemeColorResolver] = None,
    ) -> List[TextRun]:
        """Extract runs from all cell paragraphs in reading order."""
        runs: List[TextRun] = []
        for paragraph in cell.paragraphs:
            runs.extend(
                RunExtractor.extract_runs(
                    paragraph,
                    theme_color_resolver=theme_color_resolver,
                )
            )
            if runs and paragraph != cell.paragraphs[-1]:
                runs.append(TextRun(text=" "))
        return runs

    @staticmethod
    def _detect_alignments(word_table, col_count: int) -> List[str]:
        """Infer table alignments from cell paragraph alignment."""
        alignments = ["left"] * max(col_count, 0)
        for col_idx in range(col_count):
            seen: List[str] = []
            for row in word_table.rows:
                if col_idx >= len(row.cells):
                    continue
                for paragraph in row.cells[col_idx].paragraphs:
                    align = paragraph.alignment
                    if align == WD_ALIGN_PARAGRAPH.CENTER:
                        seen.append("center")
                    elif align == WD_ALIGN_PARAGRAPH.RIGHT:
                        seen.append("right")
                    elif align == WD_ALIGN_PARAGRAPH.LEFT or align == WD_ALIGN_PARAGRAPH.JUSTIFY:
                        seen.append("left")
            if seen:
                if "right" in seen:
                    alignments[col_idx] = "right"
                elif "center" in seen:
                    alignments[col_idx] = "center"
        return alignments


class CalloutDetector:
    """Detects callout boxes rendered as single-cell styled tables."""

    CALLOUT_COLORS = {
        "003366": "요약",
        "E6F0FA": "시사점",
        "FFF3CD": "주의",
        "F5F5F5": "참고",
    }
    CALLOUT_PREFIXES = {
        "요약": "요약",
        "시사점": "시사점",
        "주의": "주의",
        "참고": "참고",
        "EXECUTIVE SUMMARY": "요약",
        "KEY INSIGHT": "시사점",
        "WARNING": "주의",
        "NOTE": "참고",
    }

    @classmethod
    def detect_callout(cls, word_table) -> Optional[Blockquote]:
        """Check if a 1x1 table is actually a callout box."""
        if len(word_table.rows) != 1 or len(word_table.columns) != 1:
            return None

        cell = word_table.rows[0].cells[0]
        text = "\n".join(p.text for p in cell.paragraphs).strip()
        if not text:
            return None

        title = cls._detect_callout_by_color(cell)
        if not title:
            title = cls._detect_callout_by_content(text)
        if not title:
            return None

        content = cls._strip_callout_prefix(text, title)
        return Blockquote(title=title, text=content)

    @classmethod
    def _detect_callout_by_color(cls, cell) -> Optional[str]:
        """Detect callout type by cell background color."""
        try:
            shd = cell._tc.get_or_add_tcPr().find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd"
            )
            if shd is not None:
                fill = str(
                    shd.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill", "")
                )
                return cls.CALLOUT_COLORS.get(fill.upper())
        except Exception:
            pass
        return None

    @classmethod
    def _detect_callout_by_content(cls, text: str) -> Optional[str]:
        """Detect callout type by explicit callout prefixes only."""
        text_upper = text.strip().upper()
        for prefix, canonical_title in cls.CALLOUT_PREFIXES.items():
            if text_upper.startswith(prefix):
                return canonical_title
            if text_upper.startswith(f"[{prefix}]"):
                return canonical_title
        return None

    @classmethod
    def _strip_callout_prefix(cls, text: str, title: str) -> str:
        """Strip the visible callout title from the cell text."""
        content = text.strip()
        for prefix, canonical_title in cls.CALLOUT_PREFIXES.items():
            if canonical_title != title:
                continue
            if content.upper().startswith(f"[{prefix}]"):
                content = content[len(prefix) + 2 :].lstrip(":").lstrip()
                break
            if content.upper().startswith(prefix):
                content = content[len(prefix) :].lstrip(":").lstrip()
                break
        return content


# ═══════════════════════════════════════════════════════════════════════════════
# IMAGE EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════


class ImageExtractor:
    """Extracts images from Word documents."""

    def __init__(
        self,
        output_dir: Optional[Path] = None,
        embed_base64: bool = False,
    ):
        self.output_dir = output_dir
        self.embed_base64 = embed_base64
        self._image_counter = 0
        self._extracted_images: Dict[str, ExtractedImageAsset] = {}

    def extract_all(self, doc: DocxDocument) -> Dict[str, ExtractedImageAsset]:
        """Extract all images from document to memory and optional output directory."""
        if self.output_dir:
            self.output_dir.mkdir(parents=True, exist_ok=True)

        for rel in doc.part.rels.values():
            if not self._is_image_relationship(rel):
                continue
            try:
                image_part = rel.target_part
                image_bytes = image_part.blob

                self._image_counter += 1
                ext = Path(rel.target_ref).suffix or ".png"
                filename = f"image_{self._image_counter}{ext}"
                path_name = ""

                if self.output_dir:
                    output_path = self.output_dir / filename
                    output_path.write_bytes(image_bytes)
                    path_name = str(output_path.name)
                    logger.debug("Extracted image: %s", filename)

                mime_type = getattr(image_part, "content_type", "") or self._guess_mime_type(
                    rel.target_ref
                )
                base64_data = (
                    base64.b64encode(image_bytes).decode("ascii")
                    if self.embed_base64
                    else None
                )
                self._extracted_images[rel.rId] = ExtractedImageAsset(
                    path=path_name,
                    base64_data=base64_data,
                    mime_type=mime_type,
                    default_alt_text=self._build_default_alt_text(rel.target_ref),
                )
            except Exception as err:
                logger.warning("Failed to extract image %s: %s", rel.target_ref, err)

        return self._extracted_images

    def get_image_for_rel(self, rel_id: str, alt_text: str = "") -> Optional[Image]:
        """Get Image object for a relationship ID."""
        if rel_id in self._extracted_images:
            asset = self._extracted_images[rel_id]
            return Image(
                alt_text=alt_text or asset.default_alt_text,
                path=asset.path,
                base64_data=asset.base64_data,
                mime_type=asset.mime_type,
            )
        return None

    @staticmethod
    def _is_image_relationship(rel) -> bool:
        """Return True if the relationship points to an image part."""
        reltype = getattr(rel, "reltype", "")
        target_ref = getattr(rel, "target_ref", "")
        return "image" in reltype or "image" in target_ref

    @staticmethod
    def _build_default_alt_text(target_ref: str) -> str:
        """Generate a human-friendly fallback alt text from the file name."""
        stem = Path(target_ref).stem.replace("_", " ").replace("-", " ").strip()
        stem = re.sub(r"(?<=\D)(?=\d)|(?<=\d)(?=\D)", " ", stem)
        return stem.title() if stem else "Image"

    @staticmethod
    def _guess_mime_type(target_ref: str) -> str:
        """Infer MIME type from the image path extension."""
        mime_type, _ = mimetypes.guess_type(target_ref)
        return mime_type or "image/png"


# ═══════════════════════════════════════════════════════════════════════════════
# WORD PARSER
# ═══════════════════════════════════════════════════════════════════════════════


class WordParser:
    """Main parser for Word documents."""

    _IB_TOC_TITLE = "TABLE OF CONTENTS"
    _IB_ENDNOTES_TITLE = "ENDNOTES"
    _IB_DISCLAIMER_TITLE = "면책 조항"
    _IB_TOC_NOTE_SNIPPET = "Update Field"
    _IB_DISCLAIMER_SNIPPET = "당행은 해당 문서에 최대한 정확하고 완전한 정보를 담고자 노력하였으나"
    _IB_METADATA_LABELS = frozenset(
        {"REPORT DATE", "PREPARED BY", "INSTITUTION", "SECTOR", "PREPARED FOR"}
    )
    _IMAGE_FILENAME_RE = re.compile(
        r"^[^\\/\s]+?\.(?:png|jpe?g|gif|bmp|tiff?|svg|webp|emf|wmf)$",
        re.IGNORECASE,
    )

    def __init__(
        self,
        extract_images: bool = True,
        image_output_dir: Optional[Path] = None,
        embed_images_base64: bool = False,
    ):
        self.extract_images = extract_images
        self.image_output_dir = image_output_dir
        self.embed_images_base64 = embed_images_base64
        self.image_extractor: Optional[ImageExtractor] = None
        self.theme_color_resolver = ThemeColorResolver()
        self._omml_available = False
        try:
            from omml_latex import pre_process_docx_math  # noqa: F401

            self._omml_available = True
        except ImportError:
            pass

    def _open_document(self, source: "Union[str, BinaryIO]"):
        """Open a DOCX file or stream, pre-processing OMML equations if possible.

        Args:
            source: File path (str) or binary stream (BinaryIO).
        """
        from stream_utils import ensure_seekable, is_stream

        if self._omml_available:
            try:
                from omml_latex import pre_process_docx_math

                if is_stream(source):
                    source = ensure_seekable(source)
                    processed = pre_process_docx_math(source)
                else:
                    with open(source, "rb") as f:
                        processed = pre_process_docx_math(f)
                logger.debug("OMML pre-processing applied for %s",
                             "<stream>" if is_stream(source) else source)
                return Document(processed)
            except Exception as e:
                logger.debug("OMML pre-processing failed, opening directly: %s", e)
                # Reset stream position if we consumed it
                if is_stream(source):
                    source = ensure_seekable(source)
                    source.seek(0)
        return Document(source)

    def parse(self, source: "Union[str, BinaryIO]") -> DocumentModel:
        """Parse a Word document into DocumentModel.

        Args:
            source: File path (str) or binary stream (BinaryIO).
        """
        source_bytes = self._read_source_bytes(source)
        doc = self._open_document(io.BytesIO(source_bytes))
        self.theme_color_resolver = ThemeColorResolver.from_docx_bytes(source_bytes)
        blocks = list(self._iter_block_items(doc))
        context = ParseContext(
            profile=self._detect_document_profile(blocks),
            footnotes=self._extract_native_footnotes(doc),
        )
        metadata = MetadataExtractor.extract(doc, source_bytes=source_bytes)

        if self.extract_images or self.embed_images_base64:
            self.image_extractor = ImageExtractor(
                output_dir=self.image_output_dir if self.extract_images else None,
                embed_base64=self.embed_images_base64,
            )
            self.image_extractor.extract_all(doc)

        if context.profile == DocumentProfile.IB_GENERATED:
            self._apply_ib_rules(blocks, metadata, context)

        if metadata.title == "IB Report":
            first_heading = self._find_first_heading(blocks, context.skip_indices)
            if first_heading:
                metadata.title = first_heading

        elements = self._parse_elements(blocks, context)

        return DocumentModel(
            metadata=metadata,
            elements=elements,
            footnotes=context.footnotes,
        )

    @staticmethod
    def _extract_native_footnotes(doc: DocxDocument) -> Dict[int, str]:
        """Extract native Word footnotes from /word/footnotes.xml when present."""
        namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

        try:
            footnotes_part = doc.part.part_related_by(RT.FOOTNOTES)
        except KeyError:
            return {}

        try:
            root = ElementTree.fromstring(footnotes_part.blob)
        except ElementTree.ParseError:
            return {}

        footnotes: Dict[int, str] = {}
        for footnote in root.findall("w:footnote", namespace):
            footnote_id = footnote.attrib.get("{%s}id" % namespace["w"], "")
            footnote_type = footnote.attrib.get("{%s}type" % namespace["w"], "")
            if not footnote_id or footnote_type:
                continue
            if not footnote_id.lstrip("-").isdigit() or int(footnote_id) <= 0:
                continue

            texts = [
                text.text.strip()
                for text in footnote.findall(".//w:t", namespace)
                if text.text and text.text.strip()
            ]
            if texts:
                footnotes[int(footnote_id)] = " ".join(texts)
        return footnotes

    @staticmethod
    def _read_source_bytes(source: "Union[str, BinaryIO]") -> bytes:
        """Read the raw DOCX payload for metadata extraction and stream-safe parsing."""
        from stream_utils import ensure_seekable, is_stream

        if is_stream(source):
            stream = ensure_seekable(source)
            stream.seek(0)
            payload = stream.read()
            stream.seek(0)
            return payload

        with open(source, "rb") as file_obj:
            return file_obj.read()

    def _detect_document_profile(self, blocks: List[Union[DocxParagraph, DocxTable]]) -> DocumentProfile:
        """Detect whether the document matches the IB-generated output profile."""
        structural_signals = 0
        textual_signals = 0

        for block in blocks:
            if isinstance(block, DocxParagraph):
                text = block.text.strip()
                if text in {self._IB_TOC_TITLE, self._IB_ENDNOTES_TITLE, self._IB_DISCLAIMER_TITLE}:
                    textual_signals += 1
            elif isinstance(block, DocxTable):
                if self._is_ib_metadata_panel(block):
                    structural_signals += 1
                elif self._is_ib_cover_disclaimer_table(block):
                    structural_signals += 1

        if structural_signals >= 2:
            return DocumentProfile.IB_GENERATED
        if structural_signals >= 1 and textual_signals >= 1:
            return DocumentProfile.IB_GENERATED
        return DocumentProfile.GENERIC

    def _apply_ib_rules(
        self,
        blocks: List[Union[DocxParagraph, DocxTable]],
        metadata: DocumentMetadata,
        context: ParseContext,
    ) -> None:
        """Apply IB-specific metadata recovery and boilerplate skipping."""
        metadata_panel_idx = self._find_metadata_panel_index(blocks)
        cover_disclaimer_idx = self._find_cover_disclaimer_index(blocks)

        if metadata_panel_idx is not None:
            MetadataExtractor.apply_ib_cover_metadata(metadata, blocks, metadata_panel_idx)
            cover_end = metadata_panel_idx
            if cover_disclaimer_idx is not None and cover_disclaimer_idx >= metadata_panel_idx:
                cover_end = cover_disclaimer_idx
            context.skip_indices.update(range(0, cover_end + 1))

        toc_idx = self._find_paragraph_index(blocks, self._IB_TOC_TITLE)
        if toc_idx is not None:
            context.skip_indices.update(self._collect_toc_indices(blocks, toc_idx))

        endnotes_idx = self._find_paragraph_index(blocks, self._IB_ENDNOTES_TITLE)
        if endnotes_idx is not None:
            footnotes, skip_indices = self._extract_endnotes(blocks, endnotes_idx)
            context.footnotes.update(footnotes)
            context.skip_indices.update(skip_indices)

        disclaimer_idx = self._find_paragraph_index(blocks, self._IB_DISCLAIMER_TITLE)
        if disclaimer_idx is not None:
            context.skip_indices.update(range(disclaimer_idx, len(blocks)))

    def _find_first_heading(
        self,
        blocks: List[Union[DocxParagraph, DocxTable]],
        skip_indices: Set[int],
    ) -> Optional[str]:
        """Find the first heading outside skipped boilerplate."""
        for idx, block in enumerate(blocks):
            if idx in skip_indices or not isinstance(block, DocxParagraph):
                continue
            level = StyleDetector.detect_heading_level(block)
            if level == 1:
                return block.text.strip()
        return None

    def _parse_elements(
        self,
        blocks: List[Union[DocxParagraph, DocxTable]],
        context: ParseContext,
    ) -> List[Element]:
        """Parse all content elements after boilerplate analysis."""
        elements: List[Element] = []
        processed_tables: Set[int] = set()
        numbering_tracker = NumberingTracker()

        for idx, block in enumerate(blocks):
            if idx in context.skip_indices:
                numbering_tracker.break_sequence()
                continue

            if isinstance(block, DocxParagraph):
                parsed = self._parse_paragraph(block, numbering_tracker)
                if isinstance(parsed, list):
                    elements.extend(parsed)
                    continue
                element = parsed
            else:
                table_id = id(block._tbl)
                if table_id in processed_tables:
                    continue
                processed_tables.add(table_id)
                element = self._parse_table(block)

            if element:
                elements.append(element)

        return elements

    @staticmethod
    def _iter_block_items(doc: DocxDocument) -> Iterator[Union[DocxParagraph, DocxTable]]:
        """Yield document paragraphs and tables in source order."""
        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                yield DocxParagraph(child, doc)
            elif isinstance(child, CT_Tbl):
                yield DocxTable(child, doc)

    # Regex for LaTeX injected by OMML pre-processor
    _BLOCK_LATEX_RE = re.compile(r"^\$\$(.+?)\$\$$", re.DOTALL)
    _INLINE_LATEX_RE = re.compile(r"\$([^$]+?)\$")

    def _parse_paragraph(
        self,
        para,
        numbering_tracker: Optional[NumberingTracker] = None,
    ) -> Optional[Union[Element, List[Element]]]:
        """Parse a single paragraph."""
        image_refs = self._extract_image_references(para)
        image_element = self._parse_image_paragraph(para, image_refs=image_refs)
        text = para.text.strip()

        if image_element and not text:
            return image_element
        if not text:
            return None

        # Detect LaTeX block equations injected by OMML pre-processor
        block_match = self._BLOCK_LATEX_RE.match(text)
        if block_match:
            return Element(
                element_type=ElementType.LATEX_BLOCK,
                content=LaTeXEquation(expression=block_match.group(1).strip(), is_block=True),
                raw_text=text,
            )

        level = StyleDetector.detect_heading_level(para)
        if level:
            return Element(
                element_type=getattr(ElementType, f"HEADING_{level}", ElementType.HEADING_2),
                content=Heading(level=level, text=text),
                raw_text=text,
            )

        list_type = StyleDetector.detect_list_type(para)
        indent_level = StyleDetector.detect_list_level(para)
        if list_type == "bullet":
            content = StyleDetector.extract_bullet_text(text)
            return Element(
                element_type=ElementType.BULLET_LIST,
                content=ListItem(
                    text=content,
                    runs=RunExtractor.extract_runs(
                        para,
                        theme_color_resolver=self.theme_color_resolver,
                    ),
                    indent_level=indent_level,
                ),
                raw_text=text,
            )
        if list_type == "number":
            visible_match = StyleDetector.NUMBERED_TEXT_PATTERN.match(text)
            if visible_match:
                number = visible_match.group(1)
                content = visible_match.group(2)
            else:
                number = (
                    numbering_tracker.next_number(para, indent_level) if numbering_tracker else "1"
                )
                content = text
            return Element(
                element_type=ElementType.NUMBERED_LIST,
                content=(
                    number,
                    ListItem(
                        text=content,
                        runs=RunExtractor.extract_runs(
                            para,
                            theme_color_resolver=self.theme_color_resolver,
                        ),
                        indent_level=indent_level,
                    ),
                ),
                raw_text=text,
            )

        if image_refs:
            mixed_elements = self._parse_mixed_paragraph_elements(para, image_refs)
            if mixed_elements:
                return mixed_elements

        runs = RunExtractor.extract_runs(para, theme_color_resolver=self.theme_color_resolver)
        has_inline_latex = bool(self._INLINE_LATEX_RE.search(text))
        return Element(
            element_type=ElementType.PARAGRAPH,
            content=Paragraph(text=text, runs=runs, has_inline_latex=has_inline_latex),
            raw_text=text,
        )

    def _parse_image_paragraph(
        self,
        para,
        image_refs: Optional[List[Tuple[str, str]]] = None,
    ) -> Optional[Element]:
        """Convert an image-only paragraph into an Image element."""
        if image_refs is None:
            image_refs = self._extract_image_references(para)
        if not image_refs:
            return None

        rel_id, alt_text = image_refs[0]
        return self._build_image_element(rel_id, alt_text)

    @classmethod
    def _extract_image_references(cls, para) -> List[Tuple[str, str]]:
        """Extract embedded image relationship IDs and alt text from a paragraph."""
        refs: List[Tuple[str, str]] = []
        seen: Set[str] = set()
        for node in para._p.xpath(
            ".//*[local-name()='drawing' or local-name()='pict' or local-name()='object']"
        ):
            for rel_id, alt_text in cls._extract_image_references_from_container(node):
                if rel_id in seen:
                    continue
                seen.add(rel_id)
                refs.append((rel_id, alt_text))

        if refs:
            return refs

        for rel_id, alt_text in cls._extract_image_references_from_container(para._p):
            if rel_id in seen:
                continue
            seen.add(rel_id)
            refs.append((rel_id, alt_text))
        return refs

    @staticmethod
    def _find_rel_id(drawing) -> str:
        """Find the first embedded relationship ID in a drawing node."""
        for blip in drawing.xpath(".//*[local-name()='blip']"):
            rel_id = blip.get(qn("r:embed"))
            if rel_id:
                return str(rel_id)
        return ""

    @staticmethod
    def _extract_image_alt_text(drawing) -> str:
        """Extract human-authored alt text from a Word drawing node."""
        for doc_pr in drawing.xpath(
            ".//*[local-name()='docPr' or local-name()='cNvPr' or local-name()='imagedata']"
        ):
            for attr_name in ("descr", "title"):
                value = (doc_pr.get(attr_name) or "").strip()
                if value and not WordParser._looks_like_generated_image_name(value):
                    return value

            name = (doc_pr.get("name") or "").strip()
            if (
                name
                and not re.match(r"^Picture \d+$", name, re.I)
                and not WordParser._looks_like_generated_image_name(name)
            ):
                return name
        return ""

    @classmethod
    def _looks_like_generated_image_name(cls, value: str) -> bool:
        """Return True for filename-like placeholders Word commonly injects as alt text."""
        candidate = Path(value.strip()).name
        return bool(candidate and cls._IMAGE_FILENAME_RE.match(candidate))

    def _parse_mixed_paragraph_elements(
        self,
        para,
        image_refs: List[Tuple[str, str]],
    ) -> List[Element]:
        """Split mixed text/image paragraphs into ordered paragraph and image elements."""
        elements: List[Element] = []
        pending_runs: List[TextRun] = []
        pending_refs = list(image_refs)

        def flush_pending_text() -> None:
            trimmed_runs = self._trim_text_runs(pending_runs)
            if not trimmed_runs:
                pending_runs.clear()
                return

            text = "".join(run.text for run in trimmed_runs)
            elements.append(
                Element(
                    element_type=ElementType.PARAGRAPH,
                    content=Paragraph(
                        text=text,
                        runs=trimmed_runs,
                        has_inline_latex=bool(self._INLINE_LATEX_RE.search(text)),
                    ),
                    raw_text=text,
                )
            )
            pending_runs.clear()

        for run in para.runs:
            child_tokens = self._extract_run_tokens(run)
            if not child_tokens and run.text:
                child_tokens = [("text", run.text)]

            for token_type, value in child_tokens:
                if token_type == "text":
                    pending_runs.append(
                        TextRun(
                            text=value,
                            bold=bool(run.bold),
                            italic=bool(run.italic),
                            superscript=bool(run.font.superscript),
                            subscript=bool(run.font.subscript),
                            color_hex=RunExtractor._extract_color_hex(
                                run,
                                theme_color_resolver=self.theme_color_resolver,
                            ),
                        )
                    )
                    continue

                if token_type == "footnote":
                    pending_runs.append(TextRun(text=value, superscript=True))
                    continue

                if token_type != "image" or not pending_refs:
                    continue

                flush_pending_text()
                rel_id, alt_text = pending_refs.pop(0)
                image_element = self._build_image_element(rel_id, alt_text)
                if image_element:
                    elements.append(image_element)

        flush_pending_text()
        return elements

    def _build_image_element(self, rel_id: str, alt_text: str) -> Optional[Element]:
        """Build an image element from an extracted Word image relationship."""
        image = None
        if self.image_extractor:
            image = self.image_extractor.get_image_for_rel(rel_id, alt_text=alt_text)
        if image is None:
            image = Image(alt_text=alt_text or "Image", path="")

        return Element(
            element_type=ElementType.IMAGE,
            content=image,
            raw_text=image.path or "[IMAGE]",
        )

    @staticmethod
    def _extract_run_tokens(run) -> List[Tuple[str, str]]:
        """Extract ordered text, footnote, and image markers from a Word run."""
        tokens: List[Tuple[str, str]] = []
        for child in run._r.iterchildren():
            local_name = child.tag.split("}", 1)[-1]
            if local_name == "t" and child.text:
                tokens.append(("text", child.text))
            elif local_name == "tab":
                tokens.append(("text", "\t"))
            elif local_name in {"br", "cr"}:
                tokens.append(("text", "\n"))
            elif local_name == "footnoteReference":
                value = child.get(qn("w:id")) or child.get("id")
                if value and str(value).isdigit():
                    tokens.append(("footnote", str(value)))
            elif local_name in {"drawing", "pict", "object"}:
                tokens.append(("image", ""))
        return tokens

    @staticmethod
    def _trim_text_runs(runs: List[TextRun]) -> List[TextRun]:
        """Trim boundary whitespace while preserving inline formatting."""
        trimmed = [replace(run) for run in runs]
        while trimmed and not trimmed[0].text.strip():
            trimmed.pop(0)
        while trimmed and not trimmed[-1].text.strip():
            trimmed.pop()
        if not trimmed:
            return []

        trimmed[0] = replace(trimmed[0], text=trimmed[0].text.lstrip())
        trimmed[-1] = replace(trimmed[-1], text=trimmed[-1].text.rstrip())
        return [run for run in trimmed if run.text]

    @classmethod
    def _extract_image_references_from_container(cls, container) -> List[Tuple[str, str]]:
        """Extract image relationship IDs and alt text from a Word XML container."""
        refs: List[Tuple[str, str]] = []
        alt_text = cls._extract_image_alt_text(container)

        for blip in container.xpath(".//*[local-name()='blip']"):
            rel_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if rel_id:
                refs.append((str(rel_id), alt_text))

        for image_data in container.xpath(".//*[local-name()='imagedata']"):
            rel_id = (
                image_data.get(qn("r:id"))
                or image_data.get("{urn:schemas-microsoft-com:office:office}relid")
                or image_data.get("id")
                or image_data.get("relid")
            )
            if rel_id:
                refs.append((str(rel_id), alt_text or cls._extract_image_alt_text(image_data)))

        return refs

    def _parse_table(self, word_table) -> Optional[Element]:
        """Parse a Word table."""
        callout = CalloutDetector.detect_callout(word_table)
        if callout:
            return Element(
                element_type=ElementType.BLOCKQUOTE,
                content=callout,
                raw_text=callout.text,
            )

        table = TableExtractor.extract(
            word_table,
            theme_color_resolver=self.theme_color_resolver,
        )
        return Element(
            element_type=ElementType.TABLE,
            content=table,
            raw_text="[TABLE]",
        )

    def _find_metadata_panel_index(
        self, blocks: List[Union[DocxParagraph, DocxTable]]
    ) -> Optional[int]:
        """Find the IB metadata panel table index."""
        for idx, block in enumerate(blocks):
            if isinstance(block, DocxTable) and self._is_ib_metadata_panel(block):
                return idx
        return None

    def _find_cover_disclaimer_index(
        self, blocks: List[Union[DocxParagraph, DocxTable]]
    ) -> Optional[int]:
        """Find the cover disclaimer table index."""
        for idx, block in enumerate(blocks):
            if isinstance(block, DocxTable) and self._is_ib_cover_disclaimer_table(block):
                return idx
        return None

    def _find_paragraph_index(
        self,
        blocks: List[Union[DocxParagraph, DocxTable]],
        target_text: str,
    ) -> Optional[int]:
        """Find the first paragraph whose stripped text matches target_text."""
        for idx, block in enumerate(blocks):
            if isinstance(block, DocxParagraph) and block.text.strip() == target_text:
                return idx
        return None

    def _collect_toc_indices(
        self,
        blocks: List[Union[DocxParagraph, DocxTable]],
        toc_index: int,
    ) -> Set[int]:
        """Collect the generated TOC heading/field/note blocks."""
        indices = {toc_index}
        saw_field_paragraph = False

        for idx in range(toc_index + 1, len(blocks)):
            block = blocks[idx]
            if not isinstance(block, DocxParagraph):
                break

            text = block.text.strip()
            style_name = block.style.name if block.style else ""
            is_toc_style = style_name.upper().startswith("TOC")

            if not saw_field_paragraph:
                indices.add(idx)
                saw_field_paragraph = True
                continue

            if not text and StyleDetector.has_embedded_content(block):
                break

            if not text or is_toc_style or self._IB_TOC_NOTE_SNIPPET in text:
                indices.add(idx)
                continue

            break

        return indices

    def _extract_endnotes(
        self,
        blocks: List[Union[DocxParagraph, DocxTable]],
        start_index: int,
    ) -> Tuple[Dict[int, str], Set[int]]:
        """Extract ENDNOTES paragraphs and the indices they occupy."""
        footnotes: Dict[int, str] = {}
        indices: Set[int] = {start_index}

        for idx in range(start_index + 1, len(blocks)):
            block = blocks[idx]
            if not isinstance(block, DocxParagraph):
                indices.add(idx)
                continue

            text = block.text.strip()
            if text == self._IB_DISCLAIMER_TITLE:
                break

            indices.add(idx)
            if not text:
                continue

            endnote = self._parse_endnote_paragraph(block)
            if endnote:
                number, note_text = endnote
                footnotes[number] = note_text

        return footnotes, indices

    @staticmethod
    def _parse_endnote_paragraph(para) -> Optional[Tuple[int, str]]:
        """Parse a generated ENDNOTES paragraph into (number, text)."""
        runs = [run for run in para.runs if run.text]
        if not runs:
            return None

        first_content_run = next((run for run in runs if run.text.strip()), None)
        if first_content_run is None or not first_content_run.font.superscript:
            return None

        number_text = first_content_run.text.strip()
        if not number_text.isdigit():
            return None

        remaining_text = "".join(run.text for run in runs[1:]).strip()
        return int(number_text), remaining_text

    def _is_ib_metadata_panel(self, table: DocxTable) -> bool:
        """Return True when a table matches the generated cover metadata panel."""
        if len(table.columns) != 2 or len(table.rows) < 3:
            return False

        labels = {row.cells[0].text.strip().upper() for row in table.rows if len(row.cells) >= 2}
        return len(labels & self._IB_METADATA_LABELS) >= 3

    def _is_ib_cover_disclaimer_table(self, table: DocxTable) -> bool:
        """Return True when a table matches the generated cover disclaimer block."""
        if len(table.rows) != 1 or len(table.columns) != 1:
            return False

        text = table.rows[0].cells[0].text.strip()
        return self._IB_DISCLAIMER_SNIPPET in text


# ═══════════════════════════════════════════════════════════════════════════════
# RUN EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════


class RunExtractor:
    """Extracts TextRun objects from Word paragraph runs."""

    @staticmethod
    def extract_runs(
        para,
        theme_color_resolver: Optional[ThemeColorResolver] = None,
    ) -> List[TextRun]:
        """Extract text runs with bold, italic, superscript, and subscript flags."""
        runs: List[TextRun] = []
        for run in para.runs:
            if not run.text:
                refs = RunExtractor._extract_footnote_references(run)
                for ref in refs:
                    runs.append(TextRun(text=str(ref), superscript=True))
                continue

            runs.append(
                TextRun(
                    text=run.text,
                    bold=bool(run.bold),
                    italic=bool(run.italic),
                    superscript=bool(run.font.superscript),
                    subscript=bool(run.font.subscript),
                    color_hex=RunExtractor._extract_color_hex(
                        run,
                        theme_color_resolver=theme_color_resolver,
                    ),
                )
            )
        return runs

    @staticmethod
    def _extract_color_hex(
        run,
        theme_color_resolver: Optional[ThemeColorResolver] = None,
    ) -> Optional[str]:
        """Extract explicit RGB color from a Word run when available."""
        color_nodes = run._r.xpath(".//*[local-name()='color']")
        if color_nodes:
            color_node = color_nodes[0]
            theme_color = color_node.get(qn("w:themeColor")) or color_node.get("themeColor")
            theme_tint = color_node.get(qn("w:themeTint")) or color_node.get("themeTint")
            theme_shade = color_node.get(qn("w:themeShade")) or color_node.get("themeShade")
            if theme_color and theme_color_resolver:
                resolved = theme_color_resolver.resolve(
                    theme_color,
                    tint=theme_tint,
                    shade=theme_shade,
                )
                if resolved:
                    return resolved

        color = getattr(run.font, "color", None)
        rgb = getattr(color, "rgb", None)
        if rgb is None:
            return None

        value = str(rgb).strip()
        if not value:
            return None
        if not value.startswith("#"):
            value = f"#{value}"
        return value.upper()

    @staticmethod
    def _extract_footnote_references(run) -> List[int]:
        """Extract native footnoteReference ids from a run XML element."""
        refs: List[int] = []
        for ref in run._r.xpath(".//*[local-name()='footnoteReference']"):
            value = ref.get(qn("w:id")) or ref.get("id")
            if value and str(value).isdigit():
                refs.append(int(value))
        return refs


# ═══════════════════════════════════════════════════════════════════════════════
# PUBLIC API
# ═══════════════════════════════════════════════════════════════════════════════


def parse_word_file(
    source: "Union[str, BinaryIO]",
    extract_images: bool = True,
    image_output_dir: Optional[str] = None,
    embed_images_base64: bool = False,
) -> DocumentModel:
    """
    Parse a Word document into a DocumentModel.

    Args:
        source: Path to the .docx file (str) or a binary stream (BinaryIO).
        extract_images: Whether to extract embedded images
        image_output_dir: Directory to save extracted images
        embed_images_base64: Whether to inline embedded images as Base64 data

    Returns:
        DocumentModel containing parsed elements
    """
    parser = WordParser(
        extract_images=extract_images,
        image_output_dir=Path(image_output_dir) if image_output_dir else None,
        embed_images_base64=embed_images_base64,
    )
    return parser.parse(source)
