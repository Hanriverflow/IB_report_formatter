"""
IB Renderer Module for Word Report Generation
Handles styling and rendering of document elements in IB Bank style.

Changelog (v2):
    - Fixed header row styling (p.clear() + add_run pattern)
    - Compatible with md_parser v2 heading levels (1-4)
    - NUMBERED_HEADING rendered at correct level
    - Compiled regex for TextRenderer performance
    - Style name constants in IBStyle
    - Element-level error resilience with try-except
    - Table column count defensive checks
    - Blockquote multi-line content preserved
"""

import logging
import os
import platform
import re
import time
from dataclasses import dataclass
from typing import Dict, List, Optional, Tuple, cast

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import CONTENT_TYPE as CT
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.opc.packuri import PackURI
from docx.opc.part import XmlPart
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.parser import parse_xml
from docx.shared import Inches, Pt, RGBColor

from md_parser import (
    Blockquote,
    CodeBlock,
    DocumentModel,
    Element,
    ElementType,
    Heading,
    Image,
    LaTeXEquation,
    ListItem,
    Paragraph,
    Table,
    TableCell,
    TableRow,
    TableType,
    TextParser,
    TextRun,
)

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# STYLE CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════


@dataclass(frozen=True)
class IBStyle:
    """IB Bank styling constants"""

    # ── Colors (RGBColor) ───────────────────────────────────────────────────
    NAVY: RGBColor = RGBColor(0, 51, 102)
    DARK_GRAY: RGBColor = RGBColor(64, 64, 64)
    LIGHT_GRAY: RGBColor = RGBColor(245, 245, 245)
    ACCENT_BLUE: RGBColor = RGBColor(230, 240, 250)
    WHITE: RGBColor = RGBColor(255, 255, 255)
    RED: RGBColor = RGBColor(192, 0, 0)
    GREEN: RGBColor = RGBColor(0, 128, 0)
    ORANGE: RGBColor = RGBColor(255, 165, 0)
    MEDIUM_GRAY: RGBColor = RGBColor(128, 128, 128)
    CODE_BG: RGBColor = RGBColor(248, 249, 250)

    # ── Colors (Hex for OOXML) ──────────────────────────────────────────────
    NAVY_HEX: str = "003366"
    LIGHT_GRAY_HEX: str = "F5F5F5"
    ACCENT_BLUE_HEX: str = "E6F0FA"
    GRAY_BORDER_HEX: str = "C8C8C8"
    YELLOW_HEX: str = "FFFF00"

    # ── Fonts ───────────────────────────────────────────────────────────────
    HEADING_FONT: str = "Arial"
    BODY_FONT: str = "Calibri"
    KOREAN_FONT: str = "Malgun Gothic"
    COVER_FONT: str = "Malgun Gothic"
    TOC_FONT: str = "Malgun Gothic"

    # ── Sizes ───────────────────────────────────────────────────────────────
    H1_SIZE: Pt = Pt(14)
    H2_SIZE: Pt = Pt(12)
    H3_SIZE: Pt = Pt(11)
    H4_SIZE: Pt = Pt(10.5)
    BODY_SIZE: Pt = Pt(10.5)
    SMALL_SIZE: Pt = Pt(9)
    TABLE_HEADER_SIZE: Pt = Pt(10)
    TABLE_BODY_SIZE: Pt = Pt(10)

    # ── Spacing ─────────────────────────────────────────────────────────────
    H1_SPACE_BEFORE: Pt = Pt(18)
    H1_SPACE_AFTER: Pt = Pt(6)
    H2_SPACE_BEFORE: Pt = Pt(12)
    H2_SPACE_AFTER: Pt = Pt(4)
    H3_SPACE_BEFORE: Pt = Pt(10)
    H3_SPACE_AFTER: Pt = Pt(2)
    BODY_SPACE_AFTER: Pt = Pt(8)
    BULLET_SPACE_AFTER: Pt = Pt(4)

    # ── Line spacing ────────────────────────────────────────────────────────
    BODY_LINE_SPACING: float = 1.15

    # ── Margins ─────────────────────────────────────────────────────────────
    TOP_MARGIN: Inches = Inches(1.0)
    BOTTOM_MARGIN: Inches = Inches(0.75)
    LEFT_MARGIN: Inches = Inches(1.0)
    RIGHT_MARGIN: Inches = Inches(0.8)

    # ── Bullet ──────────────────────────────────────────────────────────────
    BULLET_INDENT: Inches = Inches(0.25)
    DEEP_LIST_INDENT: Inches = Inches(0.125)
    FULL_LIST_INDENT_LEVELS: int = 4
    MAX_LIST_INDENT: Inches = Inches(1.5)
    BULLET_CHAR: str = "■"

    # ── Custom Style Names ──────────────────────────────────────────────────
    STYLE_IB_BODY: str = "IB Body"
    STYLE_IB_BULLET: str = "IB Bullet"
    STYLE_TABLE_GRID: str = "Table Grid"


# Singleton style instance
STYLE = IBStyle()


# ═══════════════════════════════════════════════════════════════════════════════
# FONT STYLER
# ═══════════════════════════════════════════════════════════════════════════════


class FontStyler:
    """Handles font styling including East Asian fonts"""

    @staticmethod
    def set_east_asian_font(element, font_name: Optional[str] = None):
        """Set East Asian font (for Korean text) on a style or run element"""
        resolved_font = font_name or FontPolicy.resolve_korean_font()
        elm = element._element
        rPr = elm.get_or_add_rPr()
        if rPr.rFonts is None:
            rPr.get_or_add_rFonts()
        rPr.rFonts.set(qn("w:eastAsia"), resolved_font)

    @staticmethod
    def apply_run_style(
        run,
        font_name: Optional[str] = None,
        font_size: Optional[Pt] = None,
        bold: Optional[bool] = None,
        italic: Optional[bool] = None,
        color: Optional[RGBColor] = None,
        superscript: Optional[bool] = None,
    ):
        """Apply styling to a run"""
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = font_size
        if bold is not None:
            run.font.bold = bold
        if italic is not None:
            run.font.italic = italic
        if color:
            run.font.color.rgb = color
        if superscript is not None:
            run.font.superscript = superscript
        FontStyler.set_east_asian_font(run, font_name=font_name)


class FontPolicy:
    """Resolve platform-aware East Asian font defaults."""

    _resolved_fonts: Dict[str, str] = {}

    @classmethod
    def resolve_korean_font(cls, system_name: Optional[str] = None) -> str:
        """Return the preferred Korean font for the current platform."""
        system = system_name or platform.system() or "Unknown"
        cached = cls._resolved_fonts.get(system)
        if cached:
            return cached

        if system == "Darwin":
            candidates = ("Apple SD Gothic Neo", STYLE.KOREAN_FONT, "NanumGothic")
        elif system == "Windows":
            candidates = (STYLE.KOREAN_FONT, "NanumGothic", "Apple SD Gothic Neo")
        else:
            candidates = (STYLE.KOREAN_FONT, "NanumGothic", "Apple SD Gothic Neo")

        chosen = candidates[0]
        logger.debug(
            "Resolved Korean font for %s: %s (fallbacks: %s)",
            system,
            chosen,
            ", ".join(candidates[1:]) or "none",
        )
        cls._resolved_fonts[system] = chosen
        return chosen


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT STYLER
# ═══════════════════════════════════════════════════════════════════════════════


class DocumentStyler:
    """Sets up document styles"""

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def setup_document(self):
        """Set up document margins and page settings"""
        for section in self.doc.sections:
            section.top_margin = STYLE.TOP_MARGIN
            section.bottom_margin = STYLE.BOTTOM_MARGIN
            section.left_margin = STYLE.LEFT_MARGIN
            section.right_margin = STYLE.RIGHT_MARGIN

    def create_styles(self):
        """Create all custom IB styles"""
        styles = self.doc.styles

        # Heading 1
        self._setup_heading_style(
            styles["Heading 1"],
            font_size=STYLE.H1_SIZE,
            color=STYLE.NAVY,
            space_before=STYLE.H1_SPACE_BEFORE,
            space_after=STYLE.H1_SPACE_AFTER,
            add_border=True,
        )

        # Heading 2
        self._setup_heading_style(
            styles["Heading 2"],
            font_size=STYLE.H2_SIZE,
            color=STYLE.DARK_GRAY,
            space_before=STYLE.H2_SPACE_BEFORE,
            space_after=STYLE.H2_SPACE_AFTER,
        )

        # Heading 3
        self._setup_heading_style(
            styles["Heading 3"],
            font_size=STYLE.H3_SIZE,
            color=STYLE.NAVY,
            space_before=STYLE.H3_SPACE_BEFORE,
            space_after=STYLE.H3_SPACE_AFTER,
        )

        # Heading 4
        self._setup_heading_style(
            styles["Heading 4"],
            font_size=STYLE.H4_SIZE,
            color=STYLE.DARK_GRAY,
            space_before=STYLE.H3_SPACE_BEFORE,
            space_after=STYLE.H3_SPACE_AFTER,
        )

        # IB Body
        body = self._get_or_create_style(STYLE.STYLE_IB_BODY, WD_STYLE_TYPE.PARAGRAPH)
        body.font.name = STYLE.BODY_FONT
        body.font.size = STYLE.BODY_SIZE
        body.paragraph_format.line_spacing = STYLE.BODY_LINE_SPACING
        body.paragraph_format.space_after = STYLE.BODY_SPACE_AFTER
        body.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        FontStyler.set_east_asian_font(body)

        # IB Bullet
        bullet = self._get_or_create_style(STYLE.STYLE_IB_BULLET, WD_STYLE_TYPE.PARAGRAPH)
        bullet.font.name = STYLE.BODY_FONT
        bullet.font.size = STYLE.BODY_SIZE
        bullet.paragraph_format.left_indent = STYLE.BULLET_INDENT
        bullet.paragraph_format.first_line_indent = -STYLE.BULLET_INDENT
        bullet.paragraph_format.space_after = STYLE.BULLET_SPACE_AFTER
        FontStyler.set_east_asian_font(bullet)

        # Built-in TOC styles used by Word field updates
        self._setup_toc_style(
            self._get_or_create_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH),
            STYLE.BODY_SIZE,
            bold=True,
        )
        self._setup_toc_style(
            self._get_or_create_style("TOC 2", WD_STYLE_TYPE.PARAGRAPH),
            Pt(10),
            left_indent=Inches(0.2),
        )
        self._setup_toc_style(
            self._get_or_create_style("TOC 3", WD_STYLE_TYPE.PARAGRAPH),
            STYLE.SMALL_SIZE,
            left_indent=Inches(0.4),
        )
        self._setup_toc_style(
            self._get_or_create_style("TOC 4", WD_STYLE_TYPE.PARAGRAPH),
            STYLE.SMALL_SIZE,
            left_indent=Inches(0.6),
        )

    def _setup_heading_style(
        self,
        style,
        font_size: Pt,
        color: RGBColor,
        space_before: Pt,
        space_after: Pt,
        add_border: bool = False,
    ):
        """Configure a heading style"""
        style.font.name = STYLE.HEADING_FONT
        style.font.size = font_size
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = space_before
        style.paragraph_format.space_after = space_after
        FontStyler.set_east_asian_font(style)
        if add_border:
            self._add_bottom_border(style)

    def _get_or_create_style(self, name: str, style_type):
        """Get existing style or create new one"""
        try:
            return self.doc.styles.add_style(name, style_type)
        except ValueError:
            return self.doc.styles[name]

    def _setup_toc_style(
        self,
        style,
        font_size: Pt,
        bold: bool = False,
        left_indent: Inches = Inches(0),
    ) -> None:
        """Configure built-in TOC styles so Word-generated entries match Korean cover typography."""
        style.font.name = STYLE.TOC_FONT
        style.font.size = font_size
        style.font.bold = bold
        style.font.color.rgb = STYLE.DARK_GRAY
        style.paragraph_format.left_indent = left_indent
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(2 if bold else 0)
        FontStyler.set_east_asian_font(style, STYLE.TOC_FONT)

    @staticmethod
    def _add_bottom_border(style):
        """Add bottom border to a style"""
        pPr = style._element.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:color"), STYLE.NAVY_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def setup_header_footer(
        self,
        company: str = "",
        confidential: bool = True,
        show_page_numbers: bool = True,
    ):
        """
        Set up professional IB-style header and footer.

        Args:
            company: Company name to display in header
            confidential: Whether to show "CONFIDENTIAL" mark
            show_page_numbers: Whether to show page numbers in footer
        """
        for section in self.doc.sections:
            # ── Header ─────────────────────────────────────────────────────────
            header = section.header
            header.is_linked_to_previous = False

            # Clear existing content
            for para in header.paragraphs:
                para.clear()

            # Add header content
            header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

            # Left-aligned company name
            if company:
                company_run = header_para.add_run(company)
                FontStyler.apply_run_style(
                    company_run,
                    font_name=STYLE.HEADING_FONT,
                    font_size=STYLE.SMALL_SIZE,
                    color=STYLE.DARK_GRAY,
                )

            # Tab for right alignment
            header_para.add_run("\t\t")

            # Right-aligned confidential mark
            if confidential:
                conf_run = header_para.add_run("CONFIDENTIAL")
                FontStyler.apply_run_style(
                    conf_run,
                    font_name=STYLE.HEADING_FONT,
                    font_size=STYLE.SMALL_SIZE,
                    bold=True,
                    color=STYLE.RED,
                )

            # Add separator line under header
            self._add_header_border(header_para)

            # ── Footer ─────────────────────────────────────────────────────────
            footer = section.footer
            footer.is_linked_to_previous = False

            # Clear existing content
            for para in footer.paragraphs:
                para.clear()

            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if show_page_numbers:
                # Add page number field
                self._add_page_number_field(footer_para)

    def _add_header_border(self, paragraph):
        """Add bottom border to header paragraph."""
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), STYLE.GRAY_BORDER_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _add_page_number_field(self, paragraph):
        """Add page number field code to paragraph."""
        # "Page X of Y" format
        run1 = paragraph.add_run("Page ")
        FontStyler.apply_run_style(
            run1,
            font_name=STYLE.BODY_FONT,
            font_size=STYLE.SMALL_SIZE,
            color=STYLE.DARK_GRAY,
        )

        # PAGE field
        run_page = paragraph.add_run()
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")

        run_page._r.append(fldChar1)
        run_page._r.append(instrText)
        run_page._r.append(fldChar2)
        FontStyler.apply_run_style(
            run_page,
            font_name=STYLE.BODY_FONT,
            font_size=STYLE.SMALL_SIZE,
            color=STYLE.DARK_GRAY,
        )

        run2 = paragraph.add_run(" of ")
        FontStyler.apply_run_style(
            run2,
            font_name=STYLE.BODY_FONT,
            font_size=STYLE.SMALL_SIZE,
            color=STYLE.DARK_GRAY,
        )

        # NUMPAGES field
        run_total = paragraph.add_run()
        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "begin")

        instrText2 = OxmlElement("w:instrText")
        instrText2.text = "NUMPAGES"

        fldChar4 = OxmlElement("w:fldChar")
        fldChar4.set(qn("w:fldCharType"), "end")

        run_total._r.append(fldChar3)
        run_total._r.append(instrText2)
        run_total._r.append(fldChar4)
        FontStyler.apply_run_style(
            run_total,
            font_name=STYLE.BODY_FONT,
            font_size=STYLE.SMALL_SIZE,
            color=STYLE.DARK_GRAY,
        )


# ═══════════════════════════════════════════════════════════════════════════════
# TABLE STYLER
# ═══════════════════════════════════════════════════════════════════════════════


class TableStyler:
    """Handles table styling"""

    @staticmethod
    def set_cell_background(cell, hex_color: str):
        """Set cell background color"""
        tcPr = cell._element.tcPr
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            cell._element.append(tcPr)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    @staticmethod
    def set_table_borders(table):
        """Apply IB-style borders to table"""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")

        # Outer borders: Navy, thick
        for border_name in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "12")
            border.set(qn("w:color"), STYLE.NAVY_HEX)
            tblBorders.append(border)

        # Inner horizontal: Gray, thin solid
        insideH = OxmlElement("w:insideH")
        insideH.set(qn("w:val"), "single")
        insideH.set(qn("w:sz"), "4")
        insideH.set(qn("w:color"), STYLE.GRAY_BORDER_HEX)
        tblBorders.append(insideH)

        # Inner vertical: Gray, dotted for readability
        insideV = OxmlElement("w:insideV")
        insideV.set(qn("w:val"), "dotted")
        insideV.set(qn("w:sz"), "4")
        insideV.set(qn("w:color"), STYLE.GRAY_BORDER_HEX)
        tblBorders.append(insideV)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)


# ═══════════════════════════════════════════════════════════════════════════════
# ELEMENT RENDERERS
# ═══════════════════════════════════════════════════════════════════════════════


class TextRenderer:
    """Renders text with formatting"""

    # Compiled regex — class-level cache
    _BOLD_SPLIT_RE = re.compile(r"(\*\*.*?\*\*)")
    _ITALIC_SPLIT_RE = re.compile(r"(?<!\*)(\*[^*]+?\*)(?!\*)")
    _SUPERSCRIPT_RE = re.compile(r"\^([^^]+?)\^")
    _SUBSCRIPT_PATTERN = r"(?<!~)~[A-Za-z0-9]{1,8}~(?!~)"
    _VERTICAL_ALIGN_SPLIT_RE = re.compile(r"(\^[^^]+?\^|" + _SUBSCRIPT_PATTERN + r")")
    _ESCAPE_RE = re.compile(r'\\([~.*"\'()\[\]{}|_-])')

    @staticmethod
    def render_runs(
        paragraph,
        runs: List[TextRun],
        default_color: Optional[RGBColor] = None,
        font_name: Optional[str] = None,
        font_size: Optional[Pt] = None,
    ):
        """Render text runs to a paragraph"""
        font_name = font_name or STYLE.BODY_FONT
        font_size = font_size or STYLE.BODY_SIZE

        for run_data in runs:
            if run_data.is_latex:
                TextRenderer._render_inline_latex(
                    paragraph,
                    run_data.text,
                    font_size=font_size,
                )
                continue

            run = paragraph.add_run(run_data.text)
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = run_data.bold
            run.font.italic = run_data.italic
            if run_data.superscript:
                run.font.superscript = True
            if run_data.subscript:
                run.font.subscript = True
            run_color = TextRenderer._rgb_from_hex(run_data.color_hex) or default_color
            if run_color:
                run.font.color.rgb = run_color
            FontStyler.set_east_asian_font(run)

    @staticmethod
    def _render_inline_latex(paragraph, expression: str, font_size: Pt):
        """Render inline LaTeX as an inline image when possible."""
        image_path = LaTeXRenderer.render_to_image(
            expression,
            fontsize=max(int(round(font_size.pt)), 12),
            dpi=200,
        )

        if image_path:
            try:
                run = paragraph.add_run()
                run.add_picture(
                    image_path,
                    height=Pt(max(font_size.pt * 1.4, 14)),
                )
                return
            except Exception as err:
                logger.warning("Inline LaTeX image insertion failed: %s", err)
            finally:
                try:
                    os.unlink(image_path)
                except OSError:
                    pass

        fallback_run = paragraph.add_run(f"[{expression}]")
        FontStyler.apply_run_style(
            fallback_run,
            font_name="Consolas",
            font_size=font_size,
            italic=True,
            color=STYLE.DARK_GRAY,
        )

    @staticmethod
    def render_text_with_bold(
        paragraph,
        text: str,
        font_name: Optional[str] = None,
        font_size: Optional[Pt] = None,
    ):
        """Parse and render text with **bold** markers (legacy, delegates to render_text_with_formatting)."""
        TextRenderer.render_text_with_formatting(
            paragraph, text, font_name=font_name, font_size=font_size
        )

    @staticmethod
    def render_text_with_formatting(
        paragraph,
        text: str,
        font_name: Optional[str] = None,
        font_size: Optional[Pt] = None,
        default_color: Optional[RGBColor] = None,
    ):
        """Parse and render text with **bold**, *italic*, ^superscript^, and ~subscript~ markers.

        Handles nested markers in a multi-pass approach:
            1. Split on **bold** markers
            2. Within non-bold segments, split on *italic* markers
            3. Within all segments, detect ^superscript^ and ~subscript~ markers
        """
        font_name = font_name or STYLE.BODY_FONT
        font_size = font_size or STYLE.BODY_SIZE

        parsed_runs = TextParser.parse_runs_plain(text)
        if any(
            run.bold or run.italic or run.superscript or run.subscript or run.color_hex
            for run in parsed_runs
        ):
            TextRenderer.render_runs(
                paragraph,
                parsed_runs,
                default_color=default_color,
                font_name=font_name,
                font_size=font_size,
            )
            return

        # Split on bold markers first
        bold_parts = TextRenderer._BOLD_SPLIT_RE.split(text)
        for bold_part in bold_parts:
            if not bold_part:
                continue

            if bold_part.startswith("**") and bold_part.endswith("**") and len(bold_part) > 4:
                # Bold segment — check for ^superscript^ / ~subscript~ inside
                inner = TextRenderer._cleanup(bold_part[2:-2])
                TextRenderer._render_with_vertical_align(
                    paragraph, inner, font_name, font_size, bold=True, italic=False,
                    color=default_color,
                )
            else:
                # Non-bold segment — check for *italic* markers
                italic_parts = TextRenderer._ITALIC_SPLIT_RE.split(bold_part)
                for italic_part in italic_parts:
                    if not italic_part:
                        continue

                    if (
                        italic_part.startswith("*")
                        and italic_part.endswith("*")
                        and len(italic_part) > 2
                        and not italic_part.startswith("**")
                    ):
                        inner = TextRenderer._cleanup(italic_part[1:-1])
                        TextRenderer._render_with_vertical_align(
                            paragraph, inner, font_name, font_size,
                            bold=False, italic=True, color=default_color,
                        )
                    else:
                        cleaned = TextRenderer._cleanup(italic_part)
                        if cleaned:
                            TextRenderer._render_with_vertical_align(
                                paragraph, cleaned, font_name, font_size,
                                bold=False, italic=False, color=default_color,
                            )

    @staticmethod
    def _render_with_vertical_align(
        paragraph,
        text: str,
        font_name: str,
        font_size: Pt,
        bold: bool,
        italic: bool,
        color: Optional[RGBColor] = None,
    ):
        """Render text, detecting ^superscript^ and ~subscript~ markers."""
        if not text:
            return

        parts = TextRenderer._VERTICAL_ALIGN_SPLIT_RE.split(text)
        for part in parts:
            if not part:
                continue

            is_superscript = False
            is_subscript = False
            content = part

            if part.startswith("^") and part.endswith("^") and len(part) > 2:
                content = part[1:-1]
                is_superscript = True
            elif part.startswith("~") and part.endswith("~") and len(part) > 2:
                content = part[1:-1]
                is_subscript = True

            run = paragraph.add_run(content)
            run.font.name = font_name
            run.font.size = font_size
            run.font.bold = bold
            run.font.italic = italic
            if is_superscript:
                run.font.superscript = True
            if is_subscript:
                run.font.subscript = True
            if color:
                run.font.color.rgb = color
            FontStyler.set_east_asian_font(run)

    @staticmethod
    def _rgb_from_hex(color_hex: Optional[str]) -> Optional[RGBColor]:
        """Convert #RRGGBB strings to python-docx RGBColor."""
        if not color_hex:
            return None

        normalized = color_hex.lstrip("#")
        if len(normalized) != 6 or not re.fullmatch(r"[0-9A-Fa-f]{6}", normalized):
            return None
        return RGBColor.from_string(normalized.upper())

    @staticmethod
    def _cleanup(text: str) -> str:
        """Clean up escape characters (single-pass compiled regex)"""
        return TextRenderer._ESCAPE_RE.sub(r"\1", text).strip()


class CoverRenderer:
    """Renders the cover page"""

    _COVER_DISCLAIMER_TEXT = (
        "당행은 해당 문서에 최대한 정확하고 완전한 정보를 담고자 노력하였으나, 오류와 중요정보의 "
        "누락이 있을 수 있으며, 정보의 정확성, 완전성 및 적정성을 보장하지 않습니다. 이 문서는 "
        "고객의 이해를 돕기 위하여 작성된 설명자료에 불과하므로, 고객은 각자의 책임으로 개별 계약서나 "
        "공시된 정보를 통하여 거래의 내용을 숙지하여야 합니다. 이 문서는 확정적인 거래조건을 "
        "구성하지 않으며 법적인 책임을 위한 근거자료로 사용될 수 없습니다. 본 자료는 당행의 "
        "저작물로서 모든 저작권은 당행에게 있으며, 당행의 동의 없이 어떠한 경우에도 어떠한 형태로든 "
        "복제, 배포, 전송, 변경, 대여할 수 없으며, 당행의 요청 시에 즉시 반환, 파기하여 주시기 바랍니다."
    )

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, metadata):
        """Render upgraded professional IB cover page."""
        title = (metadata.title or "IB Report").strip()
        subtitle = metadata.subtitle.strip()
        ticker = metadata.ticker.strip()
        sector = metadata.sector.strip()
        analyst = metadata.analyst.strip()
        company = metadata.company.strip()
        subject_company = metadata.extra.get("subject_company", "").strip()
        date_text = metadata.extra.get("date", "").strip()
        recipient = metadata.extra.get("recipient", "").strip()

        report_type = metadata.extra.get("report_type", "").strip().upper()
        institution = subject_company or company
        identity = ticker if ticker else institution
        cover_title = title
        cover_identity = identity

        if (
            not ticker
            and self._is_meaningful_metadata_value(institution, style_default="Korea Development Bank")
            and title.startswith(institution)
        ):
            stripped_title = title[len(institution):].strip(" :-")
            if stripped_title:
                cover_identity = institution
                cover_title = stripped_title.strip(" —–-:")

        self._add_spacer(2)

        # Report classification block
        if report_type:
            type_para = self.doc.add_paragraph()
            type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            type_run = type_para.add_run(report_type)
            FontStyler.apply_run_style(
                type_run,
                font_name=STYLE.COVER_FONT,
                font_size=Pt(11),
                bold=True,
                color=STYLE.DARK_GRAY,
            )

        if self._is_meaningful_metadata_value(sector, style_default="SECTOR"):
            sector_para = self.doc.add_paragraph()
            sector_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sector_run = sector_para.add_run(sector.upper())
            FontStyler.apply_run_style(
                sector_run,
                font_name=STYLE.COVER_FONT,
                font_size=Pt(10),
                color=STYLE.MEDIUM_GRAY,
            )

        self._add_spacer(1)
        self._add_horizontal_rule(STYLE.NAVY_HEX, "14")
        self._add_spacer(2)

        # Security / company identifier
        if self._is_meaningful_metadata_value(cover_identity, style_default="Korea Development Bank"):
            id_para = self.doc.add_paragraph()
            id_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            id_run = id_para.add_run(cover_identity)
            FontStyler.apply_run_style(
                id_run,
                font_name=STYLE.COVER_FONT,
                font_size=Pt(15),
                bold=True,
                color=STYLE.NAVY,
            )

        # Main title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(cover_title)
        FontStyler.apply_run_style(
            title_run,
            font_name=STYLE.COVER_FONT,
            font_size=Pt(24),
            bold=True,
            color=STYLE.NAVY,
        )

        # Subtitle
        if subtitle:
            subtitle_para = self.doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle_para.add_run(subtitle)
            FontStyler.apply_run_style(
                subtitle_run,
                font_name=STYLE.COVER_FONT,
                font_size=Pt(12),
                italic=True,
                color=STYLE.DARK_GRAY,
            )

        self._add_spacer(2)
        self._add_horizontal_rule(STYLE.GRAY_BORDER_HEX, "6")
        self._add_spacer(2)

        # Metadata panel
        self._render_metadata_panel(
            report_date=date_text or time.strftime("%B %d, %Y"),
            analyst=analyst,
            company=institution,
            sector=sector,
            recipient=recipient,
            analysis_period=metadata.extra.get("analysis_period", "").strip(),
            analysis_basis=metadata.extra.get("analysis_basis", "").strip(),
        )

        self._add_spacer(1)
        self._render_cover_disclaimer_table()

        self.doc.add_page_break()

    def _render_cover_disclaimer_table(self):
        """Render cover disclaimer table."""
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = STYLE.STYLE_TABLE_GRID

        cell = table.rows[0].cells[0]
        TableStyler.set_cell_background(cell, STYLE.LIGHT_GRAY_HEX)

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.clear()
        paragraph.paragraph_format.line_spacing = 1.15
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)

        run = paragraph.add_run(self._COVER_DISCLAIMER_TEXT)
        FontStyler.apply_run_style(
            run,
            font_name=STYLE.COVER_FONT,
            font_size=Pt(8.5),
            color=STYLE.DARK_GRAY,
        )

        self._apply_cover_disclaimer_border(table)

    @staticmethod
    def _apply_cover_disclaimer_border(table):
        """Apply thin border around cover disclaimer table."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")

        for border_name in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "4")
            border.set(qn("w:color"), STYLE.GRAY_BORDER_HEX)
            tblBorders.append(border)

        for border_name in ("insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _render_metadata_panel(
        self,
        report_date: str,
        analyst: str,
        company: str,
        sector: str,
        recipient: str,
        analysis_period: str = "",
        analysis_basis: str = "",
    ):
        """Render a compact two-column cover metadata panel."""
        rows = []

        if report_date:
            rows.append(("Report Date", report_date))
        if analysis_period:
            rows.append(("Analysis Period", analysis_period))
        if analysis_basis:
            rows.append(("Analysis Basis", analysis_basis))

        if self._is_meaningful_metadata_value(analyst, style_default="DCM Team 1"):
            rows.append(("Prepared By", analyst))
        if self._is_meaningful_metadata_value(company, style_default="Korea Development Bank"):
            rows.append(("Institution", company))
        if self._is_meaningful_metadata_value(sector, style_default="SECTOR"):
            rows.append(("Sector", sector))

        if recipient:
            rows.append(("Prepared For", recipient))

        if not rows:
            return

        table = self.doc.add_table(rows=len(rows), cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = STYLE.STYLE_TABLE_GRID

        for idx, (label, value) in enumerate(rows):
            label_cell = table.rows[idx].cells[0]
            value_cell = table.rows[idx].cells[1]

            TableStyler.set_cell_background(label_cell, STYLE.LIGHT_GRAY_HEX)

            label_para = label_cell.paragraphs[0]
            label_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            label_para.clear()
            label_run = label_para.add_run(label.upper())
            FontStyler.apply_run_style(
                label_run,
                font_name=STYLE.COVER_FONT,
                font_size=Pt(9.5),
                bold=True,
                color=STYLE.DARK_GRAY,
            )

            value_para = value_cell.paragraphs[0]
            value_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            value_para.clear()
            value_run = value_para.add_run(value)
            FontStyler.apply_run_style(
                value_run,
                font_name=STYLE.COVER_FONT,
                font_size=STYLE.BODY_SIZE,
                color=STYLE.DARK_GRAY,
            )

        TableStyler.set_table_borders(table)

    @staticmethod
    def _is_meaningful_metadata_value(value: str, style_default: str = "") -> bool:
        """Return True when a cover metadata value is worth showing to a user."""
        normalized = (value or "").strip()
        if not normalized:
            return False
        if style_default and normalized == style_default:
            return False
        if normalized.upper() in {"SECTOR", "N/A", "IB REPORT"}:
            return False
        return True

    def _add_spacer(self, lines: int):
        """Add empty paragraphs as vertical spacers."""
        for _ in range(lines):
            self.doc.add_paragraph()

    def _add_horizontal_rule(self, color_hex: str, size: str):
        """Add a clean horizontal rule using paragraph border."""
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), size)
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), color_hex)
        pBdr.append(bottom)
        pPr.append(pBdr)


class TOCRenderer:
    """Renders the Table of Contents"""

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, model: Optional[DocumentModel] = None):
        """Insert an auto-updating TOC plus an immediate preview outline."""
        heading = self.doc.add_heading("TABLE OF CONTENTS", level=1)
        self._apply_toc_heading_style(heading)

        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()

        # TOC field
        fldChar1 = OxmlElement("w:fldChar")
        fldChar1.set(qn("w:fldCharType"), "begin")

        instrText = OxmlElement("w:instrText")
        instrText.set(qn("xml:space"), "preserve")
        instrText.text = 'TOC \\o "1-4" \\h \\z \\u'

        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "separate")

        fldChar3 = OxmlElement("w:fldChar")
        fldChar3.set(qn("w:fldCharType"), "end")

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)

        if model is not None:
            self._render_preview_entries(model)

        self.doc.add_page_break()

    @staticmethod
    def _apply_toc_heading_style(paragraph) -> None:
        """Apply dedicated typography for the TOC title."""
        for run in paragraph.runs:
            FontStyler.apply_run_style(
                run,
                font_name=STYLE.TOC_FONT,
                font_size=STYLE.H1_SIZE,
                bold=True,
                color=STYLE.NAVY,
            )

    def _render_preview_entries(self, model: DocumentModel):
        """Render a static preview TOC so the document is useful before field update."""
        entries = []
        for element in model.elements:
            if element.element_type not in {
                ElementType.HEADING_1,
                ElementType.HEADING_2,
                ElementType.HEADING_3,
                ElementType.HEADING_4,
                ElementType.NUMBERED_HEADING,
            }:
                continue
            if not isinstance(element.content, Heading):
                continue

            level = max(1, min(element.content.level, 4))
            text = element.content.text.strip()
            if text:
                entries.append((level, text))

        if not entries:
            return

        for level, text in entries:
            paragraph = self.doc.add_paragraph()
            paragraph.style = STYLE.STYLE_IB_BODY
            paragraph.paragraph_format.left_indent = Inches(0.2 * (level - 1))
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2 if level <= 2 else 0)

            run = paragraph.add_run(text)
            FontStyler.apply_run_style(
                run,
                font_name=STYLE.TOC_FONT,
                font_size=STYLE.BODY_SIZE if level == 1 else (Pt(10) if level == 2 else STYLE.SMALL_SIZE),
                bold=level == 1,
                color=STYLE.DARK_GRAY if level <= 2 else STYLE.MEDIUM_GRAY,
            )


class HeadingRenderer:
    """Renders headings"""

    # Level → (font_size, color, bold)
    _STYLE_CONFIG = {
        1: (STYLE.H1_SIZE, STYLE.NAVY, True),
        2: (STYLE.H2_SIZE, STYLE.DARK_GRAY, True),
        3: (STYLE.H3_SIZE, STYLE.NAVY, True),
        4: (STYLE.H4_SIZE, STYLE.DARK_GRAY, True),
    }

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, heading: Heading):
        """Render a heading with appropriate level"""
        # Clean markdown bold markers from heading text
        clean_text = heading.text.replace("**", "").strip()
        clean_text = TextRenderer._cleanup(clean_text)

        # Clamp level to 1-4 (Word supports Heading 1-9, but we style 1-4)
        level = max(1, min(heading.level, 4))
        p = self.doc.add_heading(clean_text, level=level)

        size, color, bold = self._STYLE_CONFIG.get(level, (STYLE.BODY_SIZE, STYLE.DARK_GRAY, True))

        for run in p.runs:
            run.font.name = STYLE.HEADING_FONT
            run.font.size = size
            run.font.bold = bold
            run.font.color.rgb = color
            FontStyler.set_east_asian_font(run)


class ParagraphRenderer:
    """Renders paragraphs"""

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, paragraph: Paragraph):
        """Render a paragraph with text runs"""
        p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)
        if paragraph.runs:
            TextRenderer.render_runs(p, paragraph.runs)
        else:
            TextRenderer.render_text_with_bold(p, paragraph.text)


class ListRenderer:
    """Renders lists"""

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render_bullet(self, item: ListItem):
        """Render a bullet list item"""
        p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BULLET)
        indent = self._resolve_indent(item.indent_level)
        p.paragraph_format.left_indent = indent
        p.paragraph_format.first_line_indent = -STYLE.BULLET_INDENT

        # Bullet character
        bullet_run = p.add_run(f"{STYLE.BULLET_CHAR}  ")
        FontStyler.apply_run_style(
            bullet_run,
            font_name=STYLE.BODY_FONT,
            font_size=STYLE.BODY_SIZE,
            color=STYLE.NAVY,
        )

        # Content
        if item.runs:
            TextRenderer.render_runs(p, item.runs)
        else:
            TextRenderer.render_text_with_bold(p, item.text)

    def render_numbered(self, number: str, item: ListItem):
        """Render a numbered list item"""
        p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)
        indent = self._resolve_indent(item.indent_level)
        p.paragraph_format.left_indent = indent
        p.paragraph_format.first_line_indent = -STYLE.BULLET_INDENT

        # Number
        num_run = p.add_run(f"{number}. ")
        num_run.font.bold = True
        num_run.font.name = STYLE.BODY_FONT
        num_run.font.size = STYLE.BODY_SIZE
        FontStyler.set_east_asian_font(num_run)

        # Content
        if item.runs:
            TextRenderer.render_runs(p, item.runs)
        else:
            TextRenderer.render_text_with_bold(p, item.text)

    @staticmethod
    def _resolve_indent(indent_level: int) -> Inches:
        """Compute a bounded hanging indent for nested lists.

        Levels 0-3 use the full 0.25in step. Deeper levels compress to a
        smaller increment so very deep lists remain readable within page
        margins instead of drifting off the page.
        """
        normalized_level = max(0, indent_level)
        full_levels = min(normalized_level + 1, STYLE.FULL_LIST_INDENT_LEVELS)
        extra_levels = max(0, normalized_level + 1 - STYLE.FULL_LIST_INDENT_LEVELS)

        indent_inches = (
            full_levels * STYLE.BULLET_INDENT.inches
            + extra_levels * STYLE.DEEP_LIST_INDENT.inches
        )
        indent_inches = min(indent_inches, STYLE.MAX_LIST_INDENT.inches)
        return Inches(indent_inches)


class TableRenderer:
    """Renders tables with type-specific styling"""

    _EMUS_PER_INCH = 914400
    _COLUMN_SAMPLE_SIZE = 3
    _TOKEN_SPLIT_RE = re.compile(r"\s+")
    _NUMERIC_LIKE_RE = re.compile(
        r"^\(?[+-]?\d[\d,]*(?:\.\d+)?\)?\s*"
        r"(?:%|bp|bps|x|배|원|천원|만원|백만원|억원|억|조|주|개|명|건)?$",
        re.IGNORECASE,
    )
    _MIN_COLUMN_WIDTH_INCHES = 0.65
    _MIN_TEXT_COLUMN_WIDTH_INCHES = 1.15
    _MAX_NUMERIC_COLUMN_WIDTH_INCHES = 1.35
    _MAX_TEXT_COLUMN_SHARE = 0.55

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, table: Table):
        """Render a table"""
        if not table.rows or table.col_count == 0:
            return

        row_count = len(table.rows)
        col_count = table.col_count

        word_table = self.doc.add_table(rows=row_count, cols=col_count)
        word_table.style = STYLE.STYLE_TABLE_GRID
        column_kinds = self._infer_column_kinds(table)
        self._apply_column_widths(word_table, table, column_kinds)

        # Render header row
        if table.rows:
            self._render_header_row(word_table, table.rows[0], col_count)

        # Render data rows based on table type
        for r_idx, row in enumerate(table.rows[1:], 1):
            self._render_data_row(
                word_table,
                row,
                r_idx,
                col_count,
                table.table_type,
                column_kinds,
            )

        # Apply borders
        TableStyler.set_table_borders(word_table)

        # Spacer paragraph after table
        self.doc.add_paragraph()

    def _apply_column_widths(self, word_table, table: Table, column_kinds: List[str]) -> None:
        """Apply content-aware column widths for more readable report tables."""
        widths = self._estimate_column_widths(
            table,
            self._get_available_table_width_inches(),
            column_kinds=column_kinds,
        )
        if not widths:
            return

        word_table.autofit = False
        word_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for c_idx, width_inches in enumerate(widths):
            width = Inches(width_inches)
            word_table.columns[c_idx].width = width
            for cell in word_table.columns[c_idx].cells:
                cell.width = width

    def _estimate_column_widths(
        self,
        table: Table,
        available_width_inches: float,
        column_kinds: Optional[List[str]] = None,
    ) -> List[float]:
        """Estimate table column widths from content density and column semantics."""
        if table.col_count <= 0:
            return []

        resolved_column_kinds = column_kinds or self._infer_column_kinds(table)
        column_infos = [
            self._build_column_info(table, col_idx, resolved_column_kinds[col_idx])
            for col_idx in range(table.col_count)
        ]
        min_widths = [
            self._minimum_column_width(column_info["kind"]) for column_info in column_infos
        ]
        max_widths = [
            self._maximum_column_width(column_info["kind"], available_width_inches)
            for column_info in column_infos
        ]
        preferred = [column_info["score"] for column_info in column_infos]
        widths = self._fit_widths_to_available_space(
            preferred,
            min_widths,
            max_widths,
            available_width_inches,
        )

        return widths

    def _build_column_info(self, table: Table, col_idx: int, column_kind: str) -> Dict[str, float]:
        """Summarize the content profile of a single column."""
        texts = []

        for row in table.rows:
            if col_idx >= len(row.cells):
                continue

            cell = row.cells[col_idx]
            text = self._cell_display_text(cell, table.table_type).strip()
            if not text:
                continue

            texts.append(text)

        representative_length = self._representative_length(texts)
        longest_token = self._longest_token_length(texts)
        is_numeric_column = column_kind == "numeric"
        score = self._column_score(is_numeric_column, representative_length, longest_token)

        return {
            "kind": column_kind,
            "score": score,
        }

    def _infer_column_kinds(self, table: Table) -> List[str]:
        """Infer each column's semantic type from the first few data cells."""
        return [self._infer_column_kind(table, col_idx) for col_idx in range(table.col_count)]

    def _infer_column_kind(self, table: Table, col_idx: int) -> str:
        """Classify a column as text or numeric using the first 2-3 body cells."""
        sample_texts = []

        for row in table.rows[1:]:
            if col_idx >= len(row.cells):
                continue

            text = self._cell_display_text(row.cells[col_idx], table.table_type).strip()
            if not text:
                continue

            sample_texts.append(text)
            if len(sample_texts) >= self._COLUMN_SAMPLE_SIZE:
                break

        if not sample_texts:
            return "text"

        numeric_like_count = sum(1 for text in sample_texts if self._is_numeric_like(text))
        return "numeric" if numeric_like_count >= (len(sample_texts) + 1) // 2 else "text"

    @classmethod
    def _is_numeric_like(cls, text: str) -> bool:
        """Return True for financial/count strings that should align right."""
        normalized = text.strip()
        if not normalized or not any(char.isdigit() for char in normalized):
            return False
        return bool(cls._NUMERIC_LIKE_RE.match(normalized))

    @classmethod
    def _representative_length(cls, texts: List[str]) -> int:
        """Return a robust content length that ignores a few extreme outliers."""
        if not texts:
            return 0

        lengths = sorted(len(text) for text in texts)
        percentile_index = int(round((len(lengths) - 1) * 0.75))
        return lengths[percentile_index]

    @classmethod
    def _longest_token_length(cls, texts: List[str]) -> int:
        """Return the length of the longest unbroken token in the sample."""
        longest = 0
        for text in texts:
            tokens = [token for token in cls._TOKEN_SPLIT_RE.split(text) if token]
            if not tokens:
                continue
            longest = max(longest, max(len(token) for token in tokens))
        return longest

    @classmethod
    def _column_score(
        cls,
        is_numeric_column: bool,
        representative_length: int,
        longest_token: int,
    ) -> float:
        """Assign a width score to a column based on how hard it is to wrap cleanly."""
        if is_numeric_column:
            return 1.0 + min(representative_length, 12) * 0.05

        return (
            1.8
            + min(representative_length, 48) * 0.06
            + min(longest_token, 24) * 0.03
        )

    @classmethod
    def _minimum_column_width(cls, kind_marker: str) -> float:
        """Return a minimum readable width for a column kind."""
        if kind_marker == "numeric":
            return cls._MIN_COLUMN_WIDTH_INCHES
        return cls._MIN_TEXT_COLUMN_WIDTH_INCHES

    @classmethod
    def _maximum_column_width(cls, kind_marker: str, available_width_inches: float) -> float:
        """Return an upper width bound for a column kind."""
        if kind_marker == "numeric":
            return cls._MAX_NUMERIC_COLUMN_WIDTH_INCHES
        return max(
            cls._MIN_TEXT_COLUMN_WIDTH_INCHES,
            available_width_inches * cls._MAX_TEXT_COLUMN_SHARE,
        )

    @classmethod
    def _fit_widths_to_available_space(
        cls,
        preferred: List[float],
        min_widths: List[float],
        max_widths: List[float],
        available_width_inches: float,
    ) -> List[float]:
        """Fit preferred widths into the available width with min/max guards."""
        col_count = len(preferred)
        if col_count == 0:
            return []

        min_total = sum(min_widths)
        if min_total >= available_width_inches:
            scale = available_width_inches / min_total
            return [width * scale for width in min_widths]

        total_preferred = sum(preferred) or float(col_count)
        widths = [
            max(min_widths[idx], available_width_inches * preferred[idx] / total_preferred)
            for idx in range(col_count)
        ]

        widths = [min(widths[idx], max_widths[idx]) for idx in range(col_count)]
        widths = cls._rebalance_widths(widths, min_widths, max_widths, preferred, available_width_inches)
        return widths

    @classmethod
    def _rebalance_widths(
        cls,
        widths: List[float],
        min_widths: List[float],
        max_widths: List[float],
        preferred: List[float],
        available_width_inches: float,
    ) -> List[float]:
        """Rebalance widths after clamping so the table uses the full line width."""
        for _ in range(8):
            total_width = sum(widths)
            gap = available_width_inches - total_width

            if abs(gap) < 0.01:
                break

            if gap > 0:
                growable = [
                    idx for idx, width in enumerate(widths) if width < max_widths[idx] - 0.01
                ]
                if not growable:
                    break
                grow_weight = sum(preferred[idx] for idx in growable) or float(len(growable))
                for idx in growable:
                    share = preferred[idx] / grow_weight if grow_weight else 1.0 / len(growable)
                    widths[idx] = min(max_widths[idx], widths[idx] + gap * share)
            else:
                shrinkable = [
                    idx for idx, width in enumerate(widths) if width > min_widths[idx] + 0.01
                ]
                if not shrinkable:
                    break
                shrink_capacity = sum(widths[idx] - min_widths[idx] for idx in shrinkable)
                if shrink_capacity <= 0:
                    break
                overflow = -gap
                for idx in shrinkable:
                    capacity = widths[idx] - min_widths[idx]
                    share = capacity / shrink_capacity
                    widths[idx] = max(min_widths[idx], widths[idx] - overflow * share)

        return widths

    def _get_available_table_width_inches(self) -> float:
        """Return the horizontal space available for body tables."""
        if not self.doc.sections:
            return 6.0

        section = self.doc.sections[-1]
        available_width = int(section.page_width) - int(section.left_margin) - int(
            section.right_margin
        )
        available_width = max(available_width, int(Inches(3.0)))
        return float(available_width) / float(self._EMUS_PER_INCH)

    def _cell_display_text(self, cell_data: TableCell, table_type: TableType) -> str:
        """Return display text used for width estimation."""
        if table_type == TableType.FINANCIAL and cell_data.is_numeric:
            return self._format_financial_number(cell_data.content)
        return cell_data.content

    def _render_header_row(self, word_table, row: TableRow, col_count: int):
        """Render header row with Navy background"""
        word_cells = word_table.rows[0].cells

        for c_idx, cell_data in enumerate(row.cells):
            if c_idx >= col_count:
                break

            cell = word_cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            TableStyler.set_cell_background(cell, STYLE.NAVY_HEX)

            # Clear default paragraph and add styled run
            p = cell.paragraphs[0]
            self._configure_cell_paragraph(p)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.clear()

            # Parse header text: strip markdown markers but preserve structure
            # Headers are always bold+white, so we parse for content only
            if cell_data.runs:
                # Use structured runs but override style for header
                for run_data in cell_data.runs:
                    run = p.add_run(run_data.text)
                    FontStyler.apply_run_style(
                        run,
                        font_name=STYLE.HEADING_FONT,
                        font_size=STYLE.TABLE_HEADER_SIZE,
                        bold=True,
                        italic=run_data.italic,
                        color=STYLE.WHITE,
                    )
            else:
                clean_text = cell_data.content.replace("**", "").strip()
                run = p.add_run(clean_text)
                FontStyler.apply_run_style(
                    run,
                    font_name=STYLE.HEADING_FONT,
                    font_size=STYLE.TABLE_HEADER_SIZE,
                    bold=True,
                    color=STYLE.WHITE,
                )

    def _render_data_row(
        self,
        word_table,
        row: TableRow,
        row_idx: int,
        col_count: int,
        table_type: TableType,
        column_kinds: List[str],
    ):
        """Render a data row with type-specific styling"""
        word_cells = word_table.rows[row_idx].cells

        for c_idx, cell_data in enumerate(row.cells):
            if c_idx >= col_count:
                break

            cell = word_cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            self._configure_cell_paragraph(p)

            # ── Alignment ───────────────────────────────────────────────────
            if c_idx < len(column_kinds) and column_kinds[c_idx] == "numeric":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ── Format content (financial number formatting) ───────────────
            display_content = cell_data.content
            if table_type == TableType.FINANCIAL and cell_data.is_numeric:
                display_content = self._format_financial_number(cell_data.content)

            # ── Render content ──────────────────────────────────────────────
            if cell_data.runs:
                # Use structured runs for full formatting fidelity
                TextRenderer.render_runs(
                    p,
                    cell_data.runs,
                    font_name=STYLE.BODY_FONT,
                    font_size=STYLE.TABLE_BODY_SIZE,
                )
            else:
                TextRenderer.render_text_with_formatting(
                    p,
                    display_content,
                    font_name=STYLE.BODY_FONT,
                    font_size=STYLE.TABLE_BODY_SIZE,
                )

            # ── Type-specific styling ───────────────────────────────────────
            self._apply_type_styling(cell, p, cell_data, row_idx, table_type)

            # ── Alternating row colors (unless special styling applied) ─────
            if row_idx % 2 == 1 and not cell_data.is_base_case:
                TableStyler.set_cell_background(cell, STYLE.LIGHT_GRAY_HEX)

    @staticmethod
    def _configure_cell_paragraph(paragraph) -> None:
        """Normalize paragraph spacing inside table cells for a tighter, cleaner grid."""
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0

    @staticmethod
    def _format_financial_number(text: str) -> str:
        """
        Format financial numbers with thousand separators.

        Handles:
            - Plain numbers: 1234567 → 1,234,567
            - Percentages: 12.5% → 12.5%
            - Negative parentheses: (1234) → (1,234)
            - Korean 억/조 units preserved
            - Already formatted numbers: 1,234 → 1,234

        Args:
            text: Cell text content

        Returns:
            Formatted text with thousand separators
        """
        import re

        text = text.strip()

        # Skip if empty or non-numeric looking
        if not text or not any(c.isdigit() for c in text):
            return text

        # Skip if already has thousand separators
        if "," in text and re.search(r"\d{1,3}(,\d{3})+", text):
            return text

        # Handle negative in parentheses: (1234567) → (1,234,567)
        paren_match = re.match(r"^\((\d+(?:\.\d+)?)\)(.*)$", text)
        if paren_match:
            num_str = paren_match.group(1)
            suffix = paren_match.group(2)
            formatted = TableRenderer._add_thousand_sep(num_str)
            return f"({formatted}){suffix}"

        # Handle negative with minus: -1234567 → -1,234,567
        neg_match = re.match(r"^-(\d+(?:\.\d+)?)(.*)$", text)
        if neg_match:
            num_str = neg_match.group(1)
            suffix = neg_match.group(2)
            formatted = TableRenderer._add_thousand_sep(num_str)
            return f"-{formatted}{suffix}"

        # Handle positive numbers with optional suffix (%, 억, 조, 원, etc.)
        pos_match = re.match(r"^(\d+(?:\.\d+)?)(.*)$", text)
        if pos_match:
            num_str = pos_match.group(1)
            suffix = pos_match.group(2)
            formatted = TableRenderer._add_thousand_sep(num_str)
            return f"{formatted}{suffix}"

        return text

    @staticmethod
    def _add_thousand_sep(num_str: str) -> str:
        """Add thousand separators to a numeric string."""
        if "." in num_str:
            integer_part, decimal_part = num_str.split(".", 1)
            formatted_int = f"{int(integer_part):,}"
            return f"{formatted_int}.{decimal_part}"
        else:
            return f"{int(num_str):,}"

    def _apply_type_styling(
        self,
        cell,
        paragraph,
        cell_data: TableCell,
        row_idx: int,
        table_type: TableType,
    ):
        """Apply table-type specific styling"""
        if table_type == TableType.FINANCIAL:
            if cell_data.is_negative:
                for run in paragraph.runs:
                    run.font.color.rgb = STYLE.RED

        elif table_type == TableType.BEP_SENSITIVITY:
            if cell_data.is_base_case:
                TableStyler.set_cell_background(cell, STYLE.YELLOW_HEX)
                for run in paragraph.runs:
                    run.font.bold = True

        elif table_type == TableType.RISK_MATRIX:
            if cell_data.risk_level:
                color_map = {
                    "high": STYLE.RED,
                    "medium": STYLE.ORANGE,
                    "low": STYLE.GREEN,
                }
                color = color_map.get(cell_data.risk_level)
                if color:
                    for run in paragraph.runs:
                        run.font.color.rgb = color
                        run.font.bold = True


class CalloutRenderer:
    """
    Renders callout boxes (blockquotes) with professional IB styling.

    Supports different callout types with distinct visual styles:
        - KEY INSIGHT: Blue accent (default)
        - EXECUTIVE SUMMARY / 요약: Navy box with prominent styling
        - WARNING / 주의: Orange accent
        - NOTE / 참고: Gray accent
    """

    # Callout type configurations: (background_hex, border_color, title_color, icon)
    _CALLOUT_STYLES = {
        # Executive Summary / Important
        "EXECUTIVE SUMMARY": (STYLE.NAVY_HEX, STYLE.NAVY, STYLE.WHITE, "▶"),
        "요약": (STYLE.NAVY_HEX, STYLE.NAVY, STYLE.WHITE, "▶"),
        "핵심": (STYLE.NAVY_HEX, STYLE.NAVY, STYLE.WHITE, "▶"),
        "SUMMARY": (STYLE.NAVY_HEX, STYLE.NAVY, STYLE.WHITE, "▶"),
        # Insights
        "KEY INSIGHT": (STYLE.ACCENT_BLUE_HEX, STYLE.NAVY, STYLE.NAVY, "▌"),
        "시사점": (STYLE.ACCENT_BLUE_HEX, STYLE.NAVY, STYLE.NAVY, "▌"),
        "결론": (STYLE.ACCENT_BLUE_HEX, STYLE.NAVY, STYLE.NAVY, "▌"),
        # Warnings
        "WARNING": ("FFF3CD", STYLE.ORANGE, STYLE.ORANGE, "⚠"),
        "주의": ("FFF3CD", STYLE.ORANGE, STYLE.ORANGE, "⚠"),
        "RISK": ("FFF3CD", STYLE.ORANGE, STYLE.ORANGE, "⚠"),
        # Notes
        "NOTE": (STYLE.LIGHT_GRAY_HEX, STYLE.DARK_GRAY, STYLE.DARK_GRAY, "ℹ"),
        "참고": (STYLE.LIGHT_GRAY_HEX, STYLE.DARK_GRAY, STYLE.DARK_GRAY, "ℹ"),
    }

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, blockquote: Blockquote):
        """Render a callout box with style based on title."""
        title_upper = blockquote.title.upper()

        # Get style configuration (default to KEY INSIGHT style)
        style_config = self._CALLOUT_STYLES.get(
            title_upper,
            self._CALLOUT_STYLES.get(
                blockquote.title, (STYLE.ACCENT_BLUE_HEX, STYLE.NAVY, STYLE.NAVY, "▌")
            ),
        )
        bg_hex, border_color, title_color, icon = style_config

        # Handle RGBColor vs hex string for border
        if isinstance(border_color, str):
            border_hex = border_color
        else:
            border_hex = (
                f"{border_color[0]:02X}{border_color[1]:02X}{border_color[2]:02X}"
                if hasattr(border_color, "__getitem__")
                else STYLE.NAVY_HEX
            )

        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        TableStyler.set_cell_background(cell, bg_hex)

        # Title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"{icon} {blockquote.title}")

        # For dark backgrounds (like Executive Summary), use white text
        if bg_hex == STYLE.NAVY_HEX:
            FontStyler.apply_run_style(
                title_run,
                font_name=STYLE.HEADING_FONT,
                font_size=Pt(12),
                bold=True,
                color=STYLE.WHITE,
            )
        else:
            FontStyler.apply_run_style(
                title_run,
                font_size=Pt(11),
                bold=True,
                color=title_color if isinstance(title_color, RGBColor) else STYLE.NAVY,
            )

        # Content — with inline formatting support (**bold**, *italic*, ^super^)
        content_text = blockquote.text.strip()
        if content_text:
            content_para = cell.add_paragraph()

            # Determine text color based on background
            content_color = STYLE.WHITE if bg_hex == STYLE.NAVY_HEX else None

            TextRenderer.render_text_with_formatting(
                content_para,
                content_text,
                font_name=STYLE.BODY_FONT,
                font_size=STYLE.BODY_SIZE,
                default_color=content_color,
            )

        # Apply border styling
        self._apply_callout_border(
            table, border_hex if isinstance(border_hex, str) else STYLE.NAVY_HEX
        )

        # Spacer
        self.doc.add_paragraph()

    def render_executive_summary(self, title: str, bullet_points: list):
        """
        Render a professional Executive Summary box.

        Args:
            title: Box title (e.g., "Executive Summary")
            bullet_points: List of key points to display
        """
        table = self.doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        TableStyler.set_cell_background(cell, STYLE.NAVY_HEX)

        # Title
        title_para = cell.paragraphs[0]
        title_run = title_para.add_run(f"▶ {title}")
        FontStyler.apply_run_style(
            title_run,
            font_name=STYLE.HEADING_FONT,
            font_size=Pt(12),
            bold=True,
            color=STYLE.WHITE,
        )

        # Bullet points
        for point in bullet_points:
            point_para = cell.add_paragraph()
            bullet_run = point_para.add_run("  •  ")
            FontStyler.apply_run_style(
                bullet_run,
                font_name=STYLE.BODY_FONT,
                font_size=STYLE.BODY_SIZE,
                color=STYLE.WHITE,
            )
            text_run = point_para.add_run(point)
            FontStyler.apply_run_style(
                text_run,
                font_name=STYLE.BODY_FONT,
                font_size=STYLE.BODY_SIZE,
                color=STYLE.WHITE,
            )

        # Full border for executive summary
        self._apply_full_border(table, STYLE.NAVY_HEX)

        # Spacer
        self.doc.add_paragraph()

    @staticmethod
    def _apply_callout_border(table, border_hex: str):
        """Apply left accent border, hide other borders."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")

        left_border = OxmlElement("w:left")
        left_border.set(qn("w:val"), "single")
        left_border.set(qn("w:sz"), "32")
        left_border.set(qn("w:color"), border_hex)
        tblBorders.append(left_border)

        for border_name in ("top", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    @staticmethod
    def _apply_full_border(table, border_hex: str):
        """Apply full border around the callout box."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tblBorders = OxmlElement("w:tblBorders")

        for border_name in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "12")
            border.set(qn("w:color"), border_hex)
            tblBorders.append(border)

        tblPr.append(tblBorders)
        if tbl.tblPr is None:
            tbl.insert(0, tblPr)


class ImageRenderer:
    """
    Renders images in Word documents.

    Supports:
        - Base64 embedded images (data:image/... URI)
        - File path images (local files)
        - Fallback to placeholder if image cannot be loaded
    """

    # Maximum image width in inches (fits within typical IB report margins)
    MAX_WIDTH_INCHES: float = 5.5

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, image: Image):
        """
        Render an image to the Word document.

        Args:
            image: Image object with either base64_data or path

        Attempts to insert actual image; falls back to placeholder on failure.
        """
        import base64
        import os
        import tempfile
        from pathlib import Path

        inserted = False
        temp_file_path = None

        try:
            # ── Case 1: Base64 embedded image ──────────────────────────────────
            if image.base64_data:
                # Decode Base64 to bytes
                img_bytes = base64.b64decode(image.base64_data)

                # Determine file extension from MIME type
                ext = self._mime_to_extension(image.mime_type)

                # Write to temporary file
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False, mode="wb") as f:
                    f.write(img_bytes)
                    temp_file_path = f.name

                # Insert into document
                self._insert_image(temp_file_path, image.alt_text)
                inserted = True
                logger.debug("Inserted Base64 image: %s", image.alt_text)

            # ── Case 2: File path image ────────────────────────────────────────
            elif image.path:
                img_path = Path(image.path)

                # Handle relative paths
                if not img_path.is_absolute():
                    # Try relative to current working directory
                    if not img_path.exists():
                        logger.warning(
                            "Image file not found: %s — inserting placeholder",
                            image.path,
                        )
                    else:
                        self._insert_image(str(img_path), image.alt_text)
                        inserted = True
                        logger.debug("Inserted file image: %s", image.path)
                elif img_path.exists():
                    self._insert_image(str(img_path), image.alt_text)
                    inserted = True
                    logger.debug("Inserted file image: %s", image.path)
                else:
                    logger.warning(
                        "Image file not found: %s — inserting placeholder",
                        image.path,
                    )

        except Exception as e:
            logger.warning(
                "Failed to insert image '%s': %s — inserting placeholder",
                image.alt_text,
                e,
            )

        finally:
            # Clean up temporary file
            if temp_file_path:
                try:
                    os.unlink(temp_file_path)
                except OSError:
                    pass

        # ── Fallback: Placeholder ──────────────────────────────────────────────
        if not inserted:
            self._render_placeholder(image.alt_text)

    def _insert_image(self, file_path: str, alt_text: str):
        """
        Insert image file into document with proper sizing.

        Args:
            file_path: Path to image file
            alt_text: Alt text for caption
        """
        # Add image with max width constraint
        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        inline_shape = run.add_picture(file_path, width=Inches(self.MAX_WIDTH_INCHES))
        self._apply_alt_text(inline_shape, alt_text)

        # Add caption below image
        if alt_text:
            caption_para = self.doc.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(alt_text)
            FontStyler.apply_run_style(
                caption_run,
                font_size=STYLE.SMALL_SIZE,
                italic=True,
                color=STYLE.DARK_GRAY,
            )

        # Spacer
        self.doc.add_paragraph()

    @staticmethod
    def _apply_alt_text(inline_shape, alt_text: str):
        """Attach descriptive alt text to a Word inline image when available."""
        if not alt_text:
            return

        doc_pr = inline_shape._inline.docPr
        doc_pr.set("descr", alt_text)
        doc_pr.set("title", alt_text)

    def _render_placeholder(self, alt_text: str):
        """Render a placeholder when image cannot be loaded."""
        p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)
        run = p.add_run(f"[Image: {alt_text}]")
        FontStyler.apply_run_style(run, italic=True, color=STYLE.RED)

    @staticmethod
    def _mime_to_extension(mime_type: str) -> str:
        """Convert MIME type to file extension."""
        mime_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/webp": ".webp",
            "image/svg+xml": ".svg",
            "image/tiff": ".tiff",
        }
        return mime_map.get(mime_type.lower(), ".png")


class FootnoteRenderer:
    """Renders footnotes section"""

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, footnotes: dict):
        """Render footnotes natively when possible, else fall back to an ENDNOTES section."""
        if not footnotes:
            return

        if self._render_native(footnotes):
            return

        self._render_endnotes(footnotes)

    def _render_native(self, footnotes: dict) -> bool:
        """Replace superscript markers with native Word footnote references."""
        run_refs = self._collect_reference_runs(footnotes)
        if not run_refs:
            return False

        footnotes_part = NativeFootnotesPart.get_or_add(self.doc.part)
        referenced_numbers = sorted({number for _, number in run_refs})
        footnotes_part.set_footnotes({number: footnotes[number] for number in referenced_numbers})

        for run, number in run_refs:
            self._replace_with_native_reference(run, number)

        return True

    def _render_endnotes(self, footnotes: dict):
        """Render footnotes as a plain ENDNOTES section when native refs are unavailable."""
        self.doc.add_page_break()
        self.doc.add_heading("ENDNOTES", level=1)

        for num, text in sorted(footnotes.items()):
            p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)

            # Superscript number
            num_run = p.add_run(str(num))
            FontStyler.apply_run_style(
                num_run,
                font_size=STYLE.SMALL_SIZE,
                color=STYLE.NAVY,
                superscript=True,
            )

            # Space and text
            p.add_run(" ")
            text_run = p.add_run(text)
            FontStyler.apply_run_style(text_run, font_size=STYLE.SMALL_SIZE)

    def _collect_reference_runs(self, footnotes: dict) -> List[Tuple[object, int]]:
        """Find superscript numeric runs that correspond to known footnotes."""
        references: List[Tuple[object, int]] = []
        for paragraph in self._iter_document_paragraphs():
            for run in paragraph.runs:
                text = run.text.strip()
                if not text.isdigit():
                    continue
                if not run.font.superscript:
                    continue
                number = int(text)
                if number in footnotes:
                    references.append((run, number))
        return references

    def _iter_document_paragraphs(self):
        """Yield paragraphs from the document body and all top-level tables."""
        for paragraph in self.doc.paragraphs:
            yield paragraph
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        yield paragraph

    @staticmethod
    def _replace_with_native_reference(run, number: int) -> None:
        """Replace a rendered superscript number run with a native footnoteReference."""
        run_element = run._r
        r_pr = run_element.get_or_add_rPr()
        for child in list(run_element):
            if child is not r_pr:
                run_element.remove(child)

        r_style = OxmlElement("w:rStyle")
        r_style.set(qn("w:val"), "FootnoteReference")
        r_pr.append(r_style)

        footnote_ref = OxmlElement("w:footnoteReference")
        footnote_ref.set(qn("w:id"), str(number))
        run_element.append(footnote_ref)


class NativeFootnotesPart(XmlPart):
    """Minimal native footnotes part used for MD→Word footnote rendering."""

    _DEFAULT_XML = (
        b'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        b'<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        b'<w:footnote w:type="separator" w:id="-1"><w:p><w:r><w:separator/></w:r></w:p></w:footnote>'
        b'<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:r><w:continuationSeparator/></w:r></w:p></w:footnote>'
        b'</w:footnotes>'
    )

    @classmethod
    def default(cls, package) -> "NativeFootnotesPart":
        """Create a new empty footnotes part with required separator entries."""
        return cls(
            PackURI("/word/footnotes.xml"),
            CT.WML_FOOTNOTES,
            parse_xml(cls._DEFAULT_XML),
            package,
        )

    @classmethod
    def get_or_add(cls, document_part) -> "NativeFootnotesPart":
        """Get the existing footnotes part or create one and relate it to the document."""
        try:
            existing = document_part.part_related_by(RT.FOOTNOTES)
            return cast("NativeFootnotesPart", existing)
        except KeyError:
            footnotes_part = cls.default(document_part.package)
            document_part.relate_to(footnotes_part, RT.FOOTNOTES)
            return footnotes_part

    def set_footnotes(self, footnotes: Dict[int, str]) -> None:
        """Replace dynamic footnotes with the provided note mapping."""
        for footnote in list(self._element):
            if not str(footnote.tag).endswith("footnote"):
                continue
            footnote_id = footnote.get(qn("w:id"))
            if footnote_id not in {"-1", "0"}:
                self._element.remove(footnote)

        for number, text in sorted(footnotes.items()):
            self._element.append(self._build_footnote(number, text))

    @staticmethod
    def _build_footnote(number: int, text: str):
        """Build a single w:footnote element with simple paragraph content."""
        xml_space = "{http://www.w3.org/XML/1998/namespace}space"

        footnote = OxmlElement("w:footnote")
        footnote.set(qn("w:id"), str(number))

        paragraph = OxmlElement("w:p")
        paragraph_props = OxmlElement("w:pPr")
        paragraph_style = OxmlElement("w:pStyle")
        paragraph_style.set(qn("w:val"), "FootnoteText")
        paragraph_props.append(paragraph_style)
        paragraph.append(paragraph_props)

        ref_run = OxmlElement("w:r")
        ref_run_props = OxmlElement("w:rPr")
        ref_run_style = OxmlElement("w:rStyle")
        ref_run_style.set(qn("w:val"), "FootnoteReference")
        ref_run_props.append(ref_run_style)
        ref_run.append(ref_run_props)
        ref_marker = OxmlElement("w:footnoteRef")
        ref_run.append(ref_marker)
        paragraph.append(ref_run)

        text_run = OxmlElement("w:r")
        text_element = OxmlElement("w:t")
        text_element.set(xml_space, "preserve")
        text_element.text = f" {text}"
        text_run.append(text_element)
        paragraph.append(text_run)

        footnote.append(paragraph)
        return footnote


class DisclaimerRenderer:
    """Renders disclaimer page"""

    _SECTIONS = [
        (
            "면책 조항",
            "본 자료는 해당 문서에 최대한 정확하고 완전한 정보를 담고자 노력하였으나, "
            "오류와 중요정보의 누락이 있을 수 있으며, 정보의 정확성, 완전성 및 적정성을 "
            "보장하지 않습니다. 이 문서는 고객의 이해를 돕기 위하여 작성된 설명자료에 "
            "불과하므로, 고객은 각자의 책임으로 개별 계약서나 공시된 정보를 통하여 "
            "거래의 내용을 숙지하여야 합니다. 이 문서는 확정적인 거래조건을 구성하지 "
            "않으며 법적인 책임을 위한 근거자료로 사용될 수 없습니다.",
        ),
        (
            "저작권 및 비밀유지",
            "본 자료는 당행의 저작물로서 모든 저작권은 당행에게 있으며, 당행의 동의 없이 "
            "어떠한 경우에도 어떠한 형태로든 복제, 배포, 전송, 변경, 대여할 수 없습니다. "
            "당행의 요청 시에 즉시 반환, 파기하여 주시기 바랍니다. 본 자료는 상기 제한에 "
            "대하여 동의하는 조건으로 제공되며 동의하지 않으시는 경우에는 즉시 파기하여 "
            "주시기 바랍니다.",
        ),
        (
            "조건부 제공",
            "본 제안서의 내용은 현재의 시장상황 및 발행구조에 대한 기초정보에 근거한 것으로 "
            "유동화대상 자산 등 구조에 대한 변경이나 기타 중대한 사유 발생시 변경될 수 있으며, "
            "당행의 내부 여신심의위원회 승인을 조건으로 합니다.",
        ),
    ]

    def __init__(self, doc: DocxDocument):
        self.doc = doc

    def render(self, company: str):
        """Render standard IB disclaimer page"""
        self.doc.add_page_break()
        self.doc.add_heading("면책 조항", level=1)

        for title, content in self._SECTIONS:
            title_para = self.doc.add_paragraph()
            title_run = title_para.add_run(title)
            FontStyler.apply_run_style(
                title_run,
                font_name=STYLE.HEADING_FONT,
                font_size=STYLE.SMALL_SIZE,
                bold=True,
                color=STYLE.MEDIUM_GRAY,
            )
            title_para.paragraph_format.space_before = Pt(4)
            title_para.paragraph_format.space_after = Pt(3)

            for line in self._split_content_lines(content):
                content_para = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)
                content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                content_para.paragraph_format.line_spacing = 1.2
                content_para.paragraph_format.space_after = Pt(6)

                content_run = content_para.add_run(line)
                FontStyler.apply_run_style(
                    content_run,
                    font_name=STYLE.BODY_FONT,
                    font_size=STYLE.SMALL_SIZE,
                    color=STYLE.MEDIUM_GRAY,
                )

        # Copyright
        self.doc.add_paragraph()

        copy_para = self.doc.add_paragraph()
        copy_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        year = time.strftime("%Y")
        copy_run = copy_para.add_run(f"(c) {year} {company}. All rights reserved.")
        FontStyler.apply_run_style(
            copy_run,
            font_size=STYLE.SMALL_SIZE,
            italic=True,
            color=STYLE.MEDIUM_GRAY,
        )

    @staticmethod
    def _split_content_lines(content: str) -> List[str]:
        """Split disclaimer content into non-empty logical paragraphs."""
        lines = [line.strip() for line in content.split("\n")]
        return [line for line in lines if line]


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT RENDERER (ORCHESTRATOR)
# ═══════════════════════════════════════════════════════════════════════════════


class IBDocumentRenderer:
    """Main renderer that orchestrates all component renderers"""

    _TREE_PREFIX_RE = re.compile(r"^([\s│├└─]+)(.*)$")
    _TREE_VALUE_RE = re.compile(r"^(.*?)(\d+(?:\.\d+)?%|실질지배)(\s+──\s+)(.*)$")

    def __init__(self, separator_mode: str = "auto"):
        self.doc: DocxDocument = Document()
        self.separator_mode = separator_mode
        self.styler = DocumentStyler(self.doc)
        self.cover_renderer = CoverRenderer(self.doc)
        self.toc_renderer = TOCRenderer(self.doc)
        self.heading_renderer = HeadingRenderer(self.doc)
        self.paragraph_renderer = ParagraphRenderer(self.doc)
        self.list_renderer = ListRenderer(self.doc)
        self.table_renderer = TableRenderer(self.doc)
        self.callout_renderer = CalloutRenderer(self.doc)
        self.image_renderer = ImageRenderer(self.doc)
        self.footnote_renderer = FootnoteRenderer(self.doc)
        self.disclaimer_renderer = DisclaimerRenderer(self.doc)

    def render(self, model: DocumentModel) -> DocxDocument:
        """
        Render a DocumentModel to a Word Document.

        Args:
            model: The parsed document model

        Returns:
            The rendered Word Document
        """
        # Setup document
        self.styler.setup_document()
        self.styler.create_styles()

        # Setup header/footer with company name
        self.styler.setup_header_footer(
            company=model.metadata.company,
            confidential=True,
            show_page_numbers=True,
        )

        # Cover page
        self.cover_renderer.render(model.metadata)

        # Table of contents
        self.toc_renderer.render(model)

        # Render elements with error resilience
        for idx, element in enumerate(model.elements):
            try:
                self._render_element(element)
            except Exception as e:
                logger.warning(
                    "Failed to render element %d (type=%s): %s — skipping",
                    idx,
                    element.element_type.name,
                    e,
                )
                # Insert a visible marker in the document so the user knows
                p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)
                err_run = p.add_run(f"[Render Error: {element.element_type.name}]")
                FontStyler.apply_run_style(err_run, italic=True, color=STYLE.RED)

        # Footnotes/Endnotes
        if model.footnotes:
            self.footnote_renderer.render(model.footnotes)

        # Disclaimer
        self.disclaimer_renderer.render(model.metadata.company)

        return self.doc

    def _render_element(self, element: Element):
        """Render a single element based on its type"""

        etype = element.element_type

        if etype in (
            ElementType.HEADING_1,
            ElementType.HEADING_2,
            ElementType.HEADING_3,
            ElementType.HEADING_4,
            ElementType.NUMBERED_HEADING,
        ):
            self.heading_renderer.render(cast(Heading, element.content))

        elif etype == ElementType.PARAGRAPH:
            self.paragraph_renderer.render(cast(Paragraph, element.content))

        elif etype == ElementType.BULLET_LIST:
            self.list_renderer.render_bullet(cast(ListItem, element.content))

        elif etype == ElementType.NUMBERED_LIST:
            content = cast(Tuple[str, ListItem], element.content)
            number, item = content
            self.list_renderer.render_numbered(number, item)

        elif etype == ElementType.TABLE:
            self.table_renderer.render(cast(Table, element.content))

        elif etype == ElementType.BLOCKQUOTE:
            self.callout_renderer.render(cast(Blockquote, element.content))

        elif etype == ElementType.IMAGE:
            self.image_renderer.render(cast(Image, element.content))

        elif etype == ElementType.LATEX_BLOCK:
            self._render_latex_block(cast(LaTeXEquation, element.content))

        elif etype == ElementType.LATEX_INLINE:
            self._render_latex_inline(cast(LaTeXEquation, element.content))

        elif etype == ElementType.SEPARATOR:
            self._render_separator(element)

        elif etype == ElementType.CODE_BLOCK:
            self._render_code_block(cast(CodeBlock, element.content))

        elif etype == ElementType.DIAGRAM:
            from md_parser import Diagram
            from diagram_renderer import DiagramRenderer
            renderer = DiagramRenderer(self.doc, theme_colors={
                "navy": f"#{STYLE.NAVY_HEX}",
            })
            renderer.render(cast(Diagram, element.content))

        elif etype == ElementType.EMPTY:
            pass  # Intentionally skip empty elements

        else:
            logger.debug("Unhandled element type: %s", etype.name)

    def _render_separator(self, element: Element):
        """Render a separator as either a horizontal rule or a page break.

        Modes:
            - rule: always render a horizontal rule
            - page-break: always render a page break
            - auto: `## ---` becomes a page break, plain `---` stays a rule
        """
        if self._resolve_separator_mode(element) == "page-break":
            self.doc.add_page_break()
            return

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create a horizontal rule via paragraph bottom border
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")  # 0.75pt line
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), STYLE.GRAY_BORDER_HEX)
        pBdr.append(bottom)
        pPr.append(pBdr)

        # Minimal spacing
        pFmt = p.paragraph_format
        pFmt.space_before = Pt(6)
        pFmt.space_after = Pt(6)

    def _resolve_separator_mode(self, element: Element) -> str:
        """Resolve separator rendering mode for a specific element."""
        if self.separator_mode in {"rule", "page-break"}:
            return self.separator_mode

        raw_text = (element.raw_text or "").strip()
        if raw_text == "## ---":
            return "page-break"
        return "rule"

    def _render_code_block(self, code_block):
        """Render a fenced code block as a monospaced shaded block."""
        if not isinstance(code_block, CodeBlock):
            return

        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        TableStyler.set_cell_background(cell, "F8F9FA")
        self._style_code_block_table(table)

        lines = code_block.code.splitlines() or [""]
        for idx, line in enumerate(lines):
            paragraph = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
            paragraph.clear()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._render_code_block_line(paragraph, line)

        self.doc.add_paragraph()

    @staticmethod
    def _style_code_block_table(table) -> None:
        """Render code blocks as clean panels rather than visible grid tables."""
        tbl = table._tbl
        tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        tbl_borders = OxmlElement("w:tblBorders")

        for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "nil")
            tbl_borders.append(border)

        tbl_pr.append(tbl_borders)
        if tbl.tblPr is None:
            tbl.insert(0, tbl_pr)

        cell = table.rows[0].cells[0]
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_mar = OxmlElement("w:tcMar")
        for edge in ("top", "left", "bottom", "right"):
            margin = OxmlElement(f"w:{edge}")
            margin.set(qn("w:w"), "180" if edge in {"left", "right"} else "140")
            margin.set(qn("w:type"), "dxa")
            tc_mar.append(margin)
        tc_pr.append(tc_mar)

    def _render_code_block_line(self, paragraph, line: str) -> None:
        """Render one line of a code block with light semantic emphasis for tree diagrams."""
        if not line.strip():
            spacer = paragraph.add_run(" ")
            FontStyler.apply_run_style(
                spacer,
                font_name="Consolas",
                font_size=Pt(9.5),
                color=STYLE.MEDIUM_GRAY,
            )
            return

        prefix_match = self._TREE_PREFIX_RE.match(line)
        if prefix_match:
            prefix, remainder = prefix_match.groups()
        else:
            prefix, remainder = "", line

        if prefix:
            prefix_run = paragraph.add_run(prefix)
            FontStyler.apply_run_style(
                prefix_run,
                font_name="Consolas",
                font_size=Pt(9.5),
                color=STYLE.MEDIUM_GRAY,
            )

        value_match = self._TREE_VALUE_RE.match(remainder)
        if value_match:
            before, value, divider, after = value_match.groups()
            if before:
                before_run = paragraph.add_run(before)
                FontStyler.apply_run_style(
                    before_run,
                    font_name="Consolas",
                    font_size=Pt(9.5),
                    color=STYLE.DARK_GRAY,
                )
            value_run = paragraph.add_run(value)
            FontStyler.apply_run_style(
                value_run,
                font_name="Consolas",
                font_size=Pt(9.5),
                bold=True,
                color=STYLE.NAVY,
            )
            divider_run = paragraph.add_run(divider)
            FontStyler.apply_run_style(
                divider_run,
                font_name="Consolas",
                font_size=Pt(9.5),
                color=STYLE.MEDIUM_GRAY,
            )
            tail_run = paragraph.add_run(after)
            FontStyler.apply_run_style(
                tail_run,
                font_name="Consolas",
                font_size=Pt(9.5),
                color=STYLE.DARK_GRAY,
            )
            return

        is_root = "사업지주회사" in remainder
        main_run = paragraph.add_run(remainder)
        FontStyler.apply_run_style(
            main_run,
            font_name="Consolas",
            font_size=Pt(9.8 if is_root else 9.5),
            bold=is_root,
            color=STYLE.NAVY if is_root else STYLE.DARK_GRAY,
        )

    def _render_latex_block(self, latex_eq):
        """
        Render a LaTeX block equation as an image.

        Uses matplotlib to render the LaTeX expression to a PNG,
        then inserts it into the document centered.
        """
        from md_parser import LaTeXEquation

        if not isinstance(latex_eq, LaTeXEquation):
            logger.warning("Invalid LaTeX equation object")
            return

        image_path = LaTeXRenderer.render_to_image(latex_eq.expression)

        if image_path:
            try:
                # Insert centered equation image
                paragraph = self.doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_path, width=Inches(4.0))

                # Spacer
                self.doc.add_paragraph()

            except Exception as e:
                logger.warning("Failed to insert LaTeX image: %s", e)
                self._render_latex_fallback(latex_eq.expression)
            finally:
                # Clean up temp file
                try:
                    import os

                    os.unlink(image_path)
                except OSError:
                    pass
        else:
            self._render_latex_fallback(latex_eq.expression)

    def _render_latex_inline(self, latex_eq):
        """Render inline LaTeX (fallback to text for now)."""
        from md_parser import LaTeXEquation

        if isinstance(latex_eq, LaTeXEquation):
            self._render_latex_fallback(latex_eq.expression, inline=True)

    def _render_latex_fallback(self, expression: str, inline: bool = False):
        """Render LaTeX as styled text when image rendering fails."""
        p = self.doc.add_paragraph(style=STYLE.STYLE_IB_BODY)
        if not inline:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(f"[{expression}]")
        FontStyler.apply_run_style(
            run,
            font_name="Consolas",
            font_size=STYLE.BODY_SIZE,
            italic=True,
            color=STYLE.DARK_GRAY,
        )


# ═══════════════════════════════════════════════════════════════════════════════
# LATEX RENDERER (NEW)
# ═══════════════════════════════════════════════════════════════════════════════


class LaTeXRenderer:
    """
    Renders LaTeX expressions to PNG images using matplotlib.

    This class provides static methods to convert LaTeX math expressions
    into image files suitable for embedding in Word documents.

    Dependencies:
        - matplotlib (optional, graceful degradation if unavailable)
    """

    # Flag to track if matplotlib is available
    _matplotlib_available: Optional[bool] = None

    @classmethod
    def is_available(cls) -> bool:
        """Check if matplotlib is available for LaTeX rendering."""
        if cls._matplotlib_available is None:
            try:
                from importlib.util import find_spec

                cls._matplotlib_available = find_spec("matplotlib") is not None
            except ImportError:
                cls._matplotlib_available = False
            if not cls._matplotlib_available:
                logger.info(
                    "matplotlib not installed — LaTeX will render as text. "
                    "Install with: pip install matplotlib"
                )
        return cls._matplotlib_available

    @classmethod
    def render_to_image(
        cls,
        expression: str,
        fontsize: int = 14,
        dpi: int = 150,
    ) -> Optional[str]:
        """
        Render a LaTeX expression to a PNG image file.

        Args:
            expression: LaTeX math expression (without $ delimiters)
            fontsize: Font size for the equation
            dpi: Resolution of the output image

        Returns:
            Path to the temporary PNG file, or None if rendering fails
        """
        if not cls.is_available():
            return None

        try:
            import matplotlib

            matplotlib.use("Agg")  # Non-interactive backend
            import tempfile

            import matplotlib.pyplot as plt

            # Create figure with transparent background
            fig, ax = plt.subplots(figsize=(0.01, 0.01))
            fig.patch.set_alpha(0)
            ax.set_axis_off()

            # Render LaTeX text
            # Wrap in display math mode for better rendering
            latex_text = f"${expression}$"

            ax.text(
                0.5,
                0.5,
                latex_text,
                fontsize=fontsize,
                ha="center",
                va="center",
                transform=ax.transAxes,
            )

            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False, mode="wb") as f:
                temp_path = f.name

            fig.savefig(
                temp_path,
                dpi=dpi,
                bbox_inches="tight",
                pad_inches=0.1,
                transparent=False,
                facecolor="white",
            )
            plt.close(fig)

            logger.debug("Rendered LaTeX to: %s", temp_path)
            return temp_path

        except Exception as e:
            logger.warning("LaTeX rendering failed: %s", e)
            return None
